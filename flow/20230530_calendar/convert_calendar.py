import docx, re, easygui, csv
from datetime import datetime, timedelta
from docx.shared import RGBColor
from thefuzz import fuzz
import paschalia

RGB_RED = RGBColor(255, 0, 0)
RGB_BLACK = RGBColor(0, 0, 0)
RGB_GRAY = RGBColor(0x3c, 0x40, 0x43)


mode = easygui.choicebox('u - Юліанський, g - Григоріанський', 'Вибір календаря', ['u','g'])

if mode == 'u':
    docx_filename='Календар УГКЦ з читаннями 2025.docx'
    csvfilename='calendar25.txt'
elif mode == 'g':
    docx_filename='Календар УГКЦ з читаннями 2025 (григоріанський).docx'
    csvfilename='calendar25n.txt'

csv_reading_validation_filename = 'reading_validation.csv'




doc = docx.Document(docx_filename)



#for 2023
#previous_pascha =  datetime(2022,4,24)
#pascha = datetime(2023,4,16)
#for 2024
#previous_pascha =  datetime(2023,4,16)
#pascha = datetime(2024,5,5)
#for 2025
#previous_pascha = datetime(2024,5,5)
#pascha = datetime(2024,4,20)

#mode='u'

month_list = {'Січень':1,
              'Лютий':2,
              'Березень':3,
              'Квітень':4,
              'Травень':5,
              'Червень':6,
              'Липень':7,
              'Серпень':8,
              'Вересень':9,
              'Жовтень':10,
              'Листопад':11,
              'Грудень':12,
              'СІЧЕНЬ': 1,
                'ЛЮТИЙ': 2,
                'БЕРЕЗЕНЬ': 3,
                'КВІТЕНЬ': 4,
                'ТРАВЕНЬ': 5,
                'ЧЕРВЕНЬ': 6,
                'ЛИПЕНЬ': 7,
                'СЕРПЕНЬ': 8,
                'ВЕРЕСЕНЬ': 9,
                'ЖОВТЕНЬ': 10,
                'ЛИСТОПАД': 11,
                'ГРУДЕНЬ': 12}
month_list_string='('+'|'.join([x.lower() for x in month_list.keys()])+')'
day_list_string="(Понеділок|Вівторок|Середа|Четвер|П’ятниця|Пʼятниця|П'ятниця|Субота|Неділя)"
workday_list= [1,2,3,4,5]

symbol_dict={
    '🕀':'#',
    '🕁':'*',
    '🕂':'+',
    '+':'+',
    '🕃r':'@', #(red)
    '🕃b':'&' #(black)
    }
symbol_dict_w_space={v:k for v,k in symbol_dict.items()}
symbol_dict_w_space[' ']='-'
symbol_hierarchy={
    ' ':5,
    '🕀':0,
    '🕁':1,
    '🕂':2,
    '+':2,
    '+':2,
    '🕃r':3, #(red)
    '🕃b':4} #(black)
reverse_symbol_hierarchy = {v:k for k,v in symbol_hierarchy.items()}

symbol_list=['🕀','🕁','🕂','🕃']
notched_color=['🕃r','🕃b']

holidays=[]
with open('holiday_dictionary.csv', newline='',encoding='utf8') as csvfile:
    spamreader = csv.reader(csvfile, delimiter=';', quotechar='"')
    for row in spamreader:
        holidays.append(row)

def process_title(day_title):
    for row in holidays:
        #print(row)
        #RE_STRING=f"{row[0]}"
        #RE_STRING=f"(🕀|\+|🕁|🕂|+|🕃r|🕃b)?( )?"
        #print(RE_STRING)
        #print(day_title)
        re_result = re.search(f"{row[0]}",day_title)
        if re_result:
            #print("FOUND:", row, day_title)
            #print(re_result.groups())
            #symbol = re_result.group(1) if re_result.group(1) else ""
            #replace_string=f"<{row[1]}>{symbol}{row[0]}</{row[1]}>"
            day_title=day_title.replace(row[0],f"<{row[1]}>{row[0]}</{row[1]}>")
    re_result = re.search(r"(🕀|🕁|🕂|\+|🕃r|🕃b)( )?(<(?:span|strong|em|i)>)",day_title)
    if re_result:
        #print(re_result.groups())
        #print(day_title)
        day_title=day_title.replace("".join(re_result.groups()),f"{re_result.group(3)}{re_result.group(1)}")
        #print(day_title)
        #print("REPLACED:", day_title)

    for k,v in symbol_dict.items():
        day_title=day_title.replace(k,v)
        
    return day_title

def find_highest_symbol(title):
    if "Предсвяття Успення Богородиці" in title:
        print(title)
    cur_rank=5
    for i in range(len(title)):
        ch=title[i:i+2] if title[i] == '🕃' else title[i]
        if "Предсвяття Успення Богородиці" in title:
                print(ch, cur_rank)
        if ch in symbol_hierarchy and symbol_hierarchy[ch]<cur_rank:
            cur_rank=symbol_hierarchy[ch]
            if "Предсвяття Успення Богородиці" in title:
                print(ch)
    return reverse_symbol_hierarchy[cur_rank]
    
def get_lent_params(date,day_symbol,day_title,paschalia_dates):
    # return lent_color, lent_symbol
    
    is_holiday = True if day_symbol=="#" else False
    #Великі свята, в які піст
    if re.search("всесвітнє .оздвиження",day_title.lower()) or re.search("всемірне .оздвиження",day_title.lower()) or re.search("усікновення чесної",day_title.lower()):
        return 1, 1
    #Великий піст
    if (date>=paschalia_dates[1]["cheesefare_sunday"]+timedelta(days=1) and date<paschalia_dates[1]["pascha"]) and date.weekday()+1 in workday_list and not is_holiday:
        if date==paschalia_dates[1]["cheesefare_sunday"]+timedelta(days=1) or date==paschalia_dates[1]["pascha"]-timedelta(days=2):
            lent_symbol = 2
        elif date.weekday()+1 in [1,3,5]:
            lent_symbol = 1
        else:
            lent_symbol = 0
        return 1, lent_symbol
    #Петрівка
    if (date> paschalia_dates[1]["pentecost"]+timedelta(days=7) and date<datetime(date.year,6,29)) and date.weekday()+1 in workday_list and not is_holiday:
        return 1, 1 if date.weekday()+1 in [3,5] else 0
    #Спасівка
    if (date> datetime(date.year,6,29) and date<datetime(date.year,6,29)) and date.weekday()+1 in workday_list and not is_holiday:
        return 1, 1 if date.weekday()+1 in [3,5] else 0
    #Пилипівка
    if (date> datetime(date.year,11,15) and date<datetime(date.year,12,24)) and date.weekday()+1 in workday_list and not is_holiday:
        return 1,1 if date.weekday()+1 in [3,5] else 0
    #П'ятниця
    if date.weekday()+1 in [5] and not is_holiday:
        return 1,1
    #Якщо не вийшли з функції на одній з минулих перевірок, отже не піст
    return 0,0

#шукаємо блок з читанням

reading_indicator_list = ['Єв\. – ',
                            'Ап\. – ',
                            'Єв\. - ',
                            'Ап\. - ',
                            'Читання на Шостому часі:',
                            'Час Шостий:',
                            'На \d-му часі',
                            'Літ.:',
                            'Вечірня з Літургією св. Василія Великого.',
                            'На Літургії св. Василія Великого з вечірнею.',
                            'Літургія Передосвячених Дарів.',
                            'Літургія св. Йоана Золотоустого.',
                            'Літургія св. Івана Золотоустого.',
                            'Літургія св. Василія Великого.',
                            'Божественної Літургії не служимо.',
                            'На вечірні:',
                            'Вечірня:',
                            'Утр:',
                            'Утр.:',
                            'На освячення води',
                            'На вмиванні:',
                            'По вмиванні:',
                            'Ряд.:',
                            'Ап.:',
                            'Апп.:',
                            'Рівноап.:',
                            'Предтечі:',
                            'Новоліттю:',
                            'Св.:',
                            'Прп.:',
                            'Ісп.:',
                            'Свщнмч:',
                            'Свщмч:',
                            'Свщмч.:',
                            'Свщнмчч.:',
                            'Мчч.:',
                            'Влкмч.:',
                            'Отців:',
                            'Отцям:',
                            'Богородиці:',
                            'Собору:',
                            'Оновлення:',
                            'Анни:',
                            'Св. Миколая:',
                            'Свята:',
                            'Суботи:',
                            'Всім святим:',
                            'За упокій:',
                        ]

reading_indicator_string='('+'|'.join(reading_indicator_list)+')'
reading_indicator_string_for_beautify='('+'|'.join(reading_indicator_list[4:])+')'

holidays_godmother = [
    [9,8],#Різдво
    [10,1],#Покров
    [11,21],#Введення
    [12,9],#Зачаття
    [12,26]#Собор

]
def get_day_color(date,day_symbol):
    for d in holidays_godmother:
        if date == datetime(date.year,d[0],d[1]):
            return 3
    return 1 if date.weekday()+1==7 or day_symbol=="#" else 0


def beautify_reading(reading):
    reading = re.sub(reading_indicator_string_for_beautify,r'<i>\g<1></i>',reading)
    return reading

def compile_echo_gospel(date,paschalia_dates):

    if date>=paschalia_dates[1]['palm_sunday']:
        return "*"

    echo = paschalia.get_echos(date,paschalia_dates)
    gospel = paschalia.get_resurrection_gospel(date,paschalia_dates)

    if date>=paschalia_dates[0]['pascha']+timedelta(days=7) and date< paschalia_dates[0]['pascha']+timedelta(days=7*2): 
        echo=None

    lst=[]
    if echo:
        echo_str = f"Гл. {echo}"
        lst.append(echo_str)
    if gospel and date.weekday()+1 == 7:
        gospel_str = f"Єв. {gospel}"
        lst.append(gospel_str)
    return '. '.join(lst)



def generate_row(cur_year,cur_month,cur_day,day_title,day_readings,glas):
    
    month=month_list[cur_month.capitalize()]
    try:
        date=datetime.strptime(f'{cur_year}-{month:02}-{cur_day:02}','%Y-%m-%d')
    except ValueError as e:
        print(f'ERROR with {cur_year}-{month:02}-{cur_day:02}')
        raise e
    paschalia_dates = paschalia.get_prev_next_pascha(date,mode)
    day_symbol= symbol_dict_w_space[find_highest_symbol(day_title)]
    day_color = get_day_color(date,day_symbol)
    #week_string=get_week(date,day_title,day_symbol)
    week_code=paschalia.get_week_code(date,day_title)
    #day_title=day_title.replace("🕃b","🕃")
    #lent_icon = get_lent_icon(date,day_symbol, day_title)
    lent_color,lent_icon = get_lent_params(date,day_symbol,day_title,paschalia_dates)

    
    if date.weekday()+1 !=7 and glas and glas!="*":
        #???
        #print(glas, date)
        glas = re.search(r"(Глас \d)",glas).group(1)

    glas = compile_echo_gospel(date,paschalia_dates)
    '''
    try:
        weekday_no=date.weekday()+1
    except ValueError as e:
        print(cur_year,cur_month,cur_day)
        print(f'{cur_year}-{month:02}-{cur_day:02}')
        print(f'ValueError, p={i}: ',p.text)
        raise e
    '''
    
    '''
    print("BEFORE:",day_title)
    for k,v in symbol_dict.items():
        day_title=day_title.replace(k,v)
    print("AFTER:",day_title)

    print(symbol_dict.items())
    raise Exception
    '''
    
    '''
    Структура рядка
    v1 - Дата
    v2 - № та тип тижня
    v3 - день тижня (1..7)
    v4 - тип свята [0,1,3]
    5 - колір посту [0,1]
    6 - рибка [0.1,2]
    '''

    
    row=[f'{cur_year}{month:02}{cur_day:02}',
         week_code,
         date.weekday()+1,
         day_color,
         lent_color,
         lent_icon,
         day_symbol,
         process_title(day_title),
         beautify_reading(day_readings),
         glas]
    '''
    row=[f'{cur_year}{month:02}{cur_day:02}',
         week_code,
         date.weekday()+1,
         beautify_reading(day_readings),
    ]
    '''
    return row




def check_day_qty(rows):
    start_date = datetime(int(rows[0][0][:4]),int(rows[0][0][4:6]), int(rows[0][0][6:]))
    end_date = datetime(int(rows[-1][0][:4]),int(rows[-1][0][4:6]), int(rows[-1][0][6:]))
    delta = end_date - start_date
    num_days = delta.days+1
    print("Кількість днів між датами", rows[0][0],rows[-1][0],":",num_days)
    if len(rows)!=num_days:
        day_matrix=[]
        date = start_date
        while date<=end_date:
            i=0
            #print(date)
            datestr=f'{date.year}{date.month:02}{date.day:02}'
            #print(datestr)
            while i<len(rows) and datestr!=rows[i][0]:
                #print(i, datestr, rows[i][0])
                #raise Exception
                i+=1
            #print(i)
            #print(lenrows])
            if i<len(rows):
                day_matrix.append([f'{date.year}{date.month:02}{date.day:02}',True])
            else:
                day_matrix.append([f'{date.year}{date.month:02}{date.day:02}',False])
            date += timedelta(days=1)

        
        print(f"Не співпадає кількість днів: {len(rows)} проти {num_days}")
        print("Відсутні дні:")
        for r in day_matrix:
            if not r[1]:
                print(r)
        #raise Exception



cur_year='2025'
cur_month=''
cur_day=''
day_title=''
day_readings=''
flag_title=False
flag_reading=False
glas=''
list_of_symbols=[]
first=True


i=0
for p in doc.paragraphs:
    for r in p.runs:
        if '🕃' in r.text:
            if r.font.color.rgb == RGB_RED:
                r.text=r.text.replace('🕃','🕃r')
            elif r.font.color.rgb == RGB_BLACK or r.font.color.rgb == RGB_GRAY or not r.font.color.rgb:
                r.text=r.text.replace('🕃','🕃b')
            else:
                print('Warning: paragraph with funky colors:')
                print(i,p.text)
                #for r in doc.paragraphs[i].runs:
               #     print(r.font.color,r.text)
                print(r.font.color.rgb,r.text)
    i+=1            
i=-1
rows=[]
reading_validation_matrix=[]
cur_month = None
for p in doc.paragraphs:
    re_result=None
    i+=1

    #if i < 4:
    #    continue
    #if i>26:
    #    break
    
    #print(i,p.text)
    #print(f'f_t={flag_title}, f_r={flag_reading}')
    
    #пропускаємо пусте
    if p.text=='':
        continue
    
    #місяць, рік
    re_result=re.search(f'{month_list_string}( \((\d+)\))?',p.text.lower())
    if re_result:
        #print(re_result.groups(),first)
        if not first:
            #print("pfqikb c.lb")
            rows.append(generate_row(cur_year,cur_month,cur_day,day_title,day_readings,glas))
            flag_reading=False
            day_readings=""
            day_title=""
        else:
            first=False
        #print(re_result.groups())
        if re_result.group(3):
            cur_year = re_result.group(3)
        cur_month=re_result.group(1)
        #print(cur_year, cur_month)
        continue
    
    #пропускаємо вступ з календарем посту.
    #починаємо працювати від першого місяця
    if not cur_month:
        continue


    #День, початок тайтлу
    re_result=re.search(r'^(\d+)( )?(.*)',p.text)
    if re_result:
        #print("Day found",i, p.text)
        cur_day=int(re_result.group(1))
        #print("cur_day =", cur_day)
        #if cur_day=='40':
        #    print(p.text)
        day_title=re_result.group(3)
        flag_title=True

        #Глас. Воскресне Євангеліє
        re_result=re.search(r'(Глас \d\. Єв\. \d{1,2}\.)',p.text)
        if re_result:
            glas=re_result.group(1)
            day_title=day_title.replace(glas,"")
            re_result=None
        #else:
        #    glas='*'
        re_result=re.search(r'Гл. (\d\. Єв\. \d{1,2}\.)',p.text)
        if re_result:
            glas = "Глас "+re_result.group(1)

        #print(cur_day, day_title, glas)
        #print("c")
        continue
    
    
    

    
    re_result=re.search(reading_indicator_string,p.text)
    if re_result:
        #print("Found readings",i,p.text)
        flag_reading=True
        flag_title=False
        if not day_readings:
            day_readings=p.text
        else:
            day_readings = day_readings + '<br>' + p.text
        continue

    #якщо в строці лише день - закриваємо попередній "день"
    re_result=re.search(f'^{day_list_string}$',p.text)
    if re_result and day_title: #and flag_reading:
        #print('APPENDING ROWS!',flag_reading)
        rows.append(generate_row(cur_year,cur_month,cur_day,day_title,day_readings,glas))
        flag_reading=False
        day_readings=""
        day_title=""
        continue
    elif re_result and not flag_reading:
        #print("Пропускаємо ", p.text)
        continue

    #інакше - ввжаємо, що продовжується тайтл.
    if flag_title:
        day_title+="<br>"+p.text
    else:
        print(f'Щось не так {i}:\n',p.text, f'(дата {cur_month}-{cur_day})')

#цикл не обробить закінчення останнього дня, викликаємо явно
rows.append(generate_row(cur_year,cur_month,cur_day,day_title,day_readings,glas))



def read_csv(path):
    lst=[]
    with open(path, newline='', encoding='utf-8') as csvfile:
        csvreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in csvreader:
            lst+=[row]
    return lst

apostol_list = read_csv('apostol.csv')
evanhelie_list= read_csv('evanhelie.csv')





def break_readings_segment(str):
    lst = []
    res = re.findall(r'(Ап\. (?:–|-) .*?)(?:<br>|<i>|\n|\||<sup>|$)',str)
    res2=[]
    for item in res:
        if ' Єв.' in item:
            #print('found', res)
            res2+=item.split(' Єв. – ')
            res2[-1]='Єв. – '+res2[-1]
            #res.remove(item)
            #print('edited', res)
        else:
            res2.append(item)
    lst+=res2
    res = re.findall(r'(Єв\. (?:–|-) .*?)(?:<br>|<i>|\n|\|<sup>|$)',str)
    lst+=res
    return lst

def process_reading(str,reading_list):
    found=False
    ratio_max=0
    item_found=None
    for item in reading_list:
        ratio=fuzz.token_sort_ratio(str,item[0])
        if ratio==100:
            found=True
            item_found=item
            #return  f'#ap{item[2]:0>3}'+f'{item[3]}'
            return [item_found[0],item_found[4]]
                        
        if ratio>ratio_max:
            ratio_max=ratio
            item_found=item
    
    if ratio_max>=85 and ratio_max<100:
        found=True
        print(f"WARNING. ratio: {ratio_max}")
        print(str, item_found[0])
        print('Input:',str)  
        #return f'#ap{item_found[2]:0>3}{item_found[3]}',item_found[4]
        return [item_found[0],item_found[4]]

    if not found:
        print(f"ERROR. ratio: {ratio_max}")
        print(str)
        print('ratio:',ratio_max)
        return 'crit_err',''
        #raise Exception

def clean_half_words(s):
    s=s.replace(' (від половини)','').replace(' (від полов.)','')
    return s


def validate_readings(rows):
    for row in rows:
        row.append(break_readings_segment(row[8]))

    for row in rows:
        row.append([])
        for item in row[10]:
            reading_details=None

            if item.startswith('Ап.'):
                reading_details =  process_reading(clean_half_words(item[6:]),apostol_list)
            elif item.startswith('Єв.'):
                reading_details =  process_reading(clean_half_words(item[6:]),evanhelie_list)

            if reading_details:
                row[11].append(reading_details[0])
                reading_validation_matrix.append([row[0],row[1],row[2],row[7],item,reading_details[1]])
            else:
                row[11].append('err')
            
            
validate_readings(rows)


#with open('c1test.txt','a',encoding='utf8') as f:
csvfile = open(csvfilename,'w',newline='',encoding='utf8')
spamwriter=csv.writer(csvfile,delimiter='|',quotechar='"', quoting=csv.QUOTE_MINIMAL)

#for row in rows:
#    row = row[0:1]+row[8:9]+row[10:]

for row in rows:
    try:
        #print(row[7])
        spamwriter.writerow(row[:10])
        #spamwriter.writerow(row)
    except Exception as e:
        print(f'Error, p={i}: ',p.text)
        print(row)
        raise e
    
csvfile.close()

with open(csv_reading_validation_filename,'w',newline='',encoding='utf8') as csvfile:
    spamwriter=csv.writer(csvfile,delimiter='|',quotechar='"', quoting=csv.QUOTE_MINIMAL)
    spamwriter.writerows(reading_validation_matrix)

with open('og_header.html','w',newline='',encoding='utf8') as csvfile:
    spamwriter=csv.writer(csvfile,delimiter='|',quotechar='"', quoting=csv.QUOTE_MINIMAL)
    spamwriter.writerows([[x[1],x[7]] for x in rows])




#перевіряємо кількість днів
#print(len(rows), rows[0])
check_day_qty(rows)

print("Done!")