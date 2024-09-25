import docx,re,csv, easygui,paschalia
from glob import glob
from docx.shared import RGBColor
from datetime import datetime
RGB_RED = RGBColor(0xff, 0x44, 0x00)

#docx_filename='2023-07-Літургії-Гр.docx'
#mode='g'
#cur_month='Липень'
month_no = datetime.now().month+1 if datetime.now().month!=12 else 1
year_no = datetime.now().year if datetime.now().month!=12 else datetime.now().year+1

month_no=10
print("WARNING. Month_no OVERRIDE", month_no)

mode = easygui.choicebox('u - Юліанський, g - Григоріанський', 'Вибір календаря', ['u','g'])

#dirs = easygui.diropenbox()
#paths = glob(f'{dirs}\\*.docx', recursive=True)
#docx_filename = paths[0]

docx_filename = easygui.fileopenbox(
    title="Select a .docx file",
    filetypes=["*.docx"],
    default="*.docx"
)

if mode == 'u':
    mode_suffix='Юл'
elif mode == 'g':
    mode_suffix='Гр'

csv_filename='Читання'+mode_suffix+'.csv'
dismissal_csv_filename='Відпусти'+mode_suffix+'.csv'


doc = docx.Document(docx_filename) 


month_list = {'Січень':1,
              'Лютий':2,
              'Березень':3,
              'Квітень':4,
              'Травень':5,
              'Червень':6,
              'Липень':7,
              'Серпень':8,
              'Вересень':9,
              'Жовтень':10,
              'Листопад':11,
              'Грудень':12}
month_list_reversed = {v:k for k,v in month_list.items()}

#cur_month_no = int(re.search('\d{4}.(\d{2})',dirs).group(1))
#cur_month= month_list_reversed[int(re.search('\d{4}.(\d{2})',dirs).group(1))]
#month_no = datetime.now().month+1 if datetime.now().month!=12 else 1
#month_no -=1
cur_month= month_list_reversed[month_no]

month_list_string='('+'|'.join(month_list.keys())+')'
day_list_string="(Понеділок|Вівторок|Середа|Четвер|П’ятниця|П'ятниця|Субота|Неділя)"

reading_type_indices={'apostol':5,
                      'evanhelie':7,
                      'vidpust':8}
'''
0 3
1 вв
2 
3 
4 
5 
6 
7 Ряд.:
8 Св.:
9 #ap249a
10 #ap334
11 #lk034
12 #iv036
'''

reading_matrix={}
with open(csv_filename, newline='', encoding='utf-8') as csvfile:
    spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
    for row in spamreader:
        if row[0]==cur_month:

            reading_matrix[int(row[1].split('.')[0])]=row[2:]
            #print(int(row[1].split('.')[0]), row[2:])

dismissal_matrix={}
with open(dismissal_csv_filename, newline='', encoding='utf-8') as csvfile:
    spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
    for row in spamreader:
        if row[0]==cur_month:
            dismissal_matrix[int(row[1].split('.')[0])]=row[2:]


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def delete_run(run):
    p = run._element
    p.getparent().remove(p)
    p._p = p._element = None


#9-apostol, 11-evanhelie
def get_readings(date,index):
    #print(date)
    reading_list=[reading_matrix[date][3]+' '+reading_matrix[date][index] if reading_matrix[date][3] else reading_matrix[date][index]]
    if reading_matrix[date][index+1]:
        reading_list.append(reading_matrix[date][3+1]+' '+reading_matrix[date][index+1])
    return reading_list
def insert_readings_day(p,date,index):
    for r in get_readings(date,index):
        p.insert_paragraph_before(r)
   
def insert_dismissal(p,date,index):
    p.insert_paragraph_before(dismissal_matrix[date][index])


for p in doc.paragraphs:
    for r in p.runs:
        if '\xa0' in r.text:
            r.text=r.text.replace('\xa0',' ')
        if '  ' in r.text:
            r.text=r.text.replace('  ',' ')
    if p.text=="":
        delete_paragraph(p)
        
            

i=0
for p in doc.paragraphs:
    if p.text=="":
        print('E', i,p.text)
    if '  ' in p.text:
        print('DS',i,p.text)
    if '\n' in p.text:
        print('N',i,p.text)
    i+=1

#тестовий блок для пошуку подвійних спейсів на границі ранів
'''
for p in doc.paragraphs:
    previous_run=None
    for r in p.runs:
        if not previous_run:
            previous_run=r
            continue
        else:
            if previous_run.text.endswith(' ') and r.text.startswith(' '):
                print('found!!!')
                print('1',previous_run.text)
                print('2',r.text)
                previous_run=r
            else:
                previous_run=r
'''   

#вставляємо читання та відпусти 
iz='святого отця нашого Йоана Золотоустого, архиєпископа Константинограда,'
vv='святого отця нашого Василія Великого, архиєпископа Кесарії Кападокійської,'

placeholder=False
cur_date=None
reading_type=None
liturgy=None
paschalia_dates = paschalia.get_prev_next_pascha(datetime(year_no, month_no,1), mode)

for p in doc.paragraphs:
    re_result=re.search(f'^{month_list_string} (\d+)',p.text)
    if re_result:
        cur_date = int(re_result.group(2))

    
    re_result=re.search(r'liturgia=(.*?)(?:\s|/)',p.text)
    if re_result:
        liturgy = re_result.group(1)

    #if re_result:
        #[день, місяць]
    #    cur_date=int(re_result.group(2))
    if re.search('<(apostol|evanhelie|vidpust)>',p.text):
        placeholder=reading_type=re.search('<(apostol|evanhelie|vidpust)>',p.text).group(1)
        continue
    if placeholder:
        if re.search(f'</{reading_type}>',p.text):
            if placeholder=='vidpust':
                print(cur_date, 'dismissal', liturgy)
                if liturgy=='lvv':
                    dismissal_matrix[cur_date][reading_type_indices[reading_type]]=dismissal_matrix[cur_date][reading_type_indices[reading_type]].replace(iz,vv)
                insert_dismissal(p,cur_date,reading_type_indices[reading_type])
            elif placeholder in ['apostol','evanhelie']:
                #print(placeholder)
                #print(cur_date)
                if not cur_date:
                    raise Exception
                insert_readings_day(p,cur_date,reading_type_indices[reading_type])
            placeholder = False
        else:
            delete_paragraph(p)

    #застаріле:
    '''
    if re.search('святий за календарем',p.text):
        print(cur_date, 'dismissal')
        
        p.insert_paragraph_before(dismissal_matrix[cur_date][2])
        delete_paragraph(p)
    '''

# вставляємо святих після дати        
#for p in doc.paragraphs:
#???
    


# робимо всі "Священик" червоними та italic
BLACK='b'
RED='r'
def add_text(p,text, color=BLACK):
    r = p.add_run(text)
    r.font.name='Times New Roman'
    r.font.size=152400
    if color == RED:
        r.font.color.rgb = RGB_RED
        r.italic = True

for p in doc.paragraphs:
    if re.search(r'Священ{1,2}ик:',p.text):
        #print('found dissmissal',p.text[:40])
        re_result=re.search(r'^(Священ{1,2}ик:)( .+?)(якого є храм)(.*?)$',p.text)
        p_bak=p.text
        p.clear()

        try:
            if re_result:
                add_text(p,re_result.group(1),color=RED)
                add_text(p,re_result.group(2))
                add_text(p,re_result.group(3),color=RED)
                add_text(p,re_result.group(4))
            else:
                re_result=re.search(r'^(Священ{1,2}ик:)(.*?)$',p_bak)
                add_text(p,re_result.group(1),color=RED)
                add_text(p,re_result.group(2))
        except:
            print(p_bak)
            raise
            
        
doc.save(docx_filename)
print("Finished successfully!")
