import docx, re, easygui
from datetime import datetime

month_dic = {'Січень':1,
              'Лютий':2,
              'Березень':3,
              'Квітень':4,
              'Травень':5,
              'Червень':6,
              'Липень':7,
              'Серпень':8,
              'Вересень':9,
              'Жовтень':10,
              'Листопад':11,
              'Грудень':12}
month_dic_string='('+'|'.join(month_dic.keys())+')'
month_dic_reversed = {v:k for k,v in month_dic.items()}

day_list_string="(Понеділок|Вівторок|Середа|Четвер|П’ятниця|П'ятниця|Субота|Неділя)"

list_of_text_templates=['liz','lvv','liz_pascha', 'liz_voznesinnia','liz_subota_po_voznesinni']

list_of_tags_mandatory=['antifon1',
                        'antifon2',
                        'antifon3',
                        'vhidne',
                        'tropari',
                        'prokimen',
                        'apostol',
                        'aleluia',
                        'evanhelie',
                        'prychasnii',
                        'vidpust']
list_of_tags_optional=['trysviate',
                       'dostoino']

def pretty(d, indent=0):
   for key, value in d.items():
      print('\t' * indent + str(key))
      if isinstance(value, dict):
         pretty(value, indent+1)
      else:
         print('\t' * (indent+1) + str(value))

def checkLiturgyIntegrity(docx_filename):
   doc = docx.Document(docx_filename)
   history={}

   #прибираємо leading / trailing пробіли
   for p in doc.paragraphs:
      if re.search(r'^.*?\s+$',p.text):
         p.text = re.sub(r'^(.*?)\s+$',r'\g<1>',p.text)
      if re.search(r'^\s+.*?$',p.text):
         p.text = re.sub(r'^\s+(.*?)$',r'\g<1>',p.text)

   #for i in range(100):
   #   print(doc.paragraphs[i].text)
   cur_date=None
   for p in doc.paragraphs:
      #if mode == 'u':
      #if mode:
      #   re_result=re.search(f'^{month_dic_string} '+r'(\d+)',p.text)
      #   if re_result:
      #      cur_date = int(re_result.group(2))

      re_result=re.search(f'^{month_dic_string} '+r'(\d+)',p.text)
      if re_result:
         cur_date = int(re_result.group(2))

      if cur_date and not cur_date in history:
         history[cur_date]={'integrity':'?'}

      
      if cur_date:
         ustav = re.search(r'^<ustav.*?liturgia=(\w+).*$', p.text)
         if ustav:
               if ustav.group(1) in list_of_text_templates:
                  history[cur_date]['ustav']='OK'
               else:
                  history[cur_date]['ustav']=ustav.group(1)
                  
         for item in list_of_tags_mandatory+list_of_tags_optional:
               #if item=='vidpust':
               #    print(p.text)
               #if re.search('^\s?<'+item+'>\s?$',p.text):
               if re.search('^<'+item+'>$',p.text):
                  if not item in history[cur_date]:
                     history[cur_date][item]='<'
                  else:
                     history[cur_date][item]+='<'
               #if re.search('^\s?</'+item+'>\s?$',p.text):
               if re.search('^</'+item+'>$',p.text):
                  if not item in history[cur_date]:
                     history[cur_date][item]='>'
                  else:
                     history[cur_date][item]+='>'

   print(f'Файл: {docx_filename}\n')

   optional_tag_summary={}
   for k, day in history.items():
      #перевірка правильності <ustav>
      if 'ustav' in day:
         if day['ustav']!='OK':
            print(f'День {k}: <ustav>, неправильний шаблон ('+day['ustav']+')')
            day['integrity']='NOK'

         #перевірка правильності основного набору міток
         for item in list_of_tags_mandatory:
            if not item in day:
               print(f'День {k}: пропущено секцію <{item}>')
               day['integrity']='NOK'
            else:
               if not '<' in day[item]:
                  print(f'День {k}: пропущено відкриваючий <{item}>')
                  day['integrity']='NOK'
               if not '>' in day[item]:
                  print(f'День {k}: пропущено закриваючий </{item}>')
                  day['integrity']='NOK'
                     
         #підсумки по опціональному набору міток
         for item in list_of_tags_optional:
            if item in day:
               #print('Found optional:',k,item)
               if day[item]!='<>':
                  if not '<' in day[item]:
                     print(f'День {k}: пропущено відкриваючий <{item}>')
                     day['integrity']='NOK'
                  if not '>' in day[item]:
                     print(f'День {k}: пропущено закриваючий </{item}>')
                     day['integrity']='NOK'
               else:
                  if not item in optional_tag_summary:
                     optional_tag_summary[item]=[k]
                  else:
                     optional_tag_summary[item]+=[k]

         
      else:
         print(f'День {k}: немає Літургії або відсутній <ustav>')
         day['integrity']='NOK'

   print("\nДні використання опціональних рубрик:")
   pretty(optional_tag_summary)

      
   for k,day in history.items():
      if day['integrity']=='NOK':
         raise Exception
   return history, optional_tag_summary


def initDateCalendar():
   month_no = datetime.now().month+1 if datetime.now().month!=12 else 1
   month_no=10
   print("WARNING. Month_no OVERRIDE", month_no)
   cur_month=month_dic_reversed[month_no]

   mode = easygui.choicebox('u - Юліанський, g - Григоріанський', 'Вибір календаря', ['u','g'])
   if mode == 'u':
      mode_suffix='НЮ'
   elif mode == 'g':
      mode_suffix='ГР'

   return month_no, cur_month, mode, mode_suffix





if __name__ == "__main__":
   month_no, cur_month, mode, mode_suffix = initDateCalendar()
   docx_filename=f'{month_no:02}-Літургія-{mode_suffix.upper()}.docx'

   checkLiturgyIntegrity(docx_filename)
   

            

            
            
      
         


                




