import easygui, calendar, os, docx, re, csv
from docx.shared import RGBColor
import paschalia
from datetime import datetime, timedelta

mode = easygui.choicebox('u - Юліанський, g - Григоріанський', 'Вибір календаря', ['u','g'])
mode_dict={'u':'Юл',
           'g':'Гр'}

paschalia.paschalia = paschalia.get_prev_next_pascha(datetime(2024,1,1),mode)

#select month
#month_no = easygui.enterbox("Введіть номер місяця", "Місяць")
month_no = datetime.now().month+1 if datetime.now().month!=12 else 1
year_no = datetime.now().year if datetime.now().month!=12 else datetime.now().year+1
#select mode
#mode = easygui.choicebox('u - Новоюліанський, g - Григоріанський', 'Вибір календаря', ['u','g'])

month_dic = {'Січень':1,
              'Лютий':2,
              'Березень':3,
              'Квітень':4,
              'Травень':5,
              'Червень':6,
              'Липень':7,
              'Серпень':8,
              'Вересень':9,
              'Жовтень':10,
              'Листопад':11,
              'Грудень':12}
month_dic_reversed = {v:k for k,v in month_dic.items()}
month_dic_string='('+'|'.join(month_dic.keys())+')'

day_dic = {"Понеділок":1,
            "Вівторок":2,
            "Середа":3,
            "Четвер":4,
            "П'ятниця":5,
            "Субота":6,
            "Неділя":7}
day_dic_reversed = {v:k for k,v in day_dic.items()}
day_dic_string='('+'|'.join(day_dic.keys())+')'

#doc = docx.Document("Літургія - Мінея\\10-Літургія-шаблони.docx") 
'''
def get_echos(date,mode):
    if mode=='u':
        return (int(date.strftime("%U"))-25) % 8 + 1
    elif mode=='g':
        return (int(date.strftime("%U"))-24) % 8 + 1
'''
def copy_paragraph(target_doc,source_paragraph):
    target_paragraph = target_doc.add_paragraph()
    for run in source_paragraph.runs:
        new_run = target_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.font.color.rgb = run.font.color.rgb
        new_run.font.highlight_color = run.font.highlight_color

def copy_paragraph_list(target_doc,paragraph_list):
    for p in paragraph_list:
        copy_paragraph(target_doc,p)

def format_line(p, handle=''):
    #handle = "bir"
    if 'b' in handle:
        p.runs[0].font.bold = True
    if 'i' in handle:
        p.runs[0].font.italic = True
    if 'r' in handle:
        p.runs[0].font.color.rgb = RGBColor(0xff, 0x44, 0x00)
    p.runs[0].font.name='Times New Roman'
    p.runs[0].font.size=152400


def get_dismissal_matrix(dismissal_csv_filename,cur_month):
    matrix={}
    with open(dismissal_csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:
            if row[0]==month_dic_reversed[cur_month]:
                matrix[int(row[1].split('.')[0])]=row[2:]
    return matrix

def get_matrix_full(csv_filename):
    matrix=[]
    with open(csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:
            matrix.append(row)
    return matrix

        
path = "Літургія - змінні частини - шаблони.docx"
def get_resurrection_template_texts(path):
    template_dic={}
    doc = docx.Document(path)
    rubric_found = cur_glas = None
    i=0
    for p in doc.paragraphs:
        #print(p.text)
        if not rubric_found and not p.text.startswith("Воскресна служба"):
            pass
        elif p.text.startswith("Воскресна служба"):
            rubric_found=True
        if rubric_found:
            re_result = re.search(r"Неділя – глас (\d)",p.text)
            if re_result:
                cur_glas = int(re_result.group(1))
                template_dic[cur_glas]=[]
                continue

        if cur_glas and p.text.startswith("</vidpust>"):
            template_dic[cur_glas].append(p)
            cur_glas=None

        if cur_glas:
            template_dic[cur_glas].append(p)
    return template_dic
            
def get_everyday_template_texts(path):
    template_dic={}
    doc = docx.Document(path)
    rubric_found = cur_glas = None
    i=0
    for p in doc.paragraphs:
        #print(p.text)
        if not rubric_found and not p.text.startswith("Повсякденна служба"):
            pass
        elif p.text.startswith("Повсякденна служба"):
            rubric_found=True
        if rubric_found:
            re_result = re.search(f"{day_dic_string}",p.text)
            if re_result:
                cur_glas = day_dic[re_result.group(1)]
                template_dic[cur_glas]=[]
                continue

        if cur_glas and p.text.startswith("</vidpust>"):
            template_dic[cur_glas].append(p)
            cur_glas=None

        if cur_glas:
            template_dic[cur_glas].append(p)
    return template_dic
        
        
def get_menaion_template_texts():
    directory_path ="Літургія - Мінея"
    all_files = os.listdir(directory_path)
    docx_file = [file for file in all_files if file.endswith(".docx") and file[:2]==f'{month_no:02}' and len(file) >= 4][0]
    print(docx_file)
    doc = docx.Document("Літургія - Мінея\\"+docx_file)
    template_found=False
    starting_tag_found=False
    template_dic={}
    for p in doc.paragraphs:
        
        re_result=re.search(f'^{month_dic_string} '+r'(\d+)',p.text)
        if re_result:
            cur_date = int(re_result.group(2))
            template_dic[cur_date]=[]
            template_found = True
        
        re_result=re.search(f'<ustav',p.text)
        if re_result:
            starting_tag_found=True
            
        if template_found and starting_tag_found:
            template_dic[cur_date].append(p)

        re_result=re.search(f'</vidpust',p.text)
        if re_result:
            #template_dic[cur_date]=template_dic[cur_date][1:]
            template_found=False
            cur_date=None
            starting_tag_found=False
            
    
    #print(template_dic)
    return template_dic

#init data
templates_menaion = get_menaion_template_texts()
templates_resurrection = get_resurrection_template_texts(path)
templates_everyday = get_everyday_template_texts(path)
saint_matrix = get_matrix_full("Місяцеслов-БД.csv")

dismissal_matrix = get_dismissal_matrix(f'Читання{mode_dict[mode]}.csv',month_no)


def insert_header_liturgy(doc,date):
    day_name = day_dic_reversed[date.weekday()+1]
    #month_name = month_dic_reversed[month_no]
    week_no=paschalia.get_week(date,"","")[:2]

    #Субота, Неділя, Тиждень etc...
    special_dates=[{"date":datetime(year_no,12,25),
                     "holiday":"Різдво",
                     "holiday_locative":"Різдві",
                     "holiday_instrumental":"Різдвом"},
                    {"date":datetime(year_no+1,1,6),
                     "holiday":"Богоявленні",
                     "holiday_locative":"Богоявленні",
                     "holiday_instrumental":"Богоявленням"}]
    for sd in special_dates:
        diff = (sd["date"]-date).days
        if abs(diff)<=7 and date.weekday()+1 in [6,7]:
            if diff>0:
                p_new = doc.add_paragraph(f"{day_name} перед {sd['holiday_instrumental']}")
            else:
                p_new = doc.add_paragraph(f"{day_name} по {sd['holiday_locative']}")
            format_line(p_new, '')
    
    if date.weekday()+1 in [6,7]:
        p_new = doc.add_paragraph(f"{day_name} {week_no} по П'ятидесятниці.")
        if date.weekday()+1 == 7:
            p_new.text+= f" Гл. "+str(paschalia.get_echos(date))+"."
            format_line(p_new, 'r')
        else:
            format_line(p_new, '')

    if date.weekday()+1 in [1]:
        p_new = doc.add_paragraph(f"Тиждень {week_no} по П'ятидесятниці.")
        format_line(p_new, '')

    #перелік святих
    lst = filter(lambda l:int(l[0])==date.month and int(l[1])==date.day,saint_matrix[1:])
    for l in lst:
        p_new=doc.add_paragraph(l[9])
        format_line(p_new, ''.join(l[2:5]))



new_doc = docx.Document()
'''
Get all days in the month
'''
for d in range(1,calendar.monthrange(year_no, month_no)[1]+1):
#for d in range (5,6):
    print(d)
    heading_text =' '.join([month_dic_reversed[month_no],str(d)+',',day_dic_reversed[datetime(year_no, month_no, d).weekday()+1]])
    new_doc.add_heading(heading_text, level=2)
    #print(heading_text)

    #REDO! in January =D
    #new_doc.add_paragraph(dismissal_matrix[d][2])
    insert_header_liturgy(new_doc,datetime(year_no,month_no,d))

    #print(dismissal_matrix[d][2])
    if datetime(year_no, month_no, d).weekday()+1==7:
        #print(d, "copied sunday ", get_echos(datetime(year_no,month_no,d)))
        copy_paragraph_list(new_doc,templates_resurrection[paschalia.get_echos(datetime(year_no,month_no,d))])
    elif d in templates_menaion.keys():
        copy_paragraph_list(new_doc,templates_menaion[d])
    else:
        print(d, datetime(year_no, month_no, d).weekday()+1,templates_everyday[datetime(year_no, month_no, d).weekday()+1][38].text[:20])
        copy_paragraph_list(new_doc,templates_everyday[datetime(year_no, month_no, d).weekday()+1])
        
    
mode_new = 'ГР' if mode=='g' else "НЮ"
new_doc.save(f'{month_no}-Літургія-{mode_new}.docx')
