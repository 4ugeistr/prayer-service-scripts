import re, easygui, calendar, glob,shutil,os,docx,csv
from datetime import datetime
from docx.shared import RGBColor
import paschalia
import get_stichera
RGB_RED = RGBColor(0xff, 0x44, 0x00)
BLACK='b'
RED='r'   

month_no = datetime.now().month+1 if datetime.now().month!=12 else 1
year_no = datetime.now().year if datetime.now().month!=12 else datetime.now().year+1

mode = easygui.choicebox('u - Юліанський, g - Григоріанський', 'Вибір календаря', ['u','g'])

day_short_dic={"ПН":1,
        "ВТ":2,
        "СР":3,
        "ЧТ":4,
        "ПТ":5,
        "СБ":6,
        "НД":7}
day_short_dic_reversed = {v:k for k,v in day_short_dic.items()}

month_dic= {'Січень':1,
              'Лютий':2,
              'Березень':3,
              'Квітень':4,
              'Травень':5,
              'Червень':6,
              'Липень':7,
              'Серпень':8,
              'Вересень':9,
              'Жовтень':10,
              'Листопад':11,
              'Грудень':12}
month_dic_reversed = {v:k for k,v in month_dic.items()}
#month_dic_string='('+'|'.join([x.lower() for x in month_dic.keys()])+')'


lpd_templates = glob.glob('ЛПД/*/*.docx')
#lpd_templates_filenames = [x.split('\\')[-1] for x in lpd_templates]


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def delete_run(run):
    p = run._element
    p.getparent().remove(p)
    p._p = p._element = None

def copy_run(target_paragraph,run):
    new_run = target_paragraph.add_run(run.text)
    #KeyError: "no style with name 'Default Paragraph Font'"
    #new_run.style = run.style.name
    new_run.bold = run.bold
    new_run.italic = run.italic
    new_run.underline = run.underline
    #new_run.font.size = run.font.size
    #new_run.font.name = run.font.name
    new_run.font.name='Times New Roman'
    new_run.font.size=152400
    new_run.font.color.rgb = run.font.color.rgb
    new_run.font.highlight_color = run.font.highlight_color

def copy_paragraph_before(paragraph_to_insert_before,source_paragraph):
    target_paragraph = paragraph_to_insert_before.insert_paragraph_before()
    try:
        target_paragraph.style = source_paragraph.style
        #target_paragraph.style.font = source_paragraph.style.font.name
    except KeyError:
        print(f"Warning. Text {source_paragraph.text[:20]} has style{source_paragraph.style}")
    for run in source_paragraph.runs:
        copy_run(target_paragraph,run)
    return target_paragraph


def stretch_texts(qty_of_verses_requested, texts):
    result = []
    repetitions_per_text, remaining_verses = divmod(qty_of_verses_requested, len(texts))
    for t in texts:
        result.append([t]*repetitions_per_text)
    for i in range(remaining_verses):
        result[i].append(result[i][0])
    return sum(result,[])


def add_text(p,text, color=BLACK):
    r = p.add_run(text)
    r.font.name='Times New Roman'
    r.font.size=152400
    if color == RED:
        r.font.color.rgb = RGB_RED
        r.italic = True

def format_line(p, handle=''):
    #handle = "bir"
    if 'b' in handle:
        p.runs[0].font.bold = True
    if 'i' in handle:
        p.runs[0].font.italic = True
    if 'r' in handle:
        p.runs[0].font.color.rgb = RGBColor(0xff, 0x44, 0x00)
    p.runs[0].font.name='Times New Roman'
    p.runs[0].font.size=152400

def insert_dismissal(path,date):
    doc = docx.Document(path)
    shoutout_found = None
    for p in doc.paragraphs:
        re_result=re.search(r"(\.|\)) Благослови\.",p.text)
        if re_result:
            shoutout_found = True
            #print(date.day, "found!")
            continue

        if shoutout_found:
            #print(date.day, "inserting")
            p_new=p.insert_paragraph_before(dismissal_matrix[date.day][-1])
            #print(date.day,p_new.text)
            format_line(p_new, '')
            delete_paragraph(p)
            shoutout_found = False

    for p in doc.paragraphs:
        if re.search(f'Священик:',p.text):
            re_result=re.search(f'^(Священик:)( .+?)(якого є храм)(.*?)$',p.text)
            p_bak=p.text
            p.clear()

            try:
                if re_result:
                    add_text(p,re_result.group(1),color=RED)
                    add_text(p,re_result.group(2))
                    add_text(p,re_result.group(3),color=RED)
                    add_text(p,re_result.group(4))
                else:
                    re_result=re.search(f'^(Священик:)(.*?)$',p_bak)
                    add_text(p,re_result.group(1),color=RED)
                    add_text(p,re_result.group(2))
            except:
                print(p_bak)
                raise
    doc.save(path)
    
def get_dismissal_matrix(dismissal_csv_filename,cur_month):
    matrix={}
    with open(dismissal_csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:
            if row[0]==month_dic_reversed[cur_month]:
                matrix[int(row[1].split('.')[0])]=row[2:]
    return matrix

stichos_list=[
    "Стих: Якщо на беззаконня зважатимеш, Господи",
    "Стих: Імени твого ради чекав я на Тебе",
    "Стих: Від ранньої сторожі до ночі",
    "Стих: Бо в Господа милість і багате в нього ізбавління",
    "Стих: Хваліть Господа, всі народи",
    "Стих: Бо утвердилася милість його на нас",
    ]


if mode == 'u':
    mode_suffix='Юл'
elif mode == 'g':
    mode_suffix='Гр'

if __name__ == "__main__":
    paschalia_dates = paschalia.get_prev_next_pascha(datetime(year_no, month_no,1), mode)
    dismissal_matrix = get_dismissal_matrix(f'Відпусти{mode_suffix}.csv',month_no)
    
    stub_dic={}
    folder= f'drafts\\{year_no}-{month_no:02}-{mode}-lpd'
    #os.makedirs('drafts', exist_ok=True)
    os.makedirs(folder, exist_ok=True)

    for d in range(1,calendar.monthrange(year_no, month_no)[1]+1):
        #stub_dic[d]=None
        day_details = paschalia.get_day_details(datetime(year_no,month_no,d),paschalia_dates)
        #expected_template_path = f"ЛПД\\тиждень-{day_details[1]}\\{day_details[1]}-{day_details[3]}-ЛПД.docx"
        expected_template_path = f"ЛПД\\{day_details[1]} тиждень\\{day_details[1]}т_{day_details[3]}-ЛПД.docx"
        if expected_template_path in lpd_templates:
            filename=folder+f'\\{d:02}.{month_no:02}-{day_details[1]}т_{day_details[3]}-ЛПД.docx'
            shutil.copy2(expected_template_path, filename)
            stub_dic[d]=filename
        #re_result = re.search("(\d)-(\d)-ЛПД.docx")
    print(f"Created {len(stub_dic)} stubs for month {month_no} mode {mode}!")

    stichera_matrix = get_stichera.get_stichera_matrix(glob.glob(f'Стихири - Мінея\\Мінея_{month_no:02}*.docx')[0])

    for k,v in stub_dic.items():
        stichera_no=None
        look_for_slava=False
        doc = docx.Document(v)
        stichera_qty = 6 if k==26 else 4
        stichos_dic = {stichos_list[-stichera_qty:][i]: i for i in range(stichera_qty) }
        stichos_dic_string='('+'|'.join(stichos_dic.keys())+')'
        try:
            texts = stretch_texts(stichera_qty,stichera_matrix[k]['gv_stichera'])
        except IndexError:
            print(f"Skipping day{k}")
            print(IndexError)
            
            continue
        for p in doc.paragraphs:
            re_result = re.search(f"{stichos_dic_string}",p.text)
            if re_result:
                #print(f"Day {k}, found string:",re_result.group(1)[:20])
                stichera_no=stichos_dic[re_result.group(1)]
                continue

            if stichera_no or stichera_no==0:
                copy_paragraph_before(p,texts[stichera_no])
                delete_paragraph(p)
                #print(f"Inserted stichera {stichera_no} for {v}")
                stichera_no=None
                look_for_slava=True
                continue
        
            re_result = re.search("^(Слава|І нині)",p.text)
            if re_result and look_for_slava:
                delete_paragraph(p)

            re_result = re.search("^Вхід",p.text)
            if re_result:
                look_for_slave=False
                if "gv_doxa" in stichera_matrix[k]:
                    print(f"Inserted doxa for {k}")
                    copy_paragraph_before(p,stichera_matrix[k]["gv_doxa"][0])
                    if re.search("слава",stichera_matrix[k]["gv_doxa"][0].text.lower()) and re.search("і нині",stichera_matrix[k]["gv_doxa"][0].text.lower()):
                        continue
        

                if datetime(year_no,month_no,k).weekday()+1 in (3,5):
                    if  len(stichera_matrix[k]["gv_theotokion"])==2:
                        print(f"Inserted theo for {k}")
                        copy_paragraph_before(p,stichera_matrix[k]["gv_theotokion"][1])
                    else:
                        print(f"Warning, failed to find theotokion for day {k}")
                        p.insert_paragraph_before("ХРЕСТОБОГОРОДИЧНИЙ")
                else:
                    print(f"Inserted theo for {k}")
                    try:
                        copy_paragraph_before(p,stichera_matrix[k]["gv_theotokion"][0])
                    except KeyError:
                        print(f"Warning, failed to find theotokion for day {k}")
                        p.insert_paragraph_before("БОГОРОДИЧНИЙ")
                            
        doc.save(v)

    for k,v in stub_dic.items():
        #print("Дата:",d,datetime(year_no, month_no, d).weekday()+1,mode)
        #insert_header(v,datetime(year_no, month_no, d))
        insert_dismissal(v,datetime(year_no, month_no, k))
        
