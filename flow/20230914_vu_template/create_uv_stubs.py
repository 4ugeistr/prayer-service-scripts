import re,glob,calendar,docx,os,easygui, shutil, csv,copy
from datetime import datetime
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_COLOR_INDEX
import paschalia,get_stichera
RGB_RED = RGBColor(0xff, 0x44, 0x00)


#filenames= glob.glob('*/*/*.txt')
mode = easygui.choicebox('u - Юліанський, g - Григоріанський', 'Вибір календаря', ['u','g'])
#paschalia.paschalia = paschalia.get_prev_next_pascha(datetime(2024,1,1),mode)



month_no = datetime.now().month+1 if datetime.now().month!=12 else 1
year_no = datetime.now().year if datetime.now().month!=12 else datetime.now().year+1

#month_no=6
#print(f"WARNING: MONTH OVERRIDE!")
print(f"Processing month: {month_no}")

day_short_dic={"ПН":1,
        "ВТ":2,
        "СР":3,
        "ЧТ":4,
        "ПТ":5,
        "СБ":6,
        "НД":7}
day_short_dic_reversed = {v:k for k,v in day_short_dic.items()}
day_short_dic_string='('+'|'.join(day_short_dic.keys())+')'

day_dic = {"Понеділок":1,
            "Вівторок":2,
            "Середа":3,
            "Четвер":4,
            "П'ятниця":5,
            "Субота":6,
            "Неділя":7}
day_dic_reversed = {v:k for k,v in day_dic.items()}
day_dic["П’ятниця"]=5
day_dic_string='('+'|'.join(day_dic.keys())+')'

month_dic= {'Січень':1,
              'Лютий':2,
              'Березень':3,
              'Квітень':4,
              'Травень':5,
              'Червень':6,
              'Липень':7,
              'Серпень':8,
              'Вересень':9,
              'Жовтень':10,
              'Листопад':11,
              'Грудень':12}
month_dic_reversed = {v:k for k,v in month_dic.items()}
month_dic_string='('+'|'.join([x.lower() for x in month_dic.keys()])+')'
month_w_offset = list(month_dic.values())[4:]+list(month_dic.values())[:4]


if mode == 'u':
    mode_suffix='Юл'
    mode_suffix2='НЮ'
elif mode == 'g':
    mode_suffix='Гр'
    mode_suffix2='ГР'
'''
def copy_paragraph_before(p_to_insert_before,source_paragraph):
    target_paragraph = p_to_insert_before.insert_paragraph_before()
    target_paragraph.style = source_paragraph.style
    target_paragraph.alignment = source_paragraph.alignment
    for run in source_paragraph.runs:
        new_run = target_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.font.color.rgb = run.font.color.rgb
        new_run.font.highlight_color = run.font.highlight_color
'''

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def delete_run(run):
    p = run._element
    p.getparent().remove(p)
    p._p = p._element = None

def copy_run(target_paragraph,run):
    new_run = target_paragraph.add_run(run.text)
    #KeyError: "no style with name 'Default Paragraph Font'"
    new_run.style = run.style.name
    new_run.bold = run.bold
    new_run.italic = run.italic
    new_run.underline = run.underline
    #new_run.font.size = run.font.size
    #new_run.font.name = run.font.name
    new_run.font.name='Times New Roman'
    new_run.font.size=152400
    new_run.font.color.rgb = run.font.color.rgb
    new_run.font.highlight_color = run.font.highlight_color

def copy_paragraph_before(paragraph_to_insert_before,source_paragraph):
    target_paragraph = paragraph_to_insert_before.insert_paragraph_before()
    try:
        target_paragraph.style = source_paragraph.style
        target_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after
        #target_paragraph.style.font = source_paragraph.style.font.name
    except KeyError:
        print(f"Warning. Text {source_paragraph.text[:20]} has style{source_paragraph.style}")
    for run in source_paragraph.runs:
        copy_run(target_paragraph,run)
    return target_paragraph

def copy_paragraph_list_before(p_to_insert_before,paragraph_list):
    #print("qty of p to insert:",len(paragraph_list))
    for p in paragraph_list:
        copy_paragraph_before(p_to_insert_before,p)




def copy_paragraph(target_doc,source_paragraph):
    target_paragraph = target_doc.add_paragraph()
    try:
        target_paragraph.style = source_paragraph.style
    except KeyError:
        print(f"Warning. Text {source_paragraph.text[:20]} has style{source_paragraph.style}")
    target_paragraph.alignment = source_paragraph.alignment
    for run in source_paragraph.runs:
        new_run = target_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.font.color.rgb = run.font.color.rgb
        new_run.font.highlight_color = run.font.highlight_color

def copy_paragraph_list(target_doc,paragraph_list):
    for p in paragraph_list:
        copy_paragraph(target_doc,p)

def get_octoechos_template_files():
    template_dic = {}
    all_files = glob.glob(r'01-Октоїх/*/*.docx')
    for f in all_files:
        re_result=re.search(r'01-Октоїх\\Глас_(\d)\\(\d)-(\w{2})-?.docx',f)
        if re_result:
            #print(re_result.group(2), re_result.group(3))
            doc = docx.Document(f)
            if int(re_result.group(2)) not in template_dic.keys():
                template_dic[int(re_result.group(2))]={}
            if re_result.group(3) in day_short_dic.keys():
                template_dic[int(re_result.group(2))][day_short_dic[re_result.group(3)]]=f
            else:
                print("Something wrong with", f, re_result.group(3))
        else:
            print("None found for:", f)
    return template_dic


irmos_list = [1,3,4,5,6,7,8,9]

def get_octoechos_kanon_texts(path):
    doc = docx.Document(path)
    matrix = {}
    for i in range(8):
        matrix[i+1]={}
    for p in doc.paragraphs:
        
        re_result = re.search(r"Глас (\d)",p.text)
        if re_result:
            echos = int(re_result.group(1))
            echos_found = True
            song_found = False
            continue

        if (echos_found or song_found) and re.search(f"{day_dic_string}",p.text):
            week_day_no = day_dic[re.search(f"{day_dic_string}",p.text).group(1)]
            matrix[echos][week_day_no]=[]
            week_day_no_found = True
            echos_found = song_found = False
            continue
        
        if (week_day_no_found or song_found) and re.search("Канон", p.text):
            matrix[echos][week_day_no].append({"label":p.text})
            label_found = True
            week_day_no_found = song_found = False
            continue

        re_result=re.search(r"Пісня (\d)",p.text)
        if (label_found or song_found) and re_result:
            song_no = int(re_result.group(1))
            matrix[echos][week_day_no][-1][song_no]=[p]
            label_found=False
            song_found = True
            continue
        
        if p.text == "":
            continue
        if song_found:
            matrix[echos][week_day_no][-1][song_no].append(p)

    return matrix

def get_kanon_litany(path):
    template_dic={}
    doc = docx.Document(path)
    litany_no = None
    for p in doc.paragraphs:
        re_result = re.search(r"Мала єктенія (\d)",p.text)
        if re_result:
            litany_no=int(re_result.group(1))
            template_dic[litany_no]=[]
            continue
        elif litany_no:
            template_dic[litany_no].append(p)
    return template_dic

def get_menaion_template_files():
    filenames = glob.glob(f'В,У - Мінея/{month_no:02}*/*.docx')
    template_dic={}
    for f in filenames:
        #print("Checking",f)
        pattern=r'(\d{2})-__-(.*?).docx'
        #print(pattern)
        re_result = re.search(pattern,f)
        doc = docx.Document(f)
        template_dic[int(re_result.group(1))]=f
    
    return template_dic

def get_menaion_templates():
    filenames = glob.glob(f'В,У - Мінея/{month_no:02}*/*.docx')
    template_dic={}
    for f in filenames:
        print("Checking",f)
        pattern=r'(\d{2})-__-(.*?).docx'
        print(pattern)
        re_result = re.search(pattern,f)
        doc = docx.Document(f)
        template_dic[int(re_result.group(1))]=doc.paragraphs
    return template_dic

def get_resurrection_troparia_texts(path):
    template_dic={}
    doc = docx.Document(path)
    echos = None
    #i=0
    for p in doc.paragraphs:
        re_result = re.search(r"Глас (\d)",p.text)
        if re_result:
            #print("found")
            echos=int(re_result.group(1))
            template_dic[echos]=[]
        #print(p.text)

        re_result = re.search(r"((Тропар|Кондак).*?\(г\. \d\)): (.*)",p.text)
        if echos and re_result:
            #template_dic[echos].append([re_result.group(1),re_result.group(3)])
            template_dic[echos].append(p)
    return template_dic

def get_menaion_troparia_texts(path):
    #print("in")
    template_dic={}
    doc = docx.Document(path)
    key = None #cur_day
    saint={}
    for p in doc.paragraphs:
        re_result = re.search(r"^(\d+) (.*)",p.text)
        if re_result:
            #print("found")
            key=int(re_result.group(1))
            cur_day_heading = re_result.group(2)
            template_dic[key]=[]

        if key:
            if re.search("^Тропар",p.text):
                re_result = re.search(r"(Тропар.*?\(г\. \d\): )(.*)",p.text)
                try:
                    #saint["troparion"]=[re_result.group(1), re_result.group(2)]
                    saint["troparion"]=p
                except IndexError as e:
                    print(p.text, re_result)
                    raise e
            elif re.search("^Кондак",p.text):
                re_result = re.search(r"(Кондак.*?\(г\. \d\): )(.*)",p.text)
                #saint["kondakion"]=[re_result.group(1), re_result.group(2)]
                saint["kondakion"]=p
            else:
                #found saint header
                saint["header"]=p.text    

        if "header" in saint.keys() and "troparion" in saint.keys() and "kondakion" in saint.keys():
            template_dic[key].append(saint)
            saint={}
        #    template_dic[echos].append([re_result.group(1),re_result.group(3)])
    return template_dic

def get_vespers_prokimenon(path):
    template_dic={}
    doc = docx.Document(path)
    day=None
    for p in doc.paragraphs:
        re_result = re.search(f"^{day_dic_string}",p.text)
        if re_result:
            day=re_result.group(1)
            template_dic[day_dic[day]]=[]
        elif day:
            template_dic[day_dic[day]].append(p)
    return template_dic

def get_matrix_full(csv_filename):
    matrix=[]
    with open(csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:
            matrix.append(row)
    return matrix

def get_matrix(csv_filename):
    matrix={}
    with open(csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:

            if row[0]==str(month_no):
                #print(row[1].split('.')[2],row)
                #print(row[2:])
                matrix[int(row[1].split('.')[2])]=row[2:]
    return matrix

#need to merge with previous one
def get_dismissal_matrix(dismissal_csv_filename,cur_month):
    matrix={}
    with open(dismissal_csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:
            if row[0]==month_dic_reversed[cur_month]:                matrix[int(row[1].split('.')[0])]=row[2:]
    return matrix



def get_theotokion_troparia_texts(path):
    service_dic = {'Вечірня':'vespers',
                   'Утреня':'orthros'}
    template_dic={}
    doc = docx.Document(path)
    for p in doc.paragraphs:
        re_result = re.search(r"Глас (\d)",p.text)
        if re_result:
            #print("found")
            echos=int(re_result.group(1))
            template_dic[echos]=[]
        #print(p.text)

        re_result = re.search(r"(Вечірня|Утреня) Відпуст",p.text)
        if re_result:
            #print(echos,re_result.group(1))
            service=service_dic[re_result.group(1)]

        re_result = re.search(f"^{day_dic_string}",p.text)
        if re_result:
            day = re_result.group(1)
            weekday_no = day_dic[day]

        re_result = re.search(f"Богородичний",p.text)
        if re_result:
            template_dic[echos].append({'service':service, 'weekday':weekday_no,'p':p})
            
    return template_dic

def insert_prokimenon(path,day):
    doc = docx.Document(path)
    #print(path)
    vespers_prokimenon_found = False
    vespers_prokimenon_end_found = True
    for p in doc.paragraphs:
        re_result = re.search("Прокімен",p.text)
        if re_result:
            vespers_prokimenon_found = True
            #delete_paragraph(p)
        re_result = re.search("(Читання|Сподоби, Господи)",p.text)
        if re_result:
            #print(day)
            for p1 in vespers_prokimenon[day]:
                copy_paragraph_before(p,p1)
            p.insert_paragraph_before()
            break
        elif vespers_prokimenon_found:
            #print(f"Deleting {p.text}")
            delete_paragraph(p)
    doc.save(path)

def insert_boh_hospod_echos(path,date):
    #print(date.day)
    doc = docx.Document(path)
    for p in doc.paragraphs:
        re_result = re.search("Бог Господь",p.text)
        if re_result:
            #print("found")
            if date.weekday()+1==7:
                echos = paschalia.get_echos(date,paschalia_dates)
            else:
                #echos = int(re.search("(\d)",templates_menaion[date.day][0]['troparion'][0]).group(1))
                echos = int(re.search(r"(\d)",templates_menaion[date.day][0]['troparion'].text).group(1))

            for r in p.runs:
                if '???' in r.text:
                    #print(r.text)
                    r.font.highlight_color=None
                    r.text = r.text.replace('???',str(echos))
                    
            break
    doc.save(path)

def insert_prefix_to_paragraph(p,text="Слава: ",formatting='b'):
    p_new = copy_paragraph_before(p,p)
    #print("iptp:    ",p_new.text)
    new_run = p_new.add_run(text)
    if 'b' in formatting:
        new_run.font.bold = True
    new_run.font.name='Times New Roman'
    new_run.font.size=152400
    
    bkp=p_new.runs[:-1]
    for r in p_new.runs:
        delete_run(r)
        
    copy_run(p_new,new_run)
    for r in bkp:
        copy_run(p_new,r)
    return p_new

def get_troparia_theotokion(echos, date, service):
    #print(echos, date, service)
    res = filter(lambda t: t['service'] == service and t['weekday'] == date.weekday()+1, templates_theotokion_dic[echos])
    l = list(res)
    p = l[0]["p"]
    return p
    #return list(filter(lambda t: t['service'] == service and t['weekday'] == date.weekday()+1, templates_theotokion_dic[echos]))[0]["p"]

def get_troparia_block(date,service):
    #if date.day!=1:
    #    return 0
    
    troparia_block=[]
    troparia=None
    if ordo_matrix[date.day][1]=='y':
        print("Something wrong")
        raise Exception("Should have a ready template")
    saint_counter=0
    #print("preparing troparia:",date.day, service)
    for i in list(range(4)):
        #if date.day==2:
        #    print(f"processing troparia for day{d}. item {i}")
        try:
            if not ordo_matrix[date.day][2:][i]:
                continue
        except IndexError as e:
            print('erred on:',i)
            raise e
            
            
        if ordo_matrix[date.day][2:][i] == 'resurrection':
            #if date.day==2:
            #    print(f"inserting troparia for day{d}. item {i}")
            troparia = templates_resurrection[paschalia.get_echos(date,paschalia_dates)][0]
        if ordo_matrix[date.day][2:][i] == 'triodion':
            if i==3:
                troparia=insert_prefix_to_paragraph(templates_triodion[date.day][0]["troparion"],text="І нині: ")
            else:
                troparia = templates_triodion[date.day][0]["troparion"]
        if ordo_matrix[date.day][2:][i] == 'saint':
            #print(date.day, i)
            try:
                if i==2:
                    #print("adding Slava")
                    #print(templates_menaion[date.day][saint_counter]["troparion"].text)
                    troparia=insert_prefix_to_paragraph(templates_menaion[date.day][saint_counter]["troparion"],text="Слава: ")
                    #print(troparia.text)
                else:
                    troparia=templates_menaion[date.day][saint_counter]["troparion"]
            except IndexError as e:
                print(date.day, e)
                print("Check if troparion is in the file")
                print("Check if menaion matrix is populated")
                raise e

            saint_counter+=1
            
        if ordo_matrix[date.day][2:][i] == 'theotokos':
            echos = int(re.search(r"(\d)",troparia_block[-1].text).group(1))
        
            if ordo_matrix[date.day][2:][2]:
                troparia=insert_prefix_to_paragraph(get_troparia_theotokion(echos,date,service),text="І нині: ")
            else:
                troparia=insert_prefix_to_paragraph(get_troparia_theotokion(echos,date,service),text="Слава, і нині: ")    
        troparia_block.append(troparia)
    
    #???
    troparia_block.append(troparia_block[-1].insert_paragraph_before())
    #print([item.text[:36] for item in troparia_block])

    for p in troparia_block:
        re_result = re.search("^(.*?)(Богородичний |Тропар |Кондак )(\(г. \d\): )(.*)",p.text)
        if re_result:
            for r in p.runs:
                delete_run(r)
            if re_result.group(1):
                r_new = p.add_run(re_result.group(1))
                format_run(r_new,'b')
            r_new = p.add_run(re_result.group(3))
            format_run(r_new,'ri')
            r_new = p.add_run(re_result.group(4))
            format_run(r_new,'')

    #print([item.text[:36] for item in troparia_block])
    return troparia_block

def insert_troparia(path,date):
    print(date.day)
    doc = docx.Document(path)
    troparia_block = {"orthros": get_troparia_block(date,'orthros'),
                      "vespers": get_troparia_block(date,'vespers')}
    #troparia_block=get_troparia_block(date,'orthros')
    troparion_found = False
    troparion_end_found = False
    troparion_label_count=0
    service = False
    for p in doc.paragraphs:
        #if troparion_found and date.day==8:
        #    print(p.text[:20])
        if re.search("ВЕЧІРНЯ",p.text):
            service = 'vespers'
        if re.search("УТРЕНЯ",p.text):
            service = 'orthros'
        
        re_result = re.search("Тропарі",p.text)
        if re_result:
            #print(f"Day {date.day}. Troparion label for service = {service} found")
            troparion_found = True
            troparion_end_found = False
            troparion_label_count+=1
            continue
        
        re_result = re.search("(Великий відпуст|.ктенія усильного благання|Мала .ктенія)",p.text)
        if re_result and troparion_found:
            troparion_end_found=True
            troparia_block_to_copy = copy.deepcopy(troparia_block[service])
            #Додаємо (2 р.) до першого тропаря на Бог Господь.
            if service == 'orthros' and troparion_label_count ==2:
                #r = troparia_block_to_copy[0].add_run("(2 р.)")
                add_text(troparia_block_to_copy[0]," (2 р.)","r")
                #print(f"ADDED {troparia_block_to_copy[0].text}")
            #copy_paragraph_list_before(p,troparia_block[service])
            copy_paragraph_list_before(p,troparia_block_to_copy)
            troparion_found=False
        elif  troparion_found and not troparion_end_found and ordo_matrix[date.day][1]!='y':
            #pass
            delete_paragraph(p)
    if troparion_label_count!=3 and date.weekday()+1 !=7:
        print(f"WARNING! Day {date.day}: troparion_label_count={troparion_label_count}")
    doc.save(path)
    
def format_line(p, handle=''):
    #handle = "bir"
    if 'b' in handle:
        p.runs[0].font.bold = True
    if 'i' in handle:
        p.runs[0].font.italic = True
    if 'r' in handle:
        p.runs[0].font.color.rgb = RGBColor(0xff, 0x44, 0x00)
    p.runs[0].font.name='Times New Roman'
    p.runs[0].font.size=152400

def format_run(r, handle=''):
    #handle = "bir"
    if 'b' in handle:
        r.font.bold = True
    if 'i' in handle:
        r.font.italic = True
    if 'r' in handle:
        r.font.color.rgb = RGBColor(0xff, 0x44, 0x00)
    r.font.name='Times New Roman'
    r.font.size=152400

def insert_header(path,date):
    #print(date)
    doc = docx.Document(path)
    for p in doc.paragraphs:
        re_result = re.search("ВЕЧІРНЯ",p.text)
        if not re_result:
            delete_paragraph(p)
        else:
            #Шапка в форматі "Травень 10 Середа"
            day_name = day_dic_reversed[date.weekday()+1]
            month_name = month_dic_reversed[month_no]
            week_no=paschalia.get_week(date,"","")[:2]
            
            txt =" ".join([month_name,str(date.day)+",",day_name])
            p_new = p.insert_paragraph_before(txt)
            format_line(p_new, '')

            #Субота, Неділя, Тиждень etc...

            special_dates=[{"date":datetime(year_no,12,25),
                             "holiday":"Різдво",
                             "holiday_locative":"Різдві",
                             "holiday_instrumental":"Різдвом"},
                            {"date":datetime(year_no+1,1,6),
                             "holiday":"Богоявленні",
                             "holiday_locative":"Богоявленні",
                             "holiday_instrumental":"Богоявленням"}]
            for sd in special_dates:
                diff = (sd["date"]-date).days
                if abs(diff)<=7 and date.weekday()+1 in [6,7]:
                    if diff>0:
                        p_new = p.insert_paragraph_before(f"{day_name} перед {sd['holiday_instrumental']}")
                    else:
                        p_new = p.insert_paragraph_before(f"{day_name} по {sd['holiday_locative']}")
                    format_line(p_new, '')
            
            if date.weekday()+1 in [6,7]:
                p_new = p.insert_paragraph_before(f"{day_name} {week_no} по П'ятидесятниці")
                format_line(p_new, '')

            if date.weekday()+1 in [1]:
                p_new = p.insert_paragraph_before(f"Тиждень {week_no} по П'ятидесятниці")
                format_line(p_new, '')

            #перелік святих
            lst = filter(lambda l:int(l[0])==date.month and int(l[1])==date.day,saint_matrix[1:])
            for l in lst:
                p_new=p.insert_paragraph_before(l[9])
                format_line(p_new, ''.join(l[2:5]))
            break
    doc.save(path)

def insert_header_from_dismissal_matrix(path,date):
    #print(date)
    doc = docx.Document(path)
    for p in doc.paragraphs:
        re_result = re.search("ВЕЧІРНЯ",p.text)
        if not re_result:
            delete_paragraph(p)
        else:
            day_name = day_dic_reversed[date.weekday()+1]
            month_name = month_dic_reversed[month_no]
            txt =" ".join([month_name,str(date.day)+",",day_name])
            p_new = p.insert_paragraph_before(txt)
            format_line(p_new, 'b')
            p_new = p.insert_paragraph_before(dismissal_matrix[date.day][2])
            format_line(p_new, '')
            break
    doc.save(path)    


def stretch_texts(qty_of_verses_requested, texts):
    result = []
    if qty_of_verses_requested > len(texts):
        repetitions_per_text, remaining_verses = divmod(qty_of_verses_requested, len(texts))
        for t in texts:
            result.append([t]*repetitions_per_text)
        for i in range(remaining_verses):
            result[i].append(result[i][0])
        return sum(result,[])
    else:
        return texts[:qty_of_verses_requested]






# робимо всі "Священик" червоними та italic
BLACK='b'
RED='r'   
def add_text(p,text, color=BLACK):
    r = p.add_run(text)
    r.font.name='Times New Roman'
    r.font.size=152400
    if color == RED:
        r.font.color.rgb = RGB_RED
        r.italic = True
        
def insert_dismissal(path,date):
    doc = docx.Document(path)
    shoutout_found = None
    for p in doc.paragraphs:
        re_result=re.search(r"\. Благослови\.",p.text)
        if re_result:
            shoutout_found = True
            #print(date.day, "found!")
            continue

        if shoutout_found:
            #print(date.day, "inserting")
            p_new=p.insert_paragraph_before(dismissal_matrix[date.day][9])
            #print(date.day,p_new.text)
            format_line(p_new, '')
            delete_paragraph(p)
            shoutout_found = False

    for p in doc.paragraphs:
        if re.search(f'Священик:',p.text):
            re_result=re.search(f'^(Священик:)( .+?)(якого є храм)(.*?)$',p.text)
            p_bak=p.text
            p.clear()

            try:
                if re_result:
                    add_text(p,re_result.group(1),color=RED)
                    add_text(p,re_result.group(2))
                    add_text(p,re_result.group(3),color=RED)
                    add_text(p,re_result.group(4))
                else:
                    re_result=re.search(f'^(Священик:)(.*?)$',p_bak)
                    add_text(p,re_result.group(1),color=RED)
                    add_text(p,re_result.group(2))
            except:
                print(p_bak)
                raise
    doc.save(path)


stichos_list=[
    "Стих: Виведи з в'язниці мою душу",
    "Стих: Мене обступлять праведники",
    "Стих: З глибин взиваю до тебе, Господи",
    "Стих: Нехай будуть твої вуха уважні",
    "Стих: Коли ти, Господи, зважатимеш на беззаконня",
    "Стих: Задля імени твого надіюсь на тебе, Господи",
    "Стих: Від ранньої сторожі до ночі",
    "Стих: Бо в Господа милість і відкуплення велике в нього",
    "Стих: Хваліте Господа всі народи",
    "Стих: Велике бо до нас його милосердя"
]

def insert_menaion_stichera(path,date,template_type):
    doc = docx.Document(path)
    stichera_no=None
    look_for_slava=False
    k=date.day
    delta = None
    dogmatikos = None
    #print(f"inserting {k}")
    #required_stichera_qty = stichera_gv_matrix[k][1]

    if date.weekday() + 1 == 6:
        #print(f"Стихири: пропускаємо {k}, бо субота, обернений порядок стихир")
        #return -2
        if stichera_gv_matrix[k][1]=='3':
            delta = 3
        #else:
        #    print(f"Стихири: пропускаємо {k}, бо субота і стихир {stichera_gv_matrix[k][1]}")
            #return -2
    
    if stichera_gv_matrix[k][1].isnumeric() and stichera_gv_matrix[k][1]!='0':
        #stichos_dic =init_stichera_dic(stichera_gv_matrix[k][1])
        n= int(stichera_gv_matrix[k][1])
        '''
        if date.weekday() + 1 in (1,2,3,4,5,7):
            stichos_dic = {stichos_list[-n:][i]: i for i in range(n) }
        if date.weekday() + 1 in (6,) and template_type=="октоїх":
            stichos_dic = {stichos_list[-n - delta :-delta][i]: i for i in range(n) }
        '''
        if date.weekday() + 1 in (6,) and template_type=="октоїх":
            stichos_dic = {stichos_list[-n - delta :-delta][i]: i for i in range(n) }
        else:
            stichos_dic = {stichos_list[-n:][i]: i for i in range(n) }
        #print(stichos_dic)

    else:
        return 0

    stichos_dic_string='('+'|'.join(stichos_dic.keys())+')'
    #print(date.day,stichos_dic_string)

    if not 'gv_stichera' in stichera_matrix[k] and stichera_gv_matrix[k][3]: 

            stichera_matrix[k]['gv_stichera'] = generic_stichera_matrix[stichera_gv_matrix[k][3]]['gv_stichera']
            stichera_matrix[k]['gv_doxa'] = generic_stichera_matrix[stichera_gv_matrix[k][3]]['gv_doxa']
            stichera_matrix[k]['gv_theotokion'] = generic_stichera_matrix[stichera_gv_matrix[k][3]]['gv_theotokion']

            for p in (stichera_matrix[k]['gv_stichera']+stichera_matrix[k]['gv_doxa']):
                for r in p.runs:
                    if re.search('(ім’я)',r.text):
                        r.text = r.text.replace('(ім’я)', stichera_gv_matrix[k][4])
                        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        r.font.italic = False
                        r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

            print(f"Стихири: використано стихири заг. служби для {k}")
    
    if 'gv_stichera' in stichera_matrix[k]:
        texts = stretch_texts(n,stichera_matrix[k]['gv_stichera'])
    else:
        print(f"Стихири: пропускаємо {k}, немає стихир ГВ.")
        return -1
    
    delete_empty = False
    for p in doc.paragraphs:
        re_result = re.search(f"{stichos_dic_string}",p.text)
        if re_result:
            #print(f"Day {k}, found string:",re_result.group(1)[:20])
            stichera_no=stichos_dic[re_result.group(1)]
            delete_empty = True
            continue

        if stichera_no or stichera_no==0:
            copy_paragraph_before(p,texts[stichera_no])
            delete_paragraph(p)
            #print(f"Inserted stichera {stichera_no} for {v}")
            stichera_no=None
            look_for_slava=True
            continue
        
        if p.text == "" and delete_empty:
            delete_paragraph(p)
            continue

        re_result = re.search("^(Слава|І нині)",p.text)
        if re_result and look_for_slava:
            if re.search('догмат',p.text):
                dogmatikos = p
            delete_paragraph(p)

        re_result = re.search("^(Вхід|Світло тихе)",p.text)
        if re_result:
            look_for_slava=False
            if "gv_doxa" in stichera_matrix[k]:
                #print(f"Inserted doxa for {k}")
                copy_paragraph_before(p,stichera_matrix[k]["gv_doxa"][0])
            #if datetime(year_no,month_no,k).weekday()+1 in (1,2,3,4,5):
            #    copy_paragraph_before(p,doc.add_paragraph("Догмат"))

            if dogmatikos:
                if dogmatikos.runs[0].text.find("Слава")!=-1 and 'gv_doxa' in stichera_matrix[k]:
                    dogmatikos.runs[0].text = dogmatikos.runs[0].text.replace("Слава і","І").replace("Слава, і","І").replace("Слава: І","І")
                copy_paragraph_before(p,dogmatikos)     
            elif datetime(year_no,month_no,k).weekday()+1 in (3,5):
                if  "gv_theotokion" in stichera_matrix[k] and len(stichera_matrix[k]["gv_theotokion"])==2:
                    #print(f"Inserted theo for {k}")
                    copy_paragraph_before(p,stichera_matrix[k]["gv_theotokion"][1])
                else:
                    p.insert_paragraph_before("ХРЕСТОБОГОРОДИЧНИЙ")
            else:
                #print(f"Inserted theo for {k}")
                try:
                    copy_paragraph_before(p,stichera_matrix[k]["gv_theotokion"][0])
                except KeyError:
                    #print("Warning, failed to find stichera_matrix[k]['gv_theotokion'][0]")
                    print(f"Стихири: для {k} не знайдено богородичний ГВ в файлі з стихирами")
                    p.insert_paragraph_before("БОГОРОДИЧНИЙ")
            break         

    doc.save(path)

def insert_kanon(path,date):
    doc = docx.Document(path)
    song8_end_found = False
    week_day = date.weekday()+1
    echos = paschalia.get_echos(date,paschalia_dates)

    #print(f"Inserting kanon for {date.day} {day_dic_reversed[date.weekday()+1]}")

    kanon_found = False
    song8_end_found = False

    for p in doc.paragraphs:
        
        if re.search("Канон",p.text):
            kanon_found = True
            #print("---знайшли початок канону")
            continue
    
        re_result = re.search ("Пісня Богородиці", p.text)
        if re_result and kanon_found:
            #print("---знайшли 9-ту пісню")
            copy_paragraph_list_before(p,templates_kanon_dic[echos][week_day][0][1])
            copy_paragraph_list_before(p,templates_kanon_dic[echos][week_day][0][3])
            copy_paragraph_list_before(p,kanon_litany_list[1])
            copy_paragraph_list_before(p,templates_kanon_dic[echos][week_day][0][4])
            copy_paragraph_list_before(p,templates_kanon_dic[echos][week_day][0][5])
            copy_paragraph_list_before(p,templates_kanon_dic[echos][week_day][0][6])
            copy_paragraph_list_before(p,kanon_litany_list[2])
            copy_paragraph_list_before(p,templates_kanon_dic[echos][week_day][0][7])
            copy_paragraph_list_before(p,templates_kanon_dic[echos][week_day][0][8])
            break
            
        if kanon_found and not song8_end_found:
            delete_paragraph(p)
    
    for p in doc.paragraphs:
        if re.search(r"(Пісня \d|Мала .ктенія)",p.text):
            p.style="Heading 3"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #p.paragraph_format.space_after = Pt(6)

    #print("WARNING: не знайшли закінчення канону!")
    doc.save(path)

old_files = glob.glob('2022/*/*.doc*')
filenames={}

i=0
for f in old_files:
    re_result=re.search(r'2022\\2022.(\d{2})\\(\d{1,2})(?:-|_)(?:.*?)(?:-|_)(.*).docx?',f)
    if re_result:
        i+=1
        #print(i,re_result.group(1),re_result.group(2),re_result.group(3))
        if not int(re_result.group(1)) in filenames.keys():
            filenames[int(re_result.group(1))]={}
        if not int(re_result.group(2)) in filenames[int(re_result.group(1))].keys():
            filenames[int(re_result.group(1))][int(re_result.group(2))]={}
        filenames[int(re_result.group(1))][int(re_result.group(2))]=re_result.group(3)


for month, month_data in filenames.items():
    day_count=0
    for day in month_data.keys():
        day_count+=1
    #print(month, day_count, calendar.monthrange(2023, month)[1])

'''
for month in range(1,13):
    for day_no in range(1,calendar.monthrange(2023, month)[1]+1):
        if not day_no in filenames[month].keys():
            print("Пропущено:",month, day_no)
'''

for m in range(1,13):
    for d, data in filenames[m].items():
        if re.search('Гл', data):
            #print(m,d,data)
            filenames[m][d]=re.sub(r'Гл.\s?\d\s?-\s?','',data)
            #print(m,d,filenames[m][d])


lent_templates = glob.glob('В,У - Пісна Тріодь/*/*.docx')
pascha_pentecost_templates = glob.glob('В,У - Квітна Тріодь/*/*.docx')

ordo_matrix = get_matrix(f"тропарі{mode_suffix2}.csv")
stichera_gv_matrix = get_matrix(f"стихириМінеї{mode_suffix2}.csv")
dismissal_matrix = get_dismissal_matrix(f'Відпусти{mode_suffix}.csv',month_no)
saint_matrix = get_matrix_full("Місяцеслов-БД.csv")
filename_triodion_matrix = get_matrix_full("Місяцеслов-БД-Тріодь.csv")
templates_octoechos_dic = get_octoechos_template_files()
templates_menaion_dic = get_menaion_template_files()
templates_resurrection = get_resurrection_troparia_texts('воскресні.docx')
#templates_menaion = get_menaion_troparia_texts(f'тропарі-{month_no:02}.docx')

templates_menaion = get_menaion_troparia_texts(glob.glob(f'Тропарі - Мінея\\{month_w_offset[month_no-1]:02}-{month_dic_reversed[month_no].upper()}.docx')[0])
if glob.glob(f'тріодь-{month_no:02}.docx'):
    templates_triodion = get_menaion_troparia_texts(glob.glob(f'тріодь-{month_no:02}.docx')[0])

vespers_prokimenon = get_vespers_prokimenon(f'прокімени.docx')
templates_theotokion_dic = get_theotokion_troparia_texts('богородичні-тропарі.docx')
templates_kanon_dic = get_octoechos_kanon_texts('05a_ОКТОЇХ_КАНОНИ.docx')
kanon_litany_list = get_kanon_litany("Канон - Мала єтенія.docx")
stichera_matrix = get_stichera.get_stichera_matrix(glob.glob(f'Стихири - Мінея\\Мінея_{month_no:02}*.docx')[0])
generic_stichera_matrix = get_stichera.get_generic_stichera_matrix('Стихири ГВ загальної служби.docx')
#dogmatika = get_stichera.get_dogmatika("Догмати.docx")


folder_name=f'drafts\\{year_no}-{month_no:02}-{mode}'
if not os.path.exists(folder_name):
    os.makedirs(folder_name)

print(year_no, month_no)

def get_saints_for_filename(month,day,multiline=False):
    index = 10
    separator = ','
    saint_string=""
    for row in saint_matrix[1:]:
        if int(row[0])==month and int(row[1])==day:
            if row[index]:
                if saint_string:
                    saint_string+=separator+row[index]
                else:
                    saint_string=row[index]
    return saint_string



paschalia_dates = paschalia.get_prev_next_pascha(datetime(year_no, month_no,1), mode)

draft_dic={}
for d in range(1,calendar.monthrange(year_no, month_no)[1]+1):
#for d in range(1,5):    
    draft_dic[d]=[]
    #print("Дата:",d,datetime(year_no, month_no, d).weekday()+1,mode)
    day_details = paschalia.get_day_details(datetime(year_no,month_no,d),paschalia_dates)
    #print(d, day_details)
    filename_triodion_row = None
    filename_menaion_string = get_saints_for_filename(month_no,d)

    for row in filename_triodion_matrix:
        if day_details[0] == row[0] and day_details[1] == int(row[1]) and day_details[3] == int(row[3]):
            filename_triodion_row = row
    
    dest_filename=f'{folder_name}\\{d:02}-{day_short_dic_reversed[datetime(year_no, month_no, d).weekday()+1]}'
    
    if datetime(year_no, month_no, d).weekday()+1==7: 
        dest_filename+=f'-Гл.{paschalia.get_echos(datetime(year_no,month_no,d),paschalia_dates)}'
    
    if filename_triodion_row:
        dest_filename+=f'-{filename_triodion_row[12]}.docx'
    elif filename_menaion_string:
        dest_filename+=f'-{filename_menaion_string}.docx'
    else:
        print(f"УВАГА: відсутнє коротке ім'я для {d}")
        dest_filename+=f'-___.docx'
    
    #print(expected_template_path)
    if day_details[0]=='lent':
        expected_lent_template_path = f"В,У - Пісна Тріодь\\тиждень-{day_details[1]}\\{day_details[1]}-{day_details[3]}-ВУ.docx"    
        if expected_lent_template_path in lent_templates and day_details[3]!=7:
            src_filename=expected_lent_template_path
            draft_dic[d].append('піст')
            #print(d, "Шаблон Великого Посту")
    elif day_details[0] in ('pascha'):
        expected_template_path = f"В,У - Квітна Тріодь\\pascha-{day_details[1]}\\{day_details[1]}-{day_details[3]}-ВУ.docx"
        #print(expected_template_path)
        if expected_template_path in pascha_pentecost_templates:
            src_filename=expected_template_path
            draft_dic[d].append('пасха')

    elif day_details[0] == 'pentecost':
        expected_template_path = f"В,У - Квітна Тріодь\\pentecost-{day_details[1]}\\{day_details[1]}-{day_details[3]}-ВУ.docx"
        #print(expected_template_path)
        if expected_template_path in pascha_pentecost_templates:
            src_filename=expected_template_path
            draft_dic[d].append('50-ця')

        elif d in templates_menaion_dic.keys() and datetime(year_no, month_no, d).weekday()+1!=7:
            src_filename=templates_menaion_dic[d]
            draft_dic[d].append('мінея')
            #print(d, "Шаблон Мінеї")
        else:        
            src_filename=templates_octoechos_dic[paschalia.get_echos(datetime(year_no,month_no,d),paschalia_dates)][datetime(year_no, month_no, d).weekday()+1]
            draft_dic[d].append('октоїх')
            #print(d, "Шаблон Октоїха")
    try:
        shutil.copy2(src_filename,dest_filename)
    except NameError as e:
        print(e)
        print("Помилка для дня ", d)
        raise e
    draft_dic[d].append(dest_filename)



#drafts = glob.glob(f'{folder_name}\\*.docx')
for d,desc in draft_dic.items():
    #print("Дата:",d,datetime(year_no, month_no, d).weekday()+1,mode)
    insert_header_from_dismissal_matrix(desc[1],datetime(year_no, month_no, d))
    insert_dismissal(desc[1],datetime(year_no, month_no, d))
    
    

    if desc[0] in ('неділя','октоїх','пасха','50-ця'):
        #вставити стихири ГВ
        #вставити тропарі вечірні
        if ordo_matrix[d][1]!='y':
            #if d==2:
            #    print(f'inserting troparia for {d}')
            insert_troparia(desc[1],datetime(year_no, month_no, d))
        #вставити глас для Бог Господь
        insert_boh_hospod_echos(desc[1],datetime(year_no, month_no, d))
        #вставити тропарі на Бог Господь
        #вставити тропарі вкінці утрені
    if desc[0] in ('октоїх') and (datetime(year_no, month_no, d).weekday()+1)!=7:
        insert_kanon(desc[1],datetime(year_no, month_no, d))
    if desc[0]=='мінея':
        #вставити прокімен
        insert_prokimenon(desc[1], datetime(year_no, month_no, d).weekday()+1)
    else:
        insert_menaion_stichera(desc[1],datetime(year_no, month_no, d),desc[0])    
    if desc[0]=='піст' and datetime(year_no, month_no, d).weekday()+1==6:
        pass

print("Завершено створення чернеток УВ!")
