import docx,os,re
#,glob,re
from datetime import datetime
import ps_docx_utils as pdu
import ps_date_utils as pdt

start_time = datetime.now()

everyday_part_list=['вечірня',
                'господи_взиваю',
                '4_перші_стихи',
                'стихири_ГВ',
                'вхід_седмиця',
                'прокімен_вечірня',
                'сподоби_господи',
                'стиховня_вечірня',
                'нині_відпускаєш',
                'тропарі_седмиця_вечірня',
                'середній_відпуст',
                'тропар_вкінці_вечірня',
                'утреня',
                'тропарі_седмиця_утреня',
                'сідальний1',
                'мала_єктенія_перед_другим_сідальним',
                'сідальний2',
                'псалом_50',
                'після_псалма_50_в_будень',
                'канон',
                'пісня_богородиці',
                'пісня9',
                'достойно',
                #'мала_єктенія_після_канона',
                'світильний',
                'хвалитні_будень',
                #'решта_стихів_хвалитних',
                'прохальна_єктенія_утрені',
                'стиховня_утреня',
                'після_стиховні',
                'потрійна_єктенія_утрені',
                'середній_відпуст',
]
sunday_part_list=['вечірня',
                  'блажен_муж',
                'господи_взиваю',
                'стихири_ГВ',
                'вхід_неділя',
                'прокімен_вечірня',
                'потрійна_єктенія',
                'сподоби_господи',
                'стиховня_вечірня',
                'нині_відпускаєш',
                'тропарі_седмиця_вечірня',
                'великий_відпуст',
                'утреня',
                'тропарі_седмиця_утреня',
                'сідальний1',
                'мала_єктенія_перед_другим_сідальним',
                'сідальний2',
                'полієлей',
                'іпакой',
                'степенна',
                'премудрість_перед_прокіменом_утрені',
                'прокімен_утреня',
                'господеві_помолімся_після_прокімена_утрені',
                'воскресіння_христове',
                'псалом_50',
                'після_псалма_50_в_неділю',
                'канон',
                'пісня_богородиці',
                'пісня9',
                #'мала_єктенія_після_канона',
                'свят_господь_бог_наш',
                'світильний',
                'хвалитні_неділя',
                'стихири_хвалитні',
                'велике_славослов’я',
                'воскресний_тропар',
                'потрійна_єктенія_утрені',
                'прохальна_єктенія_утрені',                
                'великий_відпуст',
]

octoechos_template_source_folder = '01-Октоїх'


def find_octoechos_template(day,echos, folder=octoechos_template_source_folder):

    #tmp
    hyphen = '' if day==7 else '-'

    path = folder+f'\\Глас_{echos}\\{echos}-{pdt.day_short_dic_reversed[day]}{hyphen}.docx'
    #print(path)
    if os.path.exists(path):
        return path
    else:
        return -1


list_of_lines_that_end_parts = ['Світло тихе',
                                'Вхід', 
                                'Сподоби, Господи', #кінець прокімена седмичного дня
                                'Господеві помолімся.', #кінець прокімена неділі
                                'Пісня Симеона', #кінець стиховні вечірні
                                "Утреня", #кінець вечірні
                                "УТРЕНЯ", #кінець вечірні
                                "Тропарі", #кінець початку утрені
                                'Відтак читець читає другу чергову катизму, диякон виголошує другу малу єктенію:',
                                'Полієлей', #кінець другого сідального
                                'Степенна', #кінець іпакоя
                                'Євангеліє', #кінець степенної
                                'Господеві помолімся.', #кінець прокімна
                                #'Мала єктенія',
                                'Псалом 50',
                                'Пісня Богородиці',#кінець канону 1-8
                                'Пісня 9',#кінець Пісні Богородиці
                                'Достойно',
                                '«Свят Господь Бог',
                                'Псалом 148', #кінець світильного
                                'Хвалитні', #кінець світильного
                                "Велике славослов'я",  #кінець хвалитних стихир
                                'Після стиховні читаємо:', #кінець стиховні утрені седмичного дня
                                "Тропар", #кінець Великого славослів'я
                                "Єктенія усильного благання", #кінець воскресного тропаря
                                

]

'''
dic_of_var_parts = {
    'стихири_ГВ':'Стихири',
    'стихири_ГВ':'Прокімен',
}
'''

def get_vu_octoechos_variable_parts_from_template(day,echos):
    path = find_octoechos_template(day, echos)
    doc = docx.Document(path)
    matrix = []
    service = None
    key=None
    for p in doc.paragraphs:
        
        if p.text.lower() == 'вечірня':
            service = 'вечірня'
        if p.text.lower() == 'утреня':
            service = 'утреня'

        if key:
            for item in list_of_lines_that_end_parts:
                if p.text.startswith(item):
                    key= None
                    continue

        #SPAGHETTI CODE. Redo later
        if not key and p.text =='Стихири':
            
            key = 'стихири_ГВ' if service == 'вечірня' else 'стихири_хвалитні'
            matrix.append({'key':key,'text':[p]})
            continue

        if not key and p.text.startswith('Прокімен'):
            key = f'прокімен_{service}'
            matrix.append({'key':key,'text':[p]})
            continue

        if not key and p.text.startswith('Стихири на стиховні'):
            key = f'стиховня_{service}'
            matrix.append({'key':key,'text':[p]})
            continue

        if not key and p.text.startswith('Після Вечірні'):
            key = f'тропар_вкінці_{service}'
            matrix.append({'key':key,'text':[p]})
            continue

        if not key and p.text.startswith('Перший сідальний'):
            key = f'сідальний1'
            matrix.append({'key':key,'text':[p]})
            continue
        
        if not key and p.text.startswith('Другий сідальний'):
            key = f'сідальний2'
            matrix.append({'key':key,'text':[p]})
            continue
        
        if not key and p.text.startswith('Іпакой'):
            key = f'іпакой'
            matrix.append({'key':key,'text':[p]})
            continue

        if not key and p.text.startswith('Степенна пісня'):
            key = f'степенна'
            matrix.append({'key':key,'text':[p]})
            continue
        
        if not key and p.text.startswith('Канон'):
            key = f'канон'
            matrix.append({'key':key,'text':[p]})
            continue


        if not key and p.text.startswith('Пісня 9'):
            key = f'пісня9'
            matrix.append({'key':key,'text':[p]})
            continue

        if not key and p.text.startswith('Світильний'):
            key = f'світильний'
            matrix.append({'key':key,'text':[p]})
            continue
    
        if not key and p.text==('Тропар'):
            key = f'воскресний_тропар'
            matrix.append({'key':key,'text':[p]})
            continue


        #!!!!!!!! Чи це працюватиме, з послідовними змінними частинами???

        

        if key:
            matrix[-1]['text'].append(p)

    return matrix        

def get_vu_misc_variable_parts(path):
    doc = docx.Document(path)
    matrix = []
    key=key2=None


    for p in doc.paragraphs:
        if p.style.name == 'Heading 1':
            key = p.text
            continue
        if p.style.name == 'Heading 2':
            key2 = p.text
            matrix.append({'key':key,'key2':key2,'text':[]})
            continue
    
        matrix[-1]['text'].append(p)
    return matrix


def get_vu_template_parts(path):
    doc = docx.Document(path)
    matrix = []
    key=None
    for p in doc.paragraphs:
        if '#' in p.text:
            if p.text[1:]=='кінець':
                break
            
            key = p.text[1:]
            matrix.append({'key':key, 'text':[] })
            continue
        
        else:
            matrix[-1]['text'].append(p)
    return matrix

def get_template_part_text(matrix, key):
    for item in matrix:
        if item['key']==key:
            return item['text']
    
    return -1

def build_template(path,day,echos):
    
    octoechos_texts =get_vu_octoechos_variable_parts_from_template(day,echos)

    doc = docx.Document()
    partlist = sunday_part_list if day==7 else everyday_part_list

    try:
        for item in partlist:
            
            #p = doc.add_paragraph(f'#{item}')
            #pdu.format_line(p,handle='ri')

            #octoechos_texts = list(filter(lambda d: d['echos']==echos and d['day']==day ,vu_octoechos_parts))[0]['texts']
            #if get_template_part_text(octoechos_texts, item) != -1 :
            #    pdu.copy_paragraph_list(doc, get_template_part_text(octoechos_texts, item))
            
            if get_template_part_text(octoechos_texts, item) != -1 :
                pdu.copy_paragraph_list(doc, get_template_part_text(octoechos_texts, item))

            else: 
                #octoechos_texts = list(filter(lambda d: d['echos']==echos and d['day']==day ,vu_octoechos_parts))[0]['texts']
                #octoechos_part = list(filter(lambda p: p['key']==item,octoechos_texts))[0]
                pdu.copy_paragraph_list(doc, get_template_part_text(vu_template_parts, item))

    except TypeError as e:
        print(f"Warning: {item} not found")
        print(get_template_part_text(vu_template_parts, item))
        print(e)
        raise e
    
    insert_echos_into_description(doc, "Два перші стихи 140-го")
    
    if day == 7:
        insert_echos_into_description(doc, "Бог Господь... (Пс. 117)")
        insert_echos_into_description(doc, "Два перших стихи хвалитних")

    doc.save(path)


def build_octoechos_part_matrix_full(source_folder_path):
    matrix = []
    
    #all_files = glob.glob(f'{source_folder_path}/*/*.docx')
    for echos in range(1,9):
        for day in range(1,8):

            texts = get_vu_octoechos_variable_parts_from_template(day,echos)
            matrix.append({'echos':echos,'day':day,'texts':texts})
    return matrix


#Заміна № гласу в наступних строках:
#Два перші стихи 140-го псалма, на глас ?.
#Бог Господь... (Пс. 117), глас ?
#Два перших стихи хвалитних на глас ?:

def insert_echos_into_description(doc, text):
    echos = None
    paragraph_to_update = None
    for p in doc.paragraphs:
        if p.text.startswith(text):
            paragraph_to_update = p

        if paragraph_to_update:
            result = re.search(r'\(г. (\d)\)',p.text)
            if result:
                echos = result.group(1)
                break
    if paragraph_to_update and echos:
        for r in paragraph_to_update.runs:
            if '?' in r.text:
                r.text = r.text.replace('?',echos)
    else:
        print(text, 'NO PARGARAPH TO Update')

folder_for_new_files = '01-Октоїх-new'


def build_full_octoechos(folder):
    for echos in range(1,9):
        for day in range(1,8):
            doc_filename = f"{folder}/Глас_{echos}/{echos}-{pdt.day_short_dic_reversed[day]}-.docx"
            build_template(doc_filename, day,echos)
            print(f"{datetime.now()}: Шаблон {doc_filename} побудовано!")



vu_template_parts = get_vu_template_parts('vu_template_parts.docx')
vu_misc_variable_parts = get_vu_misc_variable_parts('vu_template_misc_variable_parts.docx')

#vu_octoechos_parts = build_octoechos_part_matrix_full('01-Октоїх')

print(f"Finished building dictionaries: {(datetime.now() - start_time).total_seconds()}")

if __name__ == "__main__":
    filename = 'test.docx'
    #build_template(filename, 7,1)
    #print(f"Шаблон {filename} побудовано!")

    build_full_octoechos('01-Октоїх-new')

    end_time = datetime.now()
    elapsed_time = (end_time - start_time).total_seconds()
    print(f"Elapsed time : {elapsed_time}") 