import re,glob,calendar,docx,os,easygui, shutil, csv
from datetime import datetime
from docx.shared import RGBColor

#filenames= glob.glob('*/*/*.txt')
mode = easygui.choicebox('u - Юліанський, g - Григоріанський', 'Вибір календаря', ['u','g'])

old_files = glob.glob('2022/*/*.doc*')
filenames={}


month_no = datetime.now().month+1 if datetime.now().month!=12 else 1
year_no = datetime.now().year if datetime.now().month!=12 else datetime.now().year+1
month_no=11

day_short_dic={"ПН":1,
        "ВТ":2,
        "СР":3,
        "ЧТ":4,
        "ПТ":5,
        "СБ":6,
        "НД":7}
day_short_dic_reversed = {v:k for k,v in day_short_dic.items()}
day_short_dic_string='('+'|'.join(day_short_dic.keys())+')'

day_dic = {"Понеділок":1,
            "Вівторок":2,
            "Середа":3,
            "Четвер":4,
            "П'ятниця":5,
            "Субота":6,
            "Неділя":7}
day_dic_reversed = {v:k for k,v in day_dic.items()}
day_dic_string='('+'|'.join(day_dic.keys())+')'

month_dic= {'Січень':1,
              'Лютий':2,
              'Березень':3,
              'Квітень':4,
              'Травень':5,
              'Червень':6,
              'Липень':7,
              'Серпень':8,
              'Вересень':9,
              'Жовтень':10,
              'Листопад':11,
              'Грудень':12}
month_dic_reversed = {v:k for k,v in month_dic.items()}
month_dic_string='('+'|'.join([x.lower() for x in month_dic.keys()])+')'


if mode == 'u':
    mode_suffix='Юл'
elif mode == 'g':
    mode_suffix='Гр'

i=0
for f in old_files:
    re_result=re.search('2022\\\\2022.(\d{2})\\\\(\d{1,2})(?:-|_)(?:.*?)(?:-|_)(.*).docx?',f)
    if re_result:
        i+=1
        #print(i,re_result.group(1),re_result.group(2),re_result.group(3))
        if not int(re_result.group(1)) in filenames.keys():
            filenames[int(re_result.group(1))]={}
        if not int(re_result.group(2)) in filenames[int(re_result.group(1))].keys():
            filenames[int(re_result.group(1))][int(re_result.group(2))]={}
        filenames[int(re_result.group(1))][int(re_result.group(2))]=re_result.group(3)


for month, month_data in filenames.items():
    day_count=0
    for day in month_data.keys():
        day_count+=1
    #print(month, day_count, calendar.monthrange(2023, month)[1])

'''
for month in range(1,13):
    for day_no in range(1,calendar.monthrange(2023, month)[1]+1):
        if not day_no in filenames[month].keys():
            print("Пропущено:",month, day_no)
'''

for m in range(1,13):
    for d, data in filenames[m].items():
        if re.search('Гл', data):
            #print(m,d,data)
            filenames[m][d]=re.sub('Гл.\s?\d\s?-\s?','',data)
            #print(m,d,filenames[m][d])

def get_echos(date,mode):
    if mode=='u':
        return (int(date.strftime("%U"))-25) % 8 + 1
    elif mode=='g':
        return (int(date.strftime("%U"))-24) % 8 + 1
'''
def copy_paragraph_before(p_to_insert_before,source_paragraph):
    target_paragraph = p_to_insert_before.insert_paragraph_before()
    target_paragraph.style = source_paragraph.style
    target_paragraph.alignment = source_paragraph.alignment
    for run in source_paragraph.runs:
        new_run = target_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.font.color.rgb = run.font.color.rgb
        new_run.font.highlight_color = run.font.highlight_color
'''

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def delete_run(run):
    p = run._element
    p.getparent().remove(p)
    p._p = p._element = None

def copy_run(target_paragraph,run):
    new_run = target_paragraph.add_run(run.text)
    new_run.style = run.style.name
    new_run.bold = run.bold
    new_run.italic = run.italic
    new_run.underline = run.underline
    new_run.font.size = run.font.size
    new_run.font.name = run.font.name
    new_run.font.color.rgb = run.font.color.rgb
    new_run.font.highlight_color = run.font.highlight_color

def copy_paragraph_before(paragraph_to_insert_before,source_paragraph):
    target_paragraph = paragraph_to_insert_before.insert_paragraph_before()
    try:
        target_paragraph.style = source_paragraph.style
    except KeyError:
        print(f"Warning. Text {source_paragraph.text[:20]} has style{source_paragraph.style}")
    for run in source_paragraph.runs:
        copy_run(target_paragraph,run)
        '''
        new_run = target_paragraph.add_run(run.text)
        new_run.style = run.style.name
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.font.color.rgb = run.font.color.rgb
        new_run.font.highlight_color = run.font.highlight_color
        '''
    return target_paragraph

def copy_paragraph_list_before(p_to_insert_before,paragraph_list):
    #print("qty of p to insert:",len(paragraph_list))
    for p in paragraph_list:
        copy_paragraph_before(p_to_insert_before,p)




def copy_paragraph(target_doc,source_paragraph):
    target_paragraph = target_doc.add_paragraph()
    try:
        target_paragraph.style = source_paragraph.style
    except KeyError:
        print(f"Warning. Text {source_paragraph.text[:20]} has style{source_paragraph.style}")
    target_paragraph.alignment = source_paragraph.alignment
    for run in source_paragraph.runs:
        new_run = target_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.font.color.rgb = run.font.color.rgb
        new_run.font.highlight_color = run.font.highlight_color

def copy_paragraph_list(target_doc,paragraph_list):
    for p in paragraph_list:
        copy_paragraph(target_doc,p)

def get_octoechos_template_files():
    template_dic = {}
    all_files = glob.glob('01-Октоїх/*/*.docx')
    for f in all_files:
        re_result=re.search('01-Октоїх\\\\Глас_(\d)\\\\(\d)-(\w{2})-?.docx',f)
        if re_result:
            #print(re_result.group(2), re_result.group(3))
            doc = docx.Document(f)
            if int(re_result.group(2)) not in template_dic.keys():
                template_dic[int(re_result.group(2))]={}
            if re_result.group(3) in day_short_dic.keys():
                template_dic[int(re_result.group(2))][day_short_dic[re_result.group(3)]]=f
            else:
                print("Something wrong with", f, re_result.group(3))
        else:
            print("None found for:", f)
    return template_dic


def get_octoechos_templates():
    template_dic = {}
    all_files = glob.glob('01-Октоїх/*/*.docx')
    for f in all_files:
        re_result=re.search('01-Октоїх\\\\Глас_(\d)\\\\(\d)-(\w{2})-?.docx',f)
        if re_result:
            #print(re_result.group(2), re_result.group(3))
            doc = docx.Document(f)
            if int(re_result.group(2)) not in template_dic.keys():
                template_dic[int(re_result.group(2))]={}
            if re_result.group(3) in day_dic.keys():
                template_dic[int(re_result.group(2))][day_dic[re_result.group(3)]]=doc.paragraphs
            else:
                print("Something wrong with", f, re_result.group(3))
        else:
            print("None found for:", f)
    return template_dic

def get_menaion_template_files():
    filenames = glob.glob(f'В,У - Мінея/{month_no:02}*/*.docx')
    template_dic={}
    for f in filenames:
        #print("Checking",f)
        pattern='(\d{2})-__-(.*?).docx'
        #print(pattern)
        re_result = re.search(pattern,f)
        doc = docx.Document(f)
        template_dic[int(re_result.group(1))]=f
    
    return template_dic

def get_menaion_templates():
    filenames = glob.glob(f'В,У - Мінея/{month_no:02}*/*.docx')
    template_dic={}
    for f in filenames:
        print("Checking",f)
        pattern='(\d{2})-__-(.*?).docx'
        print(pattern)
        re_result = re.search(pattern,f)
        doc = docx.Document(f)
        template_dic[int(re_result.group(1))]=doc.paragraphs
    return template_dic

def get_resurrection_troparia_texts(path):
    template_dic={}
    doc = docx.Document(path)
    echos = None
    #i=0
    for p in doc.paragraphs:
        re_result = re.search("Глас (\d)",p.text)
        if re_result:
            #print("found")
            echos=int(re_result.group(1))
            template_dic[echos]=[]
        #print(p.text)

        re_result = re.search("((Тропар|Кондак).*?\(г\. \d\)): (.*)",p.text)
        if echos and re_result:
            #template_dic[echos].append([re_result.group(1),re_result.group(3)])
            template_dic[echos].append(p)
    return template_dic

def get_menaion_troparia_texts(path):
    #print("in")
    template_dic={}
    doc = docx.Document(path)
    key = None #cur_day
    saint={}
    for p in doc.paragraphs:
        re_result = re.search(f"^(\d+) (.*)",p.text)
        if re_result:
            #print("found")
            key=int(re_result.group(1))
            cur_day_heading = re_result.group(2)
            template_dic[key]=[]

        if key:
            if re.search("^Тропар",p.text):
                re_result = re.search("(Тропар.*?\(г\. \d\): )(.*)",p.text)
                try:
                    #saint["troparion"]=[re_result.group(1), re_result.group(2)]
                    saint["troparion"]=p
                except IndexError as e:
                    print(p.text, re_result)
                    raise e
            elif re.search("^Кондак",p.text):
                re_result = re.search("(Кондак.*?\(г\. \d\): )(.*)",p.text)
                #saint["kondakion"]=[re_result.group(1), re_result.group(2)]
                saint["kondakion"]=p
            else:
                #found saint header
                saint["header"]=p.text    

        if "header" in saint.keys() and "troparion" in saint.keys() and "kondakion" in saint.keys():
            template_dic[key].append(saint)
            saint={}
        #    template_dic[echos].append([re_result.group(1),re_result.group(3)])
    return template_dic

def get_vespers_prokimenon(path):
    template_dic={}
    doc = docx.Document(path)
    day=None
    for p in doc.paragraphs:
        re_result = re.search(f"^{day_dic_string}",p.text)
        if re_result:
            day=re_result.group(1)
            template_dic[day_dic[day]]=[]
        elif day:
            template_dic[day_dic[day]].append(p)
    return template_dic

def get_matrix_full(csv_filename):
    matrix=[]
    with open(csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:
            matrix.append(row)
    return matrix

def get_matrix(csv_filename):
    matrix={}
    with open(csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:

            if row[0]==str(month_no):
                matrix[int(row[1].split('.')[2])]=row[2:]
    return matrix

#need to merge with previous one
def get_dismissal_matrix(dismissal_csv_filename,cur_month):
    matrix={}
    with open(dismissal_csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:
            if row[0]==month_dic_reversed[cur_month]:
                matrix[int(row[1].split('.')[0])]=row[2:]
    return matrix



def get_theotokion_troparia_texts(path):
    service_dic = {'Вечірня':'vespers',
                   'Утреня':'orthros'}
    template_dic={}
    doc = docx.Document(path)
    for p in doc.paragraphs:
        re_result = re.search("Глас (\d)",p.text)
        if re_result:
            #print("found")
            echos=int(re_result.group(1))
            template_dic[echos]=[]
        #print(p.text)

        re_result = re.search(f"(Вечірня|Утреня) Відпуст",p.text)
        if re_result:
            #print(echos,re_result.group(1))
            service=service_dic[re_result.group(1)]

        re_result = re.search(f"^{day_dic_string}",p.text)
        if re_result:
            day = re_result.group(1)
            weekday_no = day_dic[day]

        re_result = re.search(f"Богородичний",p.text)
        if re_result:
            template_dic[echos].append({'service':service, 'weekday':weekday_no,'p':p})
            
    return template_dic


def get_kanon_litany(path):
    template_dic={}
    doc = docx.Document(path)
    litany_no = None
    for p in doc.paragraphs:
        re_result = re.search("Мала єктенія (\d)",p.text)
        if re_result:
            litany_no=int(re_result.group(1))
            template_dic[litany_no]=[]
            continue
        elif litany_no:
            template_dic[litany_no].append(p)
    return template_dic
        

def get_kanon_texts(path):
    template_dic={}
    doc = docx.Document(path)
    
    kanon_found=False
    kanon_end_found=False
    for p in doc.paragraphs:
        re_result = re.search("Глас (\d)",p.text)
        if re_result:
            #print("found")
            echos=int(re_result.group(1))
            template_dic[echos]={}
        #print(p.text)
        re_result = re.search(f"^{day_dic_string}",p.text)
        if re_result:
            day = re_result.group(1)
            weekday_no = day_dic[day]
            template_dic[echos][weekday_no]=[]

        re_result = re.search(f"Канон*(П)",p.text)
        if re_result:
            kanon_found=True
            kanon_end_found=False
            template_doc[echos]=[]
            
        re_result = re.search(f"Пісня 9",p.text)
        if re_result:
            kanon_found = False
            kanon_end_found = True
            
        if kanon_found and not kanon_end_found:
            template_doc[echos].append(p)
            
    return template_dic

ordo_matrix = get_matrix("тропарі.csv")
dismissal_matrix = get_matrix(f'Відпусти{mode_suffix}.csv')
saint_matrix = get_matrix_full("Місяцеслов-БД.csv")
templates_octoechos_dic = get_octoechos_template_files()
templates_menaion_dic = get_menaion_template_files()
templates_resurrection = get_resurrection_troparia_texts('воскресні.docx')
templates_menaion = get_menaion_troparia_texts(f'тропарі-{month_no}.docx')
vespers_prokimenon = get_vespers_prokimenon(f'прокімени.docx')
templates_theotokion_dic = get_theotokion_troparia_texts('богородичні-тропарі.docx')
templates_kanon_dic = get_kanon_texts('05a_ОКТОЇХ_КАНОНИ.docx')
litany_list = get_kanon_litany("Канон - Мала єтенія.docx")



folder_name=f'drafts\\{year_no}-{month_no}-{mode}'
if not os.path.exists(folder_name):
    os.makedirs(folder_name)

print(year_no, month_no)



draft_dic={}
for d in range(1,calendar.monthrange(year_no, month_no)[1]+1):
    draft_dic[d]=[]
    #print(d,datetime(year_no, month_no, d).weekday()+1)
    dest_filename=f'{folder_name}\\{d:02}-{day_short_dic_reversed[datetime(year_no, month_no, d).weekday()+1]}'
    if datetime(year_no, month_no, d).weekday()+1==7:
        #echos=
        dest_filename+=f'-Гл.{get_echos(datetime(year_no,month_no,d),mode)}'
        #draft_dic[d].append('неділя')
    dest_filename+=f'-{filenames[month_no][d]}.docx'
    if d in templates_menaion_dic.keys() and datetime(year_no, month_no, d).weekday()+1!=7:
        src_filename=templates_menaion_dic[d]
        draft_dic[d].append('мінея')
    else:        
        src_filename=templates_octoechos_dic[get_echos(datetime(year_no,month_no,d),mode)][datetime(year_no, month_no, d).weekday()+1]
        draft_dic[d].append('октоїх')
    shutil.copy(src_filename,dest_filename)
    draft_dic[d].append(dest_filename)



def insert_prokimenon(path,day):
    doc = docx.Document(path)
    #print(path)
    vespers_prokimenon_found = False
    vespers_prokimenon_end_found = True
    for p in doc.paragraphs:
        re_result = re.search("Прокімен",p.text)
        if re_result:
            vespers_prokimenon_found = True
            #delete_paragraph(p)
        re_result = re.search("(Читання|Сподоби, Господи)",p.text)
        if re_result:
            #print(day)
            for p1 in vespers_prokimenon[day]:
                copy_paragraph_before(p,p1)
            p.insert_paragraph_before()
            break
        elif vespers_prokimenon_found:
            #print(f"Deleting {p.text}")
            delete_paragraph(p)
    doc.save(path)

def insert_boh_hospod_echos(path,date):
    #print(date.day)
    doc = docx.Document(path)
    for p in doc.paragraphs:
        re_result = re.search("Бог Господь",p.text)
        if re_result:
            #print("found")
            if date.weekday()+1==7:
                echos = get_echos(date,mode)
            else:
                #echos = int(re.search("(\d)",templates_menaion[date.day][0]['troparion'][0]).group(1))
                echos = int(re.search("(\d)",templates_menaion[date.day][0]['troparion'].text).group(1))

            for r in p.runs:
                if '???' in r.text:
                    #print(r.text)
                    r.font.highlight_color=None
                    r.text = r.text.replace('???',str(echos))
                    
            break
    doc.save(path)

def insert_prefix_to_paragraph(p,text="Слава: ",formatting='b'):
    p_new = copy_paragraph_before(p,p)
    new_run = p_new.add_run(text)
    if 'b' in formatting:
        new_run.font.bold = True
    new_run.font.name='Times New Roman'
    new_run.font.size=152400
    
    bkp=p_new.runs[:-1]
    for r in p_new.runs:
        delete_run(r)
        
    copy_run(p_new,new_run)
    for r in bkp:
        copy_run(p_new,r)
    return p_new




def get_troparia_theotokion(echos, date, service):
    #print(echos, date, service)
    res = filter(lambda t: t['service'] == service and t['weekday'] == date.weekday()+1, templates_theotokion_dic[echos])
    l = list(res)
    p = l[0]["p"]
    return p
    #return list(filter(lambda t: t['service'] == service and t['weekday'] == date.weekday()+1, templates_theotokion_dic[echos]))[0]["p"]

def get_troparia_block(date,service):
    troparia_block=[]
    
    if ordo_matrix[date.day][1]=='y':
        print("Something wrong")
        raise Exception("Should have a ready template")
    saint_counter=0
    
    for i in list(range(4)):
        #if date.day==5:
            #print("i",i)
        if not ordo_matrix[date.day][2:][i]:
            continue
        if ordo_matrix[date.day][2:][i] == 'resurrection':
            troparia = templates_resurrection[get_echos(date,mode)][0]
        if ordo_matrix[date.day][2:][i] == 'saint':
            #print(date.day, i)
            if i==2:
                #print("adding Slava")
                troparia=insert_prefix_to_paragraph(templates_menaion[date.day][saint_counter]["troparion"],text="Слава: ")
                #print(troparia.text)
            else:
                troparia=templates_menaion[date.day][saint_counter]["troparion"]
            #troparia_block.append(troparia)
            saint_counter+=1
        #if i==3 and ordo_matrix[date.day][2:][i] == 'saint':
        #    troparia_block.append(templates_resurrection[date.day][troparion_count]["troparion"])
        if ordo_matrix[date.day][2:][i] == 'theotokos':
            #echos = int(re.search("(\d)",templates_resurrection[date.day][saint_counter]["troparion"]).group(1))
            #print(troparia_block[-1], troparia_block[-1].text)
            echos = int(re.search("(\d)",troparia_block[-1].text).group(1))
            if ordo_matrix[date.day][2:][2]:
                #print("І нині")
                troparia=insert_prefix_to_paragraph(get_troparia_theotokion(echos,date,service),text="І нині: ")
                #print("to insert",troparia.text)
            else:
                #print("Слава і нині")
                troparia=insert_prefix_to_paragraph(get_troparia_theotokion(echos,date,service),text="Слава, і нині: ")
                #print("to insert",troparia.text)
        #if date.day==5:
            #print("inserting",i,ordo_matrix[date.day])
            #print(ordo_matrix[date.day][2:][i])
            #print(troparia.text)        
        troparia_block.append(troparia)
    troparia_block.append(troparia_block[-1].insert_paragraph_before())        
    return troparia_block

def insert_troparia(path,date):
    #print(date.day)
    doc = docx.Document(path)
    troparia_block = {"orthros": get_troparia_block(date,'orthros'),
                      "vespers": get_troparia_block(date,'vespers')}
    #troparia_block=get_troparia_block(date,'orthros')
    troparion_found = False
    troparion_end_found = False
    service = False
    for p in doc.paragraphs:
        if re.search("ВЕЧІРНЯ",p.text):
            service = 'vespers'
        if re.search("УТРЕНЯ",p.text):
            service = 'orthros'
        
        re_result = re.search("Тропарі",p.text)
        if re_result:
            troparion_found = True
            troparion_end_found = False
            continue
        
        re_result = re.search("(Великий відпуст|Єктенія усильного благання|Мала єктенія)",p.text)
        if re_result and troparion_found:
            troparion_end_found=True
            copy_paragraph_list_before(p,troparia_block[service])
            troparion_found=False
        elif  troparion_found and not troparion_end_found:
            delete_paragraph(p)
            #continue
    doc.save(path)          
def format_line(p, handle=''):
    #handle = "bir"
    if 'b' in handle:
        p.runs[0].font.bold = True
    if 'i' in handle:
        p.runs[0].font.italic = True
    if 'r' in handle:
        p.runs[0].font.color.rgb = RGBColor(0xff, 0x44, 0x00)
    p.runs[0].font.name='Times New Roman'
    p.runs[0].font.size=152400

def insert_header(path,date):
    #print(date)
    doc = docx.Document(path)
    for p in doc.paragraphs:
        re_result = re.search("ВЕЧІРНЯ",p.text)
        if not re_result:
            delete_paragraph(p)
        else:
            #Шапка в форматі "Травень 10 Середа"
            day_name = day_dic_reversed[date.weekday()+1]
            month_name = month_dic_reversed[month_no]
            txt =" ".join([month_name,str(date.day),day_name])
            p_new = p.insert_paragraph_before(txt)
            format_line(p_new, '')

            #TODO: Субота, Неділя, Тиждень etc...
            if date.weekday()+1 in [6,7]:
                p_new = p.insert_paragraph_before(day_name+" ЯКАСЬ ТАМ")
                format_line(p_new, '')

            #TODO: перелік святих
            lst = filter(lambda l:int(l[0])==date.month and int(l[1])==date.day,saint_matrix[1:])
            for l in lst:
                p_new=p.insert_paragraph_before(l[9])
                format_line(p_new, ''.join(l[2:5]))
            break


            
    doc.save(path)


#drafts = glob.glob(f'{folder_name}\\*.docx')
for d,desc in draft_dic.items():
    insert_header(desc[1],datetime(year_no, month_no, d))
    if desc[0]=='неділя' or desc[0]=='октоїх':
        pass
        #вставити стихири ГВ
        #вставити тропарі вечірні
        if ordo_matrix[d][1]!='y':
            insert_troparia(desc[1],datetime(year_no, month_no, d))
        #вставити глас для Бог Господь
        insert_boh_hospod_echos(desc[1],datetime(year_no, month_no, d))
        #вставити тропарі на Бог Господь
        #вставити тропарі вкінці утрені
    if desc[0]=='мінея':
        #вставити прокімен
        insert_prokimenon(desc[1], datetime(year_no, month_no, d).weekday()+1)
        pass

    

