import docx,re

day_dic = {"Понеділок":1,
            "Вівторок":2,
            "Середа":3,
            "Четвер":4,
            "П'ятниця":5,
            "Субота":6,
            "Неділя":7}
day_dic_reversed = {v:k for k,v in day_dic.items()}
day_dic_string='('+'|'.join(day_dic.keys())+')'

def get_kanon_litany(path):
    template_dic={}
    doc = docx.Document(path)
    litany_no = None
    for p in doc.paragraphs:
        re_result = re.search("Мала єктенія (\d)",p.text)
        if re_result:
            litany_no=int(re_result.group(1))
            template_dic[litany_no]=[]
            continue
        elif litany_no:
            template_dic[litany_no].append(p)
    return template_dic
        

def get_kanon_texts(path):
    template_dic={}
    doc = docx.Document(path)
    
    kanon_found=False
    kanon_end_found=False
    i=0
    for p in doc.paragraphs:
        p.text = re.sub('П.ятниця',"П'ятниця",p.text)
        if i<20:
            #print(kanon_found, kanon_end_found, p.text[:40])
            pass
        re_result = re.search("^Глас (\d)",p.text)
        if re_result:
            print(i, "found echos",re_result.group(1), p.text[:40])
            echos=int(re_result.group(1))
            if echos not in template_dic.keys():
                template_dic[echos]={}
        #print(p.text)
        re_result = re.search(f"^{day_dic_string}",p.text)
        if re_result:
            day = re_result.group(1)
            print(echos,"found weekday",re_result.group(1))
            weekday_no = day_dic[day]
            template_dic[echos][weekday_no]=[]
            print(i, "created list stub for", echos, weekday_no,template_dic[echos][weekday_no])

        re_result = re.search(f"^Канон.*\(П\)",p.text)
        if re_result:
            kanon_found=True
            kanon_end_found=False
            
        re_result = re.search(f"Пісня 9",p.text)
        if re_result:
            kanon_found = False
            kanon_end_found = True
            
        if kanon_found and not kanon_end_found:
            #print(i, "attempting to append", echos, weekday_no)
            template_dic[echos][weekday_no].append(p)
        i+=1    
    return template_dic


templates_kanon_dic = get_kanon_texts('05a_ОКТОЇХ_КАНОНИ.docx')
litany_list = get_kanon_litany("Канон - Мала єтенія.docx")
