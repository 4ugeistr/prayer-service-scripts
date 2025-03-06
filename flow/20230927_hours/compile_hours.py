import re, docx, csv, easygui, glob, os, calendar, shutil
from datetime import datetime
import paschalia
from docx.shared import RGBColor, Pt
RGB_RED = RGBColor(0xff, 0x00, 0x00)

mode = easygui.choicebox('u - Юліанський, g - Григоріанський', 'Вибір календаря', ['u','g'])
mode_dic = {'НЮ':'u',
            'ГР':'g',}

if mode == 'u':
    mode_suffix='Юл'
    mode_suffix2='НЮ'
elif mode == 'g':
    mode_suffix='Гр'
    mode_suffix2='ГР'

mode_dic_reversed = {v:k for k,v in mode_dic.items()}

month_no = datetime.now().month+1 if datetime.now().month!=12 else 1
year_no = datetime.now().year if datetime.now().month!=12 else datetime.now().year+1

month_no = 3
print("WARNING. Month_no OVERRIDE", month_no)

month_dic= {'Січень':1,
              'Лютий':2,
              'Березень':3,
              'Квітень':4,
              'Травень':5,
              'Червень':6,
              'Липень':7,
              'Серпень':8,
              'Вересень':9,
              'Жовтень':10,
              'Листопад':11,
              'Грудень':12}
month_dic_reversed = {v:k for k,v in month_dic.items()}
month_dic_string='('+'|'.join([x.lower() for x in month_dic.keys()])+')'
month_w_offset = list(month_dic.values())[4:]+list(month_dic.values())[:4]


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

'''
def get_echos(date,mode):
    if mode=='u':
        return (int(date.strftime("%U"))-25) % 8 + 1
    elif mode=='g':
        return (int(date.strftime("%U"))-24) % 8 + 1
'''

def get_matrix(csv_filename):
    matrix={}
    with open(csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:
            if row[0]==str(month_no):
                matrix[int(row[1].split('.')[2])]=row[2:]
    return matrix


def copy_paragraph(target_doc,source_paragraph):
    target_paragraph = target_doc.add_paragraph()
    for run in source_paragraph.runs:
        new_run = target_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.font.color.rgb = run.font.color.rgb
        new_run.font.highlight_color = run.font.highlight_color

def copy_paragraph_before(paragraph,source_paragraph):
    target_paragraph = paragraph.insert_paragraph_before()
    for run in source_paragraph.runs:
        new_run = target_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.font.color.rgb = run.font.color.rgb
        new_run.font.highlight_color = run.font.highlight_color


def get_resurrection_template_texts(path):
    template_dic={}
    doc = docx.Document(path)
    #rubric_found = cur_glas = None
    #i=0
    key = None
    for p in doc.paragraphs:
        re_result = re.search(r"Глас (\d)", p.text)
        if re_result:
            key=int(re_result.group(1))
            if not key in template_dic:
                template_dic[key]=[]
        if key and p.text.find("Тропар")!=-1:
            template_dic[key].append(p)
        if key and p.text.find("Кондак")!=-1:
            template_dic[key].append(p)
    return template_dic

def get_menaion_template_texts(path):
    template_dic={}
    doc = docx.Document(path)
    key = None
    for p in doc.paragraphs:
        re_result = re.search(r"^(\d{1,2})",p.text)
        if re_result:
            key = int(re_result.group(1))
            if not key in template_dic:
                template_dic[key]=[]
        if key and p.text.find("Тропар")!=-1:
            template_dic[key].append(p)
        if key and p.text.find("Кондак")!=-1:
            template_dic[key].append(p)
    return template_dic

def get_feast_template_texts(path):
    template_dic={}
    try:
        doc = docx.Document(path)
    except docx.opc.exceptions.PackageNotFoundError:
        return template_dic
    key = None
    for p in doc.paragraphs:
        re_result = re.search(r"^(\d{1,2})",p.text)
        if re_result:
            key = int(re_result.group(1))
            if not key in template_dic:
                template_dic[key]=[]
        if key and p.text.find("Тропар")!=-1:
            template_dic[key].append(p)
        if key and p.text.find("Кондак")!=-1:
            template_dic[key].append(p)
    return template_dic


def get_template_files(path):
    files = glob.glob(path)
    #print(files)
    template_dic={}
    for f in files:
        re_result = re.search(r"docx_templates\\hours-template-(\d)",f)
        template_dic[int(re_result.group(1))] = f
    return template_dic

def get_hours_matrix():
    hours_matrix={}
    for d in range(1,calendar.monthrange(year_no, month_no)[1]+1):
        #print(d)
        if ordo_matrix[d][1] =="":
            if len(templates_menaion[d])<=2:
                hours_matrix[d] = ["","","",
                                "",templates_menaion[d][0],templates_menaion[d][1],
                                "","","",
                                "",templates_menaion[d][0],templates_menaion[d][1]]
            else:
                #print(d)
                hours_matrix[d] = ["","","",
                                "",templates_menaion[d][0],templates_menaion[d][1],
                                "","","",
                                "",templates_menaion[d][2],templates_menaion[d][3]]
        elif ordo_matrix[d][1] =="n":
            pass
        elif ordo_matrix[d][1] =="y":
            hours_matrix[d]=[]
            i=0
            for i in range(12):
                o = ordo_matrix[d][2:][i]
                if o=="":
                    hours_matrix[d].append("")
                elif o =='resurrection' and (i+1)%3!=0:
                    hours_matrix[d].append(templates_resurrection[paschalia.get_echos(datetime(year_no,month_no,d),paschalia_dates)][0])
                elif o =='resurrection' and (i+1)%3==0:
                    hours_matrix[d].append(templates_resurrection[paschalia.get_echos(datetime(year_no,month_no,d),paschalia_dates)][1])
                elif o =='triodion' and (i+1)%3!=0:
                    hours_matrix[d].append(troparia_lent_triodion[d][0])
                elif o =='triodion' and (i+1)%3==0:
                    hours_matrix[d].append(troparia_lent_triodion[d][1])
                elif o =='feast' and (i+1)%3!=0:
                    hours_matrix[d].append(templates_feast[d][0])
                elif o =='feast' and (i+1)%3==0:
                    hours_matrix[d].append(templates_feast[d][1])
            
                elif o =='saint' and (i+1)%3!=0:
                    hours_matrix[d].append(templates_menaion[d][0])
                elif o =='saint' and (i+1)%3==0:
                    hours_matrix[d].append(templates_menaion[d][1])
        '''    
        if month_no == 1 and d==1:
            hours_matrix[d] =   [templates_menaion[d][0],templates_menaion[d][1],templates_menaion[d][2],
                                templates_menaion[d][0],templates_menaion[d][1],templates_menaion[d][3],
                                templates_menaion[d][0],templates_menaion[d][1],templates_menaion[d][2],
                                templates_menaion[d][0],templates_menaion[d][1],templates_menaion[d][3]]
        '''
    #print("paragraphs gathered")          
    '''
    for p in hours_matrix[2]:
        if type(p)==str:
            print("")
        else:
            print(p.text[:40])
    '''
    return hours_matrix


def format_line(p, handle=''):
    #handle = "bir"
    if 'b' in handle:
        p.runs[0].font.bold = True
    if 'i' in handle:
        p.runs[0].font.italic = True
    if 'r' in handle:
        p.runs[0].font.color.rgb = RGBColor(0xff, 0x44, 0x00)
    p.runs[0].font.name='Times New Roman'
    p.runs[0].font.size=152400

# робимо всі "Священик" червоними та italic
BLACK='b'
RED='r'   
def add_text(p,text, color=BLACK):
    r = p.add_run(text)
    r.font.name='Times New Roman'
    r.font.size=152400
    if color == RED:
        r.font.color.rgb = RGB_RED
        r.italic = True


def get_dismissal_matrix(dismissal_csv_filename,cur_month):
    matrix={}
    with open(dismissal_csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:
            if row[0]==month_dic_reversed[cur_month]:                matrix[int(row[1].split('.')[0])]=row[2:]
    return matrix

def insert_dismissal(path,date):
    doc = docx.Document(path)
    shoutout_found = None
    for p in doc.paragraphs:
        re_result=re.search(r"\. Благослови\.",p.text)
        if re_result:
            shoutout_found = True
            #print(date.day, "found!")
            continue

        if shoutout_found:
            #print(date.day, "inserting")
            p_new=p.insert_paragraph_before(dismissal_matrix[date.day][9])
            p_new.paragraph_format.space_after = Pt(6)
            #print(date.day,p_new.text)
            format_line(p_new, '')
            delete_paragraph(p)
            shoutout_found = False

    for p in doc.paragraphs:
        if re.search('Священ{1,2}ик: Христос',p.text):
            re_result=re.search('^(Священ{1,2}ик:)( .+?)(якого є храм)(.*?)$',p.text)
            p_bak=p.text
            p.clear()

            try:
                if re_result:
                    add_text(p,re_result.group(1),color=RED)
                    add_text(p,re_result.group(2))
                    add_text(p,re_result.group(3),color=RED)
                    add_text(p,re_result.group(4))
                else:
                    re_result=re.search(f'^(Священ{1,2}ик:)(.*?)$',p_bak)
                    add_text(p,re_result.group(1),color=RED)
                    add_text(p,re_result.group(2))
            except:
                print(p_bak)
                raise
    doc.save(path)



if __name__== "__main__":
    ordo_matrix = get_matrix(f"Часи_{mode_dic_reversed[mode]}.csv")
    lent_triodion_templates = glob.glob('lent-triodion/*/*.docx')
    pentecostarion_templates = glob.glob('pentecost/*/*.docx')
    pascha_templates = glob.glob('pascha/*/*.docx')
    templates_resurrection = get_resurrection_template_texts('воскресні.docx')
    #templates_menaion = get_menaion_troparia_texts(glob.glob(f'Тропарі - Мінея\\{month_w_offset[month_no-1]:02}-{month_dic_reversed[month_no].upper()}.docx')[0])
    templates_menaion = get_menaion_template_texts(glob.glob(f'Тропарі - Мінея\\{month_w_offset[month_no-1]:02}-{month_dic_reversed[month_no].upper()}.docx')[0])
    templates_feast = get_feast_template_texts(f'свято-{month_no:02}-{mode_dic_reversed[mode]}.docx')
    troparia_lent_triodion = get_feast_template_texts(f'піст-тріодь-{month_no:02}-{mode_dic_reversed[mode]}.docx')
    template_file_list = get_template_files('docx_templates/hours-template-*.docx')

    dismissal_matrix = get_dismissal_matrix(f'Відпусти{mode_suffix}.csv',month_no)

    paschalia_dates = paschalia.get_prev_next_pascha(datetime(year_no, month_no,1), mode)
    hours_matrix = get_hours_matrix()

    if not os.path.exists('drafts'):
        os.makedirs('drafts')
    folder_name=f'drafts\\{year_no}-{month_no:02}-{mode}'
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)

        
    #print(year_no, month_no)
    for d in range(1,calendar.monthrange(year_no, month_no)[1]+1):
        #print(d,datetime(year_no, month_no, d).weekday()+1)
        day_details = paschalia.get_day_details(datetime(year_no,month_no,d),paschalia_dates)
        expected_triodion_template_path=f"lent-triodion\\тиждень-{day_details[1]}\\{day_details[1]}-{day_details[3]}.docx"
        expected_pascha_template_path=f"pascha\\тиждень-{day_details[1]}\\{day_details[1]}-{day_details[3]}.docx"
        expected_pentecostarion_template_path=f"pentecost\\тиждень-{day_details[1]}\\{day_details[1]}-{day_details[3]}.docx"

        dest_filename=f'{folder_name}\\{d:02}.{month_no:02}.docx'
        template=None


        #AI: add cheese-fare week logic here

        if day_details[0]=='lent' and expected_triodion_template_path in lent_triodion_templates and (day_details[3] in (1,2,3,4,5,6) or (day_details[1] in (6,7) and day_details[3]==7)):
            template = expected_triodion_template_path
            #shutil.copy2(expected_triodion_template_path,dest_filename)
        elif day_details[0]=='pascha' and expected_pascha_template_path in pascha_templates:
            template = expected_pascha_template_path
            #shutil.copy2(expected_pentecostarion_template_path,dest_filename)
        elif day_details[0]=='pentecost' and expected_pentecostarion_template_path in pentecostarion_templates:
            template = expected_pentecostarion_template_path
            #shutil.copy2(expected_pentecostarion_template_path,dest_filename)

        elif ordo_matrix[d][1]=='y' or datetime(year_no, month_no, d).weekday()+1==7:
            template = template_file_list[7]
            #shutil.copy2(template_file_list[7],dest_filename)
        else:
            template = template_file_list[datetime(year_no, month_no, d).weekday()+1]
            #shutil.copy2(template_file_list[datetime(year_no, month_no, d).weekday()+1],dest_filename)
        #print(f"{d:02}", day_details, expected_pascha_template_path if day_details=='pascha' else expected_pentecostarion_template_path)
        #print("template: ",template)
        shutil.copy2(template,dest_filename)
    print("Stub files created.")





    hours_dic = {1:"ПЕРШИЙ",
                3:"ТРЕТІЙ",
                6:"ШОСТИЙ",
                9:"ДЕВ'ЯТИЙ"}
    hours_dic_reversed = {v:k for k,v in hours_dic.items()}
    hours_dic_reversed_string='('+'|'.join(hours_dic_reversed.keys())+')'

    hours_translate = {1:1,
                    3:2,
                    6:3,
                    9:4}


    for d in range(1,calendar.monthrange(year_no, month_no)[1]+1):
        #print(d)
        dest_filename=f'{folder_name}\\{d:02}.{month_no:02}.docx'
        doc = docx.Document(dest_filename)
        hour=None
        found_slava=False
        for p in doc.paragraphs:
            
            re_result = re.search(f"^ЧАС {hours_dic_reversed_string}",p.text)
            if re_result:
                hour = int(hours_dic_reversed[re_result.group(1)])
                found_slava = False
            #if d == 22:
                #print(hour)
                #print(d, "found hour", hour)
            re_result = re.search(r"^Слава Отцю, і Сину, і Святому Духові\.$",p.text)
            if re_result and hour and ordo_matrix[d][1]!='n':
                found_slava=True
                #print(f"found СН: {d} {hour} {ordo_matrix[d][1]} {3*hours_translate[hour]-3}")
                value=hours_matrix[d][3*hours_translate[hour]-3]
                if not type(value)==str:
                    #print('inserting slava')
                    copy_paragraph_before(p,value)

            re_result = re.search(r"^Тропар(\s)?",p.text)
            if re_result and hour and ordo_matrix[d][1]!='n':
                #print("found тропар")
                value=hours_matrix[d][3*hours_translate[hour]-2]
                if not found_slava:
                    #print(f'deleting {p.text[:24]}')
                    delete_paragraph(p)
                    continue
                if value and found_slava: #and value.text !=p.text
                    #print(f'inserting troparion {value.text[:24]} before {p.text[:24]}')
                    copy_paragraph_before(p,value)
                    delete_paragraph(p)
                elif value:
                    print("warning troparion",d,hour,value.text[:24])
                

            re_result = re.search(r"^Кондак(\s)?",p.text)
            if re_result and hour and ordo_matrix[d][1]!='n':
                #print("found кондак")
                value=hours_matrix[d][3*hours_translate[hour]-1]
                if value:
                    #print('inserting kondakion')
                    copy_paragraph_before(p,value)
                    delete_paragraph(p)
                elif hours_matrix[d][1]=='y':
                    print("warning kondakion",d,hour)



        doc.save(dest_filename)
        #print(f"Inserting dismissal for file: {dest_filename}")
        insert_dismissal(dest_filename,datetime(year_no,month_no,d))
            
    #for k,v in hours_matrix.items():
    #    print(k,len(v))
    print("Done!")
