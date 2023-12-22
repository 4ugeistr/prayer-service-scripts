import docx, re, csv, calendar, time
import paschalia
from docx.shared import RGBColor
from datetime import datetime
start_time = datetime.now()

COLOR_RED = RGBColor(0xff, 0x00, 0x00)
COLOR_CINNEBAR = RGBColor(0xff, 0x44, 0x00)

docx_filename = "2024NJUL.docx"

year_no=2024

month_list = {'Січень':1,
              'Лютий':2,
              'Березень':3,
              'Квітень':4,
              'Травень':5,
              'Червень':6,
              'Липень':7,
              'Серпень':8,
              'Вересень':9,
              'Жовтень':10,
              'Листопад':11,
              'Грудень':12}
'''
month_list = dict(month_list_lower)
for k,v in month_list_lower.items():
    month_list[k.upper()]=v
'''

reading_indicator_string = r'^(Єв\. –|'\
                                r'Ап\. –|'\
                                r'Єв\. -|'\
                                r'Ап\. -|'\
                                'Читання на Шостому часі:|'\
                                'Час Шостий:|'\
                                'Літ.:|'\
                                'Вечірня з Літургією св. Василія Великого.|'\
                                'На Літургії св. Василія Великого з вечірнею.|'\
                                'Літургія Передосвячених Дарів.|'\
                                'Літургія св. Йоана Золотоустого.|'\
                                'Літургія св. Василія Великого.|'\
                                'На вечірні:|'\
                                'Вечірня:|'\
                                'Утр:|'\
                                'Утр.:|'\
                                'На освячення води|'\
                                'На вмиванні:|'\
                                'По вмиванні:|'\
                                'Ряд.:|'\
                                'Ап.:|'\
                                'Апп.:|'\
                                'Рівноап.:|'\
                                'Предтечі:|'\
                                'Новоліттю:|'\
                                'Св.:|'\
                                'Свв.:|'\
                                'Прп.:|'\
                                'Ісп.:|'\
                                'Мч.:|'\
                                'Мчч.:|'\
                                'Мчц.:|'\
                                'Свщнмч:|'\
                                'Свщмч:|'\
                                'Свщмч.:|'\
                                'Свщнмчч.:|'\
                                'Влкмч.:|'\
                                r'На \d-му часі|'\
                                'Отців:|'\
                                'Отцям:|'\
                                'Богородиці:|'\
                                'Собору:|'\
                                'Оновлення:|'\
                                'Анни:|'\
                                'Св. Миколая:|'\
                                'Свята:|'\
                                'Суботи:|'\
                                'Всім святим:|'\
                                'За упокій:'\
                                ')'


month_list_string='('+'|'.join([x.lower() for x in month_list.keys()])+')'

day_dic = {"Понеділок":1,
            "Вівторок":2,
            "Середа":3,
            "Четвер":4,
            "П'ятниця":5,
            "Субота":6,
            "Неділя":7}
day_dic_reversed = {v:k for k,v in day_dic.items()}
#day_dic_string='('+'|'.join(day_dic.keys())+')'
day_list_string="(Понеділок|Вівторок|Середа|Четвер|П’ятниця|П'ятниця|Субота|Неділя)"
doc = docx.Document(docx_filename)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def get_matrix_full(csv_filename):
    matrix=[]
    with open(csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:
            matrix.append(row)
    return matrix

#matrix[0]["days"][0]["header"]
def get_day_dic(month,day):
    for m in matrix:
        if month == m["month"]:
            for d in m["days"]:
                if day == d["day"]:
                    return d
def format_line(p, handle=''):
    #handle = "bir"
    if 'b' in handle:
        p.runs[0].font.bold = True
    if 'i' in handle:
        p.runs[0].font.italic = True
    if 'r' in handle:
        p.runs[0].font.color.rgb = COLOR_RED
    #p.runs[0].font.name='Times New Roman'
    p.runs[0].font.name='Book Antiqua'
    #p.runs[0].font.size=152400
    p.runs[0].font.size=177800

def format_run(r, handle=''):
    #handle = "bir"
    if 'b' in handle:
        r.font.bold = True
    if 'i' in handle:
        r.font.italic = True
    if 'r' in handle:
        r.font.color.rgb = COLOR_RED
    #p.runs[0].font.name='Times New Roman'
    r.font.name='Book Antiqua'
    #p.runs[0].font.size=152400
    r.font.size=177800
    
def get_saints(month,day):
    saint_string=""
    for row in saint_matrix[1:]:
        if int(row[0])==month and int(row[1])==day:
            if saint_string:
                saint_string+='\n'+row[8]
            else:
                saint_string=row[8]
    return saint_string


ending_fem_dic = {'1':"-ша",
                  '2':"-га",
                  '3':"-тя",
                  '4':"-та",
                  '5':"-та",
                  '6':"-та",
                  '7':"-ма",
                  '8':"-ма",
                  '9':"-та",
                  '0':"-та",
                  '40':"-ва"}

ending_masc_dic = {'1':"-ий",
                  '2':"-ий",
                  '3':"-ій",
                  '4':"-ий",
                  '5':"-ий",
                  '6':"-ий",
                  '7':"-ий",
                  '8':"-ий",
                  '9':"-ий",
                  '0':"-ий"}
    
special_sat_sun_dates=[
    {"date":datetime(year_no,1,6),
     "holiday":"Богоявленні",
     "holiday_locative":"Богоявленні",
     "holiday_instrumental":"Богоявленням"},
    {"date":datetime(year_no,9,14),
     "holiday":"Воздвиження",
     "holiday_locative":"Воздвиженні",
     "holiday_instrumental":"Воздвиженням"},
    {"date":datetime(year_no,12,25),
     "holiday":"Різдво",
     "holiday_locative":"Різдві",
     "holiday_instrumental":"Різдвом"}]

special_named_dates=[
    {"date":datetime(year_no,1,5),
     "day":"Навечір'я Богоявлення."},
    {"date":datetime(year_no,12,24),
     "day":"Навечір'я Різдва"}
    ,]

def insert_special_header(p, date):
    #print)_
    day_name = day_dic_reversed[date.weekday()+1]
    #week_no=paschalia.get_week(date,"","")[:2]
    #ending_fem = ending_fem_dic[str(week_no)[-1]]
    #ending_masc = ending_masc_dic[str(week_no)[-1]]
    
    
    '''
    for sd in special_named_dates:
        if date==sd["date"]:
            #print(date,sd["date"])
            p_new = p.insert_paragraph_before(sd["day"])
            format_line(p_new, '')
    '''        
    #day_label=paschalia.get_day_label(date)
        
    p_new = None
    if date.weekday()+1 in [1,7]:
        p_new = p.insert_paragraph_before(paschalia.get_day_label(date))
        if date.weekday()+1 ==7:
            format_line(p_new, 'r')
        else:
            format_line(p_new, '')


    for sd in special_sat_sun_dates:
        diff = (sd["date"]-date).days
        if abs(diff)<=7 and date.weekday()+1 in [6,7]:
            #p_new=None
            if diff>0:
                if not p_new:
                    p_new = p.insert_paragraph_before(f"{day_name} перед {sd['holiday_instrumental']}.")
                    format_line(p_new, '')
                else:
                    p_new.text = p_new.text[:-1] + f", перед {sd['holiday_instrumental']}."
            elif diff<0:
                if not p_new:
                    p_new = p.insert_paragraph_before(f"{day_name} по {sd['holiday_locative']}.")
                    format_line(p_new, '')
                else:
                    p_new.text = p_new.text[:-1] + f", по {sd['holiday_locative']}."
            if date.weekday()+1 ==7:
                format_line(p_new, 'r')

    if date.weekday()+1 in [7]:
        if p_new:
            r_new = p_new.add_run(f' Гл. {paschalia.get_echos(date)}. Єв. {paschalia.get_resurrection_gospel(date)}')
            format_run(r_new, 'ri')
        else:
            print('WTF')
           
    '''
    if date.weekday()+1 in [6,7]:
        p_new = p.insert_paragraph_before(f"{day_name} {week_no}{ending_fem} по Зісланні Святого Духа.")
        format_line(p_new, '')

    if date.weekday()+1 in [1]:
        p_new = p.insert_paragraph_before(f"Тиждень {week_no}{ending_masc} по Зісланні Святого Духа.")
        format_line(p_new, '')
    '''



saint_matrix = get_matrix_full("Місяцеслов-БД.csv")


matrix = []

month_no = None
day_no = None
header_found = False
header = None
reading_found = False
reading = None
header_inserted = None
day = None
i=0
for p in doc.paragraphs:
#for p in doc.paragraphs[:60]:
    #print(i,p.text[:40])
    i+=1
    re_result = re.search(month_list_string,p.text.lower())
    if re_result:
        month_no = month_list[re_result.group(1).capitalize()]

        #якщо перше квітня - закінчуємо процедуру.
        if month_no==4:
            break

        matrix.append({"month":month_no,"days":[]})
        month = matrix[-1]["days"]
         


    #знаходимо день
    re_result=re.search('^'+day_list_string+'$',p.text)
    if re_result:
        reading_found = False
        reading_first = False
        week_day = re_result.group(1)
        delete_paragraph(p)

        continue
    
        
    #знаходимо початок блоку дня. ігноруємо день тижня
    re_result = re.search(r"^(\d{1,2}) (.*)$",p.text)
    if re_result:
        month.append({"day":int(re_result.group(1)),"header":re_result.group(2)})
        day_no = int(re_result.group(1))
        header = month[-1]
        header_found = True
        reading_found = False
        header_inserted = None

        p_new = p.insert_paragraph_before(re_result.group(1)+" "+week_day)
        p_new.style="Heading 2"
        delete_paragraph(p)
        
        continue

    
    #знаходимо початок блоку читань
    re_result=re.search(reading_indicator_string,p.text)
    if re_result:
        if "reading" in header:
            header["reading"]+='\n'+p.text
        else:
            header["reading"]=p.text
        header_found = False
        reading_found = True
        
    if reading_found and not header_inserted:
        header_inserted=True
        #print("!inserting headers!")
        #header["reading"]+='\n'+p.text
        insert_special_header(p,datetime(year_no, month_no,day_no))

        lst = filter(lambda l:int(l[0])==month_no and int(l[1])==day_no,saint_matrix[1:])
        for l in lst:
            p_new=p.insert_paragraph_before(l[8]+'.')
            format_line(p_new, ''.join(l[2:5]))
        '''
        saints_string=get_saints(month_no, day_no)
        saints_list = saints_string.split('\n')
        for item in saints_list:
            p.insert_paragraph_before(item)
            format_line(p_new, ''.join(l[2:5]))
            #print("!inserted",item)
        '''

    
    if header_found and p.text:
        header["header"] +='\n'+p.text
        delete_paragraph(p)
        continue
    if reading_found and p.text:
        header["header"] +='\n'+p.text
        continue

        
    '''    
    elif reading_first and p.text:
        print("!inserting headers!")
        header["reading"]+='\n'+p.text
        saints_string=get_saints(month_no, day_no)
        saints_list = saints_string.split('\n')
        for item in saints_list:
            p.insert_paragraph_before(item)
            print("!inserted",item)
'''


    
doc.save("2024NJUL_updated.docx")

csv_data=[]
year_no=2024

#for month_no in range(1,13):
for month_no in range(1,4):
    for d in range(1,calendar.monthrange(year_no, month_no)[1]+1):
        csv_data.append([month_no,d])
        #print(month_no,d)
        csv_data[-1].append(get_day_dic(month_no,d)["header"].replace('\n',' '))
        csv_data[-1].append(get_saints(month_no,d).replace('\n',' '))
        
csvfile = open("test.csv",'w',newline='\n',encoding='utf8')
spamwriter=csv.writer(csvfile,delimiter='|',quotechar="\"", quoting=csv.QUOTE_MINIMAL)
spamwriter.writerows(csv_data)
csvfile.close()        

end_time = datetime.now()
elapsed_time = (end_time - start_time).total_seconds()
print(f"Elapsed time : {elapsed_time}") 
