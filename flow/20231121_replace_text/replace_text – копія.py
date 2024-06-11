import docx, re, easygui, glob

#doc_input = docx.Document("input.docx") 
doc_input = easygui.fileopenbox(
    title="Select a .docx file",
    filetypes=["*.docx"],
    default="*.docx"
)

path = easygui.diropenbox()
files = glob.glob(f'{path}\\*.docx')

def get_texts(path):
    doc = docx.Document(path) 
    lst = [[],[]]
    search = False
    replace = False
    i=0
    for p in doc.paragraphs:
        if re.search("Search",p.text):
            #if not lst or not "id" in lst[-1]:
            #    lst.append({"id":i,"search":[],"replace":[]})
            search = True
            replace = False
            i+=1
            continue
        if re.search("Replace",p.text):
            search = False
            replace = True
            continue
        if search:
            lst[0].append(p)
        if replace:
            lst[1].append(p)
    return lst

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def copy_run(target_paragraph,run):
    new_run = target_paragraph.add_run(run.text)
    #KeyError: "no style with name 'Default Paragraph Font'"
    #new_run.style = run.style.name
    new_run.bold = run.bold
    new_run.italic = run.italic
    new_run.underline = run.underline
    #new_run.font.size = run.font.size
    #new_run.font.name = run.font.name
    new_run.font.name='Times New Roman'
    new_run.font.size=152400
    new_run.font.color.rgb = run.font.color.rgb
    new_run.font.highlight_color = run.font.highlight_color

def copy_paragraph_before(paragraph_to_insert_before,source_paragraph):
    target_paragraph = paragraph_to_insert_before.insert_paragraph_before()
    try:
        target_paragraph.style = source_paragraph.style
        #target_paragraph.style.font = source_paragraph.style.font.name
    except KeyError:
        print(f"Warning. Text {source_paragraph.text[:20]} has style{source_paragraph.style}")
    for run in source_paragraph.runs:
        copy_run(target_paragraph,run)
    return target_paragraph

def copy_paragraph_list_before(p_to_insert_before,paragraph_list):
    #print("qty of p to insert:",len(paragraph_list))
    for p in paragraph_list:
        copy_paragraph_before(p_to_insert_before,p)



docs = ["12-Літургія-ГР.docx","12-Літургія-НЮ.docx","Літургія - змінні частини - шаблони.docx"]


def copy_run(target_paragraph,run):
    new_run = target_paragraph.add_run(run.text)
    '''
    try:
        new_run.style = run.style.name
    except KeyError:
        print(f"Warning: can't find style {run.style.name}")
        new_run.style = 'normal'
    '''
    new_run.bold = run.bold
    new_run.italic = run.italic
    new_run.underline = run.underline
    #new_run.font.size = run.font.size
    #new_run.font.name = run.font.name
    new_run.font.name='Times New Roman'
    new_run.font.size=152400
    new_run.font.color.rgb = run.font.color.rgb
    new_run.font.highlight_color = run.font.highlight_color

def copy_paragraph_before(paragraph_to_insert_before,source_paragraph):
    target_paragraph = paragraph_to_insert_before.insert_paragraph_before()
    try:
        target_paragraph.style = source_paragraph.style
        #target_paragraph.style.font = source_paragraph.style.font.name
    except KeyError:
        print(f"Warning. Text {source_paragraph.text[:20]} has style{source_paragraph.style}")
    for run in source_paragraph.runs:
        copy_run(target_paragraph,run)
    return target_paragraph

def copy_paragraph_list_before(p_to_insert_before,paragraph_list):
    #print("qty of p to insert:",len(paragraph_list))
    for p in paragraph_list:
        copy_paragraph_before(p_to_insert_before,p)
        

text_matrix = get_texts(doc_input)

for docx_filename_full in files:
    filename = docx_filename_full.split('\\')[-1]
    print(f"Filename: {filename}")
    doc = docx.Document(docx_filename_full)

    found_start = False
    found_end = False
    i=0
    for p in doc.paragraphs:
        if i<len(text_matrix[0]) and p.text == text_matrix[0][i].text:
            if not found_start and i==0:
                found_start = True
            if found_start:
                delete_paragraph(p)
            i+=1
        elif found_start:
            found_start = False
            found_end = True

        if found_end:
            copy_paragraph_list_before(p,text_matrix[1])
            found_end = False
            i=0
            print(f'{filename}: inserted {text_matrix[1][0].text[:20]}')
    doc.save(docx_filename_full)

print("Done.")

