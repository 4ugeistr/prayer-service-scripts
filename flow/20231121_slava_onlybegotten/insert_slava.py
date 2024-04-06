import docx, re, easygui, glob

doc_input = docx.Document("input.docx") 
doc_input = easygui.fileopenbox(
    title="Select a .docx file",
    filetypes=["*.docx"],
    default="*.docx"
)

path = easygui.diropenbox()
files = glob.glob(f'{path}\\*.docx')

def get_texts(doc):
    lst = [[],[]]
    search = False
    replace = False
    i=0
    for p in doc.paragraphs:
        if re.search("Search",p.text):
            if not lst or not "id" in lst[-1]:
                lst.append({"id":i,"search":[],"replace":[]})
            search = True
            replace = False
            i+=1
            continue
        if re.search("Replace",p.text):
            search = False
            replace = True
            continue
        if search:
            lst[0].append(p)
        if replace:
            lst[1].append(p)
    return lst






docs = ["12-Літургія-ГР.docx","12-Літургія-НЮ.docx","Літургія - змінні частини - шаблони.docx"]


def copy_run(target_paragraph,run):
    new_run = target_paragraph.add_run(run.text)
    new_run.style = run.style.name
    new_run.bold = run.bold
    new_run.italic = run.italic
    new_run.underline = run.underline
    #new_run.font.size = run.font.size
    #new_run.font.name = run.font.name
    new_run.font.name='Times New Roman'
    new_run.font.size=152400
    new_run.font.color.rgb = run.font.color.rgb
    new_run.font.highlight_color = run.font.highlight_color

def copy_paragraph_before(paragraph_to_insert_before,source_paragraph):
    target_paragraph = paragraph_to_insert_before.insert_paragraph_before()
    try:
        target_paragraph.style = source_paragraph.style
        #target_paragraph.style.font = source_paragraph.style.font.name
    except KeyError:
        print(f"Warning. Text {source_paragraph.text[:20]} has style{source_paragraph.style}")
    for run in source_paragraph.runs:
        copy_run(target_paragraph,run)
    return target_paragraph

def copy_paragraph_list_before(p_to_insert_before,paragraph_list):
    #print("qty of p to insert:",len(paragraph_list))
    for p in paragraph_list:
        copy_paragraph_before(p_to_insert_before,p)
        

text_matrix = get_texts(doc_input)

for docx_filename in files:
    doc = docx.Document(docx_filename)

    texts = None
    for p in doc.paragraphs:
        if p.text == text_matrix[0][0].text:
            print(f"Знайдено: {p.text[:20]}")
        '''
        if re.search(r"<antifon3>",p.text):
            print(docx_filename, "FOUND")
            copy_paragraph_list_before(p,doc_input.paragraphs)
        '''
    doc.save(docx_filename)



