import docx


day_short_dic={"ПН":1,
        "ВТ":2,
        "СР":3,
        "ЧТ":4,
        "ПТ":5,
        "СБ":6,
        "НД":7}
day_short_dic_reversed = {v:k for k,v in day_short_dic.items()}
day_short_dic_string='('+'|'.join(day_short_dic.keys())+')'

everyday_part_list=[




]
sunday_part_list=['вечірня',
                'господи_взиваю',
                '4_перші_стихи',
                'світе_тихий_седмиця',
                'сподоби_господи',
                'нині_відпускаєш',
                'середній_відпуст',
                'утреня',
                'тропарі_седмиця',
                'мала_єктенія_перед_другим_сідальним',
                'псалом_50',
                'молитва_9',
                'канон',
                'достойно',
                'мала_єктенія_після_канона',
                'хвалитні',
                'решта_стихів_хвалитних',
                'прохальна_єктенія_утрені',
                'після_стиховні',
                'потрійна_єктенія_утрені',
                'середній_відпуст']

def get_misc_variable_parts(path):
    doc = docx.Document(path)
    matrix = {}
    return matrix


def get_misc_variable_parts(path):
    doc = docx.Document(path)
    matrix = []
    key=None
    for p in doc.paragraphs:
        if '#' in p.text:
            if p.text[1:]=='кінець':
                break
            
            key = p.text[1:]
            matrix.append({'key':key, 'text':[] })
            continue
        
        else:
            matrix[-1]['text'].append(p)
    return matrix


def build_template(path,day,echos):
    doc = docx.Document()
    partlist = sunday_part_list if day==7 else everyday_part_list
    for item in partlist:
        

if __name__ == "__main__":
    build_template('test.docx', 1,1)