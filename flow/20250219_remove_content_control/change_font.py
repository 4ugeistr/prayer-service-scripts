import os
import easygui
from pathlib import Path
from docx import Document  # Ensure python-docx is installed
from docx.shared import Pt


def select_directory():
    return easygui.diropenbox("Select a directory")


def set_font_style(docx_path):
    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    doc.save(docx_path)


def has_content_controls(docx_path):
    doc = Document(docx_path)
    for element in doc.element.body.iter():
        if element.tag == qn('w:sdt'):
            return True
    return False


def process_files(directory):
    for root, _, files in os.walk(directory):
        for file in files:
            if file.endswith(".docx"):
                docx_path = Path(root) / file
                print("Processing:", docx_path)     
                # Set font style in the new docx file
                set_font_style(docx_path)
                


def main():
    directory = select_directory()
    if directory:
        process_files(directory)
        print("Conversion completed!")
    else:
        print("No directory selected.")


if __name__ == "__main__":
    main()
