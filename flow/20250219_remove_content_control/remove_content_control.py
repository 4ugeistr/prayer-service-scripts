import os
import easygui
from pathlib import Path
from win32com import client  # Ensure pywin32 is installed
from docx import Document  # Ensure python-docx is installed
from docx.oxml.ns import qn
from docx.shared import Pt


def select_directory():
    return easygui.diropenbox("Select a directory")


def convert_docx_to_doc(docx_path, doc_path):
    word = client.Dispatch("Word.Application")
    doc = word.Documents.Open(str(docx_path))
    doc.SaveAs(str(doc_path), FileFormat=0)  # 0 is for .doc format
    doc.Close()
    word.Quit()


def convert_doc_to_docx(doc_path, docx_path):
    word = client.Dispatch("Word.Application")
    doc = word.Documents.Open(str(doc_path))
    doc.SaveAs(str(docx_path), FileFormat=12)  # 12 is for .docx format
    doc.Close()
    word.Quit()


def set_font_style(docx_path):
    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    doc.save(docx_path)


def has_content_controls(docx_path):
    doc = Document(docx_path)
    for element in doc.element.body.iter():
        if element.tag == qn('w:sdt'):
            return True
    return False


def process_files(directory):
    for root, _, files in os.walk(directory):
        for file in files:
            if file.endswith(".docx"):
                
                print("Checking:", file)
                docx_path = Path(root) / file
                if not has_content_controls(docx_path):
                    print("Skipping (no content controls):", file)
                    continue
                print("Processing:", file)
                doc_path = docx_path.with_suffix(".doc")
                
                # Convert docx to doc
                convert_docx_to_doc(docx_path, doc_path)
                
                # Convert doc back to docx
                new_docx_path = doc_path.with_suffix(".docx")
                convert_doc_to_docx(doc_path, new_docx_path)
                
                # Set font style in the new docx file
                set_font_style(new_docx_path)
                
                # Delete the .doc file
                os.remove(doc_path)


def main():
    directory = select_directory()
    if directory:
        process_files(directory)
        print("Conversion completed!")
    else:
        print("No directory selected.")


if __name__ == "__main__":
    main()
