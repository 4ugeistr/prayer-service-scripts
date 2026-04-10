import os,re

import ps_docx_utils as pdu
import easygui as eg

from datetime import datetime

import docx
from docx.enum.text import WD_COLOR_INDEX

from sys import exit

day_name_sequence = ["НД","ПН","ВВ","СР","ЧТ","ПТ","СБ",]

day_short_dic={"ПН":1,
        "ВВ":2,
        "СР":3,
        "ЧТ":4,
        "ПТ":5,
        "СБ":6,
        "НД":7}
day_short_dic_reversed = {v:k for k,v in day_short_dic.items()}
day_short_dic_string='('+'|'.join(day_short_dic.keys())+')'

day_dic = {"Понеділок":1,
            "Вівторок":2,
            "Середа":3,
            "Четвер":4,
            "П'ятниця":5,
            "Субота":6,
            "Неділя":7}
day_dic_reversed = {v:k for k,v in day_dic.items()}
day_dic["П’ятниця"]=5
day_dic_string='('+'|'.join(day_dic.keys())+')'



def find_docx_files(start_dir):

    docx_files = []
    for root, dirs, files in os.walk(start_dir):
        for file in files:
            if file.endswith('.docx') and not file.startswith('~'):
                full_path = os.path.join(root, file)
                docx_files.append(full_path)
    return docx_files

def get_octoechos_file(docx_files, echos,weekday_abr):
    for file_path in docx_files:
        #print(file_path)
        file_name = os.path.basename(file_path)
        #print(file_name)
        if file_name[0] == str(echos) and file_name[2:4]==weekday_abr:
            return file_path
    else:
        print(f"File not found for {echos}, {weekday_abr}")
        raise Exception
        #exit(-1)
        #return
    #return docx_files[0]

def format_as_code(par):
    par.runs[0].font.highlight_color=WD_COLOR_INDEX.GRAY_50

def build_sources(docx_files):
    for echos in range(1,9):
        print(f"Обробляємо глас {echos}")
        if echos>1:
            continue
        filename = f'Джерело. Глас {echos}.docx'
        new_doc = docx.Document()
        new_doc.add_heading(f'ГЛАС {echos}', level=1)

        for day in day_name_sequence:
            new_doc.add_heading(f'Глас {echos} - {day}', level=1)
            par = new_doc.add_paragraph(f"/октоїх/{echos}/{day}")
            format_as_code(par)
            
            old_template = get_octoechos_file(docx_files,echos,day)
            old_doc = docx.Document(old_template)

            new_doc.add_heading('ВЕЧІРНЯ', level=2)
            
            par = new_doc.add_heading(f"//вечірня/господи_взиваю", level=3)
            format_as_code(par)
            
            #[Вечірня] Стихири на Господи Взиваю
            paragraphs = get_stichera_gv(old_doc)
            for i,p in enumerate(paragraphs):
                par = new_doc.add_paragraph(f"//вечірня/господи_взиваю/стихири/{i+1}")
                format_as_code(par)
                pdu.copy_paragraph(new_doc,p)
            #ps.copy_paragraph_list(new_doc,paragraphs)

            #[Вечірня] Стихири на стиховні
            paragraphs = get_vespers_stichera_stkh(old_doc)
            for i,p in enumerate(paragraphs):
                par = new_doc.add_paragraph(f"//вечірня/стиховня/стихири/{i+1}")
                format_as_code(par)
                pdu.copy_paragraph(new_doc,p)
                
            new_doc.add_heading('УТРЕНЯ', level=2)
            
            #[Утреня] Сідальні після Першого славослов'я
            paragraphs = get_first_sessional_hymns(old_doc)
            for i,p in enumerate(paragraphs):
                par = new_doc.add_paragraph(f"//утреня/сідальні_по_першому_славословї/{i+1}")
                format_as_code(par)
                pdu.copy_paragraph(new_doc,p)

            #[Утреня] Сідальні після Другого славослов'я
            paragraphs = get_second_sessional_hymns(old_doc)
            for i,p in enumerate(paragraphs):
                par = new_doc.add_paragraph(f"//утреня/сідальні_по_другому_славословї/{i+1}")
                format_as_code(par)
                pdu.copy_paragraph(new_doc,p)

            #[Утреня] Іпакой
            paragraphs = get_hypakoe(old_doc)
            if paragraphs:
                #for i, p in enumerate(paragraphs):
                    par = new_doc.add_paragraph(f"//утреня/іпакой")
                    format_as_code(par)
                    pdu.copy_paragraph(new_doc, paragraphs)

            #[Утреня] Степенна пісня
            paragraphs = get_anavathmoi(old_doc)
            if paragraphs:
                par = new_doc.add_paragraph(f"//утреня/степенна")
                format_as_code(par)
                pdu.copy_paragraph_list(new_doc,paragraphs)


            #[Утреня] Прокімен
            paragraphs = get_prokimenon(old_doc)
            if paragraphs:
                par = new_doc.add_paragraph(f"//утреня/прокімен")
                format_as_code(par)
                pdu.copy_paragraph_list(new_doc, paragraphs)

            # [Утреня] Пісня Канону
            if day_short_dic[day] == 7:
                for n in kanon_ode_numbers:
                    paragraphs = get_kanon_ode(old_doc, str(n))
                    if paragraphs:
                        par = new_doc.add_paragraph(f"//утреня/канон/пісня_{n}")
                        format_as_code(par)
                        pdu.copy_paragraph_list(new_doc, paragraphs)
            else:
                for n in kanon_ode_numbers:
                   if n in templates_kanon_dic[echos][day_short_dic[day]][1]:
                       par = new_doc.add_paragraph(f"//утреня/канон/пісня_{n}")
                       format_as_code(par)
                       pdu.copy_paragraph_list(new_doc, templates_kanon_dic[echos][day_short_dic[day]][1][n])

            #[Утреня] Світильні
            paragraphs = get_exapostolaria(old_doc)
            for i,p in enumerate(paragraphs):
                par = new_doc.add_paragraph(f"//утреня/світильні/{i+1}")
                format_as_code(par)
                pdu.copy_paragraph(new_doc,p)

            #[Утреня] Стихири хвалитні
            paragraphs = get_stichera_lauds(old_doc)
            if paragraphs:
                for i,p in enumerate(paragraphs):
                    par = new_doc.add_paragraph(f"//утреня/хвалитні/стихири/{i+1}")
                    format_as_code(par)
                    pdu.copy_paragraph(new_doc,p)
            elif day=="НД":
                print(f"Утреня, стихири на хвалитних не знайдено для {echos}-{day}")
              
   
            #[Утреня] Стихири на стиховні
            paragraphs = get_matins_stichera_stkh(old_doc)
            if paragraphs:
                for i,p in enumerate(paragraphs):
                    par = new_doc.add_paragraph(f"//утреня/стиховня/стихири/{i+1}")
                    format_as_code(par)
                    pdu.copy_paragraph(new_doc,p)
            #else:
            #    print(f"Утреня, стихири на стиховні не знайдено для {echos}-{day}")
            
        new_doc.save(f"Джерело. Глас {echos}.docx")

def process_template(filepath):
    doc = docx.Document(filepath)


def get_stichera_gv(doc):

    found_header = False
    paragraphs_found = []

    for p in doc.paragraphs:

        # Шукаємо фрагмент документу, що нас цікавить. Наприклад - за заголовком
        if p.text == "Стихири":
            found_header = True
            continue

        # Задаємо строки, які ігноруємо
        if found_header and (p.text.startswith("Стих:") or p.text.startswith("Псалом")):
            continue

        # Обробляємо потрібні параграфи
        if found_header and p.text:
            paragraphs_found.append(p)
            continue

        # Виходимо, коли закінчили опрацьовувати потрібний фрагмент
        if found_header and (not p.text or p.text in ("Вхід","Світло тихе")):
            return paragraphs_found

def get_vespers_stichera_stkh(doc):

    found_header = False
    paragraphs_found = []

    for p in doc.paragraphs:

        # Шукаємо початок фрагменту документу, що нас цікавить. Наприклад - за заголовком
        if p.text == "Стихири на стиховні":
            found_header = True
            continue

        # Пропускаємо стихи для Стиховні Октоїха
        # Якщо МІНЕЯ - треба буде дописати
        if found_header and (p.text.startswith("Стих:")):
            continue

        # Обробляємо потрібні параграфи
        if found_header and p.text:
            paragraphs_found.append(p)
            continue

        # Виходимо, коли закінчили опрацьовувати потрібний фрагмент
        if found_header and (not p.text or p.text in ("Пісня Симеона")):
            return paragraphs_found
        
def get_first_sessional_hymns(doc):

    found_header = False
    paragraphs_found = []

    for p in doc.paragraphs:

        # Шукаємо початок фрагменту документу, що нас цікавить. Наприклад - за заголовком
        if p.text == "Перший сідальний":
            found_header = True
            continue

        # Обробляємо потрібні параграфи
        if found_header and p.text:
            paragraphs_found.append(p)
            continue

        # Виходимо, коли закінчили опрацьовувати потрібний фрагмент
        if found_header and (not p.text or p.text in ("Мала єктенія") or p.text.startswith("Відтак")):
            return paragraphs_found

def get_second_sessional_hymns(doc):

    found_header = False
    paragraphs_found = []

    for p in doc.paragraphs:

        # Шукаємо початок фрагменту документу, що нас цікавить. Наприклад - за заголовком
        if p.text == "Другий сідальний":
            found_header = True
            continue

        # Обробляємо потрібні параграфи
        if found_header and p.text:
            paragraphs_found.append(p)
            continue

        # Виходимо, коли закінчили опрацьовувати потрібний фрагмент
        if found_header and (not p.text or p.text in ("Псалом 50")):
            return paragraphs_found

def get_hypakoe(doc):

    found_header = False
    paragraphs_found = []

    for p in doc.paragraphs:

        # Шукаємо початок фрагменту документу, що нас цікавить. Наприклад - за заголовком
        if p.text.startswith("Іпакой"):
            paragraphs_found.append(p)
            found_header = True
            continue

        # Обробляємо потрібні параграфи. Одразу виходимо, бо параграф завжди один
        if found_header and p.text:
            paragraphs_found.append(p)
            return paragraphs_found

def get_anavathmoi(doc):

    found_header = False
    paragraphs_found = []

    for p in doc.paragraphs:

        # Шукаємо початок фрагменту документу, що нас цікавить. Наприклад - за заголовком
        if p.text.startswith("Степенна пісня"):
            found_header = True
            paragraphs_found.append(p)
            continue

        # Обробляємо потрібні параграфи
        if found_header and p.text:
            paragraphs_found.append(p)
            continue

        # Виходимо, коли закінчили опрацьовувати потрібний фрагмент
        if found_header and (not p.text or p.text in ("Євангеліє")):
            return paragraphs_found

def get_prokimenon(doc):

    found_header = service_matins = False
    paragraphs_found = []

    for p in doc.paragraphs:

        if p.text == "УТРЕНЯ":
            service_matins = True
            continue

        # Шукаємо початок фрагменту документу, що нас цікавить. Наприклад - за заголовком
        if service_matins and p.text.startswith("Прокімен"):
            paragraphs_found.append(p)
            found_header = True
            continue

        # Обробляємо потрібні параграфи
        if found_header and p.text:
            paragraphs_found.append(p)
            continue

        # Виходимо, коли закінчили опрацьовувати потрібний фрагмент
        if found_header and (not p.text or p.text in ("Господеві помолімся")):
            return paragraphs_found

kanon_ode_numbers = list(range(1,10))

def get_kanon_ode(doc,no):

    found_header = False
    paragraphs_found = []

    for p in doc.paragraphs:

        # Шукаємо початок фрагменту документу, що нас цікавить. Наприклад - за заголовком
        if p.text == f"Пісня {no}":
            found_header = True
            paragraphs_found.append(p)
            continue

        # Обробляємо потрібні параграфи
        if found_header and p.text:
            paragraphs_found.append(p)
            continue

        # Виходимо, коли закінчили опрацьовувати потрібний фрагмент
        if found_header and (not p.text or p.text.startswith("Пісня") or p.text in ("Мала єктенія")):
            return paragraphs_found


def get_octoechos_kanon_texts(path):
    doc = docx.Document(path)
    matrix = {}
    for i in range(8):
        matrix[i + 1] = {}
    for p in doc.paragraphs:

        re_result = re.search(r"Глас (\d)", p.text)
        if re_result:
            echos = int(re_result.group(1))
            echos_found = True
            song_found = False
            continue

        if (echos_found or song_found) and re.search(f"{day_dic_string}", p.text):
            week_day_no = day_dic[re.search(f"{day_dic_string}", p.text).group(1)]
            matrix[echos][week_day_no] = []
            week_day_no_found = True
            echos_found = song_found = False
            continue

        if (week_day_no_found or song_found) and re.search("Канон", p.text):
            matrix[echos][week_day_no].append({"label": p.text})
            label_found = True
            week_day_no_found = song_found = False
            continue

        re_result = re.search(r"Пісня (\d)", p.text)
        if (label_found or song_found) and re_result:
            song_no = int(re_result.group(1))
            matrix[echos][week_day_no][-1][song_no] = [p]
            label_found = False
            song_found = True
            continue

        if p.text == "":
            continue
        if song_found:
            matrix[echos][week_day_no][-1][song_no].append(p)

    return matrix

def get_exapostolaria(doc):

    found_header = False
    paragraphs_found = []

    for p in doc.paragraphs:

        # Шукаємо початок фрагменту документу, що нас цікавить. Наприклад - за заголовком
        if p.text == "Світильний":
            found_header = True
            continue

        # Обробляємо потрібні параграфи
        if found_header and p.text and not p.text.startswith("ВСТАВИТИ"):
            paragraphs_found.append(p)
            continue

        # Виходимо, коли закінчили опрацьовувати потрібний фрагмент
        if found_header and (not p.text or p.text in ("Хвалитні")):
            return paragraphs_found

def get_stichera_lauds(doc):

    found_header = service_matins = False
    paragraphs_found = []

    for p in doc.paragraphs:

        # Шукаємо початок фрагменту документу, що нас цікавить. Наприклад - за заголовком
        if p.text == "УТРЕНЯ":
            service_matins = True
            continue
        
        # Шукаємо початок фрагменту документу, що нас цікавить. Наприклад - за заголовком
        if service_matins and p.text == "Стихири":
            found_header = True
            continue

        # Задаємо строки, які ігноруємо
        if found_header and (p.text.startswith("Стих:") or p.text.startswith("Псалом")):
            continue

        # Обробляємо потрібні параграфи
        if found_header and p.text:
            paragraphs_found.append(p)
            continue

        # Виходимо, коли закінчили опрацьовувати потрібний фрагмент
        if found_header and (not p.text or p.text in ("Стихира євангельська")):
            return paragraphs_found

def get_matins_stichera_stkh(doc):

    found_header = service_matins = False
    paragraphs_found = []

    for p in doc.paragraphs:

        # Шукаємо початок фрагменту документу, що нас цікавить. Наприклад - за заголовком
        if p.text == "УТРЕНЯ":
            service_matins = True
            continue
        
        # Шукаємо початок фрагменту документу, що нас цікавить. Наприклад - за заголовком
        if service_matins and p.text == "Стихири на стиховні":
            found_header = True
            continue

        # Пропускаємо стихи для Стиховні Октоїха
        # Якщо МІНЕЯ - треба буде дописати
        if found_header and (p.text.startswith("Стих:")):
            continue


        # Обробляємо потрібні параграфи
        if found_header and p.text:
            paragraphs_found.append(p)
            continue

        # Виходимо, коли закінчили опрацьовувати потрібний фрагмент
        if found_header and (not p.text or p.text.startswith("Після стиховні")):
            return paragraphs_found



if __name__ == "__main__":
    start_time = datetime.now()


    target_dir = eg.diropenbox(title="Оберіть початкову директорію")
    docx_files = find_docx_files(target_dir)

    templates_kanon_dic = get_octoechos_kanon_texts('05a_ОКТОЇХ_КАНОНИ.docx')




    build_sources(docx_files)

    print("Done!")

    
    '''
    doc = docx.Document(get_octoechos_file(docx_files,1,"ВВ"))

    paragraphs_found = get_gv_stichera(doc)

    print(f"Found {len(paragraphs_found)} paragraphs with stichera:")
    for i, p in enumerate(paragraphs_found):
        print(f"[{i}]: {p.text}")
    '''
    print(f"Time spent: {(datetime.now() - start_time).total_seconds()}")