import easygui, calendar, os, docx, re, csv
from docx.shared import RGBColor
import paschalia
from datetime import datetime, timedelta
import cal_generate as generate_calendar
import ps_docx_utils as pdu


mode = easygui.choicebox('u - Юліанський, g - Григоріанський', 'Вибір календаря', ['u','g'])
mode_dict={'u':'Юл',
           'g':'Гр'}

#paschalia.paschalia = paschalia.get_prev_next_pascha(datetime(2024,1,1),mode)

#select month
month_no = datetime.now().month+1 if datetime.now().month!=12 else 1
year_no = datetime.now().year if datetime.now().month!=12 else datetime.now().year+1
#month_no=3
#print("WARNING. Month_no OVERRIDE", month_no)

month_dic = {'Січень':1,
              'Лютий':2,
              'Березень':3,
              'Квітень':4,
              'Травень':5,
              'Червень':6,
              'Липень':7,
              'Серпень':8,
              'Вересень':9,
              'Жовтень':10,
              'Листопад':11,
              'Грудень':12}
month_dic_reversed = {v:k for k,v in month_dic.items()}
month_dic_string='('+'|'.join(month_dic.keys())+')'

day_dic = {"Понеділок":1,
            "Вівторок":2,
            "Середа":3,
            "Четвер":4,
            "П'ятниця":5,
            "Субота":6,
            "Неділя":7}
day_dic_reversed = {v:k for k,v in day_dic.items()}
day_dic_string='('+'|'.join(day_dic.keys())+')'

#doc = docx.Document("Літургія - Мінея\\10-Літургія-шаблони.docx") 
'''
#DEPRECATED -  see paschalia
def get_echos(date,mode):
    if mode=='u':
        return (int(date.strftime("%U"))-25) % 8 + 1
    elif mode=='g':
        return (int(date.strftime("%U"))-24) % 8 + 1
'''

'''
#DEPRECATED -  see docx_utils
def copy_paragraph(target_doc,source_paragraph):
    target_paragraph = target_doc.add_paragraph()
    for run in source_paragraph.runs:
        new_run = target_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.font.color.rgb = run.font.color.rgb
        new_run.font.highlight_color = run.font.highlight_color
        if new_run.text.startswith('<'):
            new_run.font.name = 'Consolas'
            new_run.font.highlight_color = 15
        else:
            new_run.font.name = 'Times New Roman'
            
#DEPRECATED -  see docx_utils
def copy_paragraph_list(target_doc,paragraph_list):
    for p in paragraph_list:
        copy_paragraph(target_doc,p)
        
#DEPRECATED -  see docx_utils
def format_line(p, handle=''):
    #handle = "bir"
    if 'b' in handle:
        p.runs[0].font.bold = True
    if 'i' in handle:
        p.runs[0].font.italic = True
    if 'r' in handle:
        p.runs[0].font.color.rgb = RGBColor(0xff, 0x44, 0x00)
    p.runs[0].font.name='Times New Roman'
    p.runs[0].font.size=152400
'''

def get_dismissal_matrix(dismissal_csv_filename,cur_month):
    matrix={}
    with open(dismissal_csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:
            if row[0]==month_dic_reversed[cur_month]:
                matrix[int(row[1].split('.')[0])]=row[2:]
    return matrix

def get_matrix_full(csv_filename):
    matrix=[]
    with open(csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:
            matrix.append(row)
    return matrix

        
template_path = "docx_resources\\Літургія\\Літургія - змінні частини - шаблони.docx"
def get_resurrection_template_texts(path):
    template_dic={}
    doc = docx.Document(path)
    rubric_found = cur_glas = None
    i=0
    for p in doc.paragraphs:
        #print(p.text)
        if not rubric_found and not p.text.startswith("Воскресна служба"):
            pass
        elif p.text.startswith("Воскресна служба"):
            rubric_found=True
        if rubric_found:
            re_result = re.search(r"Неділя – глас (\d)",p.text)
            if re_result:
                cur_glas = int(re_result.group(1))
                template_dic[cur_glas]=[]
                continue

        if cur_glas and p.text.startswith("</vidpust>"):
            template_dic[cur_glas].append(p)
            cur_glas=None

        if cur_glas:
            template_dic[cur_glas].append(p)
    return template_dic
            
def get_everyday_template_texts(path):
    template_dic={}
    doc = docx.Document(path)
    rubric_found = cur_glas = None
    i=0
    for p in doc.paragraphs:
        #print(p.text)
        if not rubric_found and not p.text.startswith("Повсякденна служба"):
            pass
        elif p.text.startswith("Повсякденна служба"):
            rubric_found=True
        if rubric_found:
            re_result = re.search(f"{day_dic_string}",p.text)
            if re_result:
                cur_glas = day_dic[re_result.group(1)]
                template_dic[cur_glas]=[]
                continue

        if cur_glas and p.text.startswith("</vidpust>"):
            template_dic[cur_glas].append(p)
            cur_glas=None

        if cur_glas:
            template_dic[cur_glas].append(p)
    return template_dic     
        
def get_menaion_template_texts():
    template_dic={}
    directory_path ="docx_resources\\Літургія\\Літургія - Мінея"
    all_files = os.listdir(directory_path)
    try:
        docx_file = [file for file in all_files if file.endswith(".docx") and file[:2]==f'{month_no:02}' and len(file) >= 4][0]
    except IndexError:
        print(f"Warning: menaion file not found for month {month_no}")
        return template_dic
    print(docx_file)
    doc = docx.Document("docx_resources\\Літургія\\Літургія - Мінея\\"+docx_file)
    template_found=False
    starting_tag_found=False
    
    for p in doc.paragraphs:
        
        re_result=re.search(f'^{month_dic_string} '+r'(\d+)',p.text)
        if re_result:
            cur_date = int(re_result.group(2))
            template_dic[cur_date]=[]
            template_found = True
        
        re_result=re.search(f'<ustav',p.text)
        if re_result:
            starting_tag_found=True
            
        if template_found and starting_tag_found:
            template_dic[cur_date].append(p)

        re_result=re.search(f'</vidpust',p.text)
        if re_result:
            #template_dic[cur_date]=template_dic[cur_date][1:]
            template_found=False
            cur_date=None
            starting_tag_found=False
            
    
    #print(template_dic)
    return template_dic

def get_triodion_template_texts(path):
    template_dic={}
    doc = docx.Document(path)
    
    template_found=False
    starting_tag_found=False
    triodion_week=triodion_day=None

    for p in doc.paragraphs:
        
        re_result=re.search(r'^Тиждень (\d), День (\d)',p.text)
        if re_result:
            triodion_week = int(re_result.group(1))
            triodion_day = int(re_result.group(2))
            #cur_date = int(re_result.group(2))
            template_dic[(triodion_week,triodion_day)]=[]
            template_found = True
            continue
        
        re_result=re.search(f'<ustav',p.text)
        if re_result:
            starting_tag_found=True
            
        if template_found and starting_tag_found:
            template_dic[(triodion_week,triodion_day)].append(p)

        re_result=re.search(f'</vidpust',p.text)
        if re_result:
            #template_dic[cur_date]=template_dic[cur_date][1:]
            template_found=False
            #cur_date=None
            starting_tag_found=False

    #print(template_dic)
    return template_dic


'''
OBSOLETE
moved to cal_generate.py

'''

def insert_header_liturgy(doc,date):
    day_name = day_dic_reversed[date.weekday()+1]
    #month_name = month_dic_reversed[month_no]
    week_no=paschalia.get_day_details(date,mode)

    #Субота, Неділя, Тиждень etc...
    special_dates=[{"date":datetime(year_no,12,25),
                     "holiday":"Різдво",
                     "holiday_locative":"Різдві",
                     "holiday_instrumental":"Різдвом"},
                    {"date":datetime(year_no+1,1,6),
                     "holiday":"Богоявленні",
                     "holiday_locative":"Богоявленні",
                     "holiday_instrumental":"Богоявленням"}]
    for sd in special_dates:
        diff = (sd["date"]-date).days
        if abs(diff)<=7 and date.weekday()+1 in [6,7]:
            if diff>0:
                p_new = doc.add_paragraph(f"{day_name} перед {sd['holiday_instrumental']}")
            else:
                p_new = doc.add_paragraph(f"{day_name} по {sd['holiday_locative']}")
            pdu.format_line(p_new, '')


    #if date.weekday()+1 in [6,7]:
    if date.weekday()+1:
        #print(date)
        p_new=doc.add_paragraph(reading_matrix[date.day-1][4])
        if date.weekday()+1 == 7:
            p_new.text += f" Глас {paschalia.get_echos(date,mode)}."
        #print(reading_matrix[date.day-1][:5])

    


    '''
    if date.weekday()+1 in [6,7]:
        #p_new = doc.add_paragraph(f"{day_name} {week_no} по П'ятидесятниці.")
        p_new = doc.add_paragraph(paschalia.get_day_label_legacy(date))
        if date.weekday()+1 == 7:
            p_new.text+= f" Гл. "+str(paschalia.get_echos(date,paschalia_dates))+"."
            format_line(p_new, 'r')
        else:
            format_line(p_new, '')
    '''
    '''
    if date.weekday()+1 in [1]:
        #p_new = doc.add_paragraph(f"Тиждень {week_no} по П'ятидесятниці.")
        p_new = doc.add_paragraph(paschalia.get_day_label_legacy(date))
        format_line(p_new, '')
    '''
    '''
    #перелік святих
    lst = filter(lambda l:int(l[0])==date.month and int(l[1])==date.day,saint_matrix[1:])
    for l in lst:
        p_new=doc.add_paragraph(l[9])
        format_line(p_new, ''.join(l[2:5]))
    '''
if mode == 'u':
    mode_suffix='Юл'
elif mode == 'g':
    mode_suffix='Гр'


if __name__ == "__main__":
    #init data
    #paschalia_dates = paschalia.get_prev_next_pascha(datetime(year_no, month_no,1), mode)
    templates_menaion = get_menaion_template_texts()
    templates_lent = get_triodion_template_texts("docx_resources\\Літургія\\Літургія - Тріодь\\Літургія - Піст.docx")
    templates_pascha = get_triodion_template_texts("docx_resources\\Літургія\\Літургія - Тріодь\\Літургія - Пасха.docx")
    templates_pentecost = get_triodion_template_texts("docx_resources\\Літургія\\Літургія - Тріодь\\Літургія - 50-ця.docx")
    
    templates_resurrection = get_resurrection_template_texts(template_path)
    templates_everyday = get_everyday_template_texts(template_path)
    saint_matrix = get_matrix_full("matrices\\Місяцеслов-БД.csv")
    reading_matrix = get_matrix_full(f"matrices\\Читання{mode_suffix}.csv")
    reading_matrix = list(filter(lambda x: (x[0]== month_dic_reversed[month_no]), reading_matrix))  

    #dismissal_matrix = get_dismissal_matrix(f'Читання{mode_dict[mode]}.csv',month_no)


    new_doc = docx.Document()
    '''
    Get all days in the month
    '''
    for d in range(1,calendar.monthrange(year_no, month_no)[1]+1):
    #for d in range (5,6):
        #print(d)
        day_details = paschalia.get_day_details(datetime(year_no,month_no,d),mode)
        print(d, day_details)
        if day_details[0]=='lent':
            if not day_details[3] in (6,7):
                continue
            elif (day_details[1]==7 and day_details[3]==6):
                continue
            else:
                #TODO: insert vv
                pass

        heading_text =' '.join([month_dic_reversed[month_no],str(d)+',',day_dic_reversed[datetime(year_no, month_no, d).weekday()+1]])
        new_doc.add_heading(heading_text, level=2)




        #insert_header_liturgy(new_doc,datetime(year_no,month_no,d))

        pdu.copy_paragraph_list(new_doc,generate_calendar.prepare_header_for_docx(datetime(year_no,month_no,d),mode))

        #print(dismissal_matrix[d][2])
        if day_details[0] == 'lent' and (day_details[1], day_details[3]) in templates_lent.keys():
            pdu.copy_paragraph_list(new_doc, templates_lent[(day_details[1], day_details[3])])
        elif day_details[0]=='pascha' and (day_details[1], day_details[3]) in templates_pascha.keys():
            pdu.copy_paragraph_list(new_doc,templates_pascha[(day_details[1], day_details[3])])
            #print("template: pascha")
        elif day_details[0]=='pentecost' and (day_details[1], day_details[3]) in templates_pentecost.keys():
            pdu.copy_paragraph_list(new_doc,templates_pentecost[(day_details[1], day_details[3])])
            #print("template: pentecost")
        elif datetime(year_no, month_no, d).weekday()+1==7:
            #print(d, "copied sunday ", get_echos(datetime(year_no,month_no,d)))
            pdu.copy_paragraph_list(new_doc,templates_resurrection[paschalia.get_echos(datetime(year_no,month_no,d),mode)])
            #print("template: resurrection")
        elif d in templates_menaion.keys():
            pdu.copy_paragraph_list(new_doc,templates_menaion[d])
            #print("template: menaion")
        else:
            #print(d, datetime(year_no, month_no, d).weekday()+1,templates_everyday[datetime(year_no, month_no, d).weekday()+1][38].text[:20])
            pdu.copy_paragraph_list(new_doc,templates_everyday[datetime(year_no, month_no, d).weekday()+1])
            #print("template: everyday")

    #for p in new_doc.paragraphs:
    #    re_result = re.search('',p.text)


        
    mode_new = 'ГР' if mode=='g' else "НЮ"
    new_doc.save(f'drafts\\liturgy\\{month_no:02}-Літургія-{mode_new}.docx')
