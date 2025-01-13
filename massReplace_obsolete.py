import glob,docx,re
from docx.shared import Pt

searchFor = 'безплотих'
replaceWith = "безплотних"

def simpleReplaceText(filename):
    doc = docx.Document(filename)
    for p in doc.paragraphs:
        for r in p.runs:
            #if 'НД' in filename:
            #    print(r.text)
            if re.search(searchFor,r.text):
                r.text=re.sub(searchFor, replaceWith, r.text)
                r.font.name='Times New Roman'
                r.font.size=152400
                print(f'Updated: {filename}')
                doc.save(filename)

def specialReplaceText(filename):
    doc = docx.Document(filename)
    for p in doc.paragraphs:
        #print(p.space_after)
        if re.search('Ти, що береш гріхи світу,\* прийми нашу молитву\. Ти, що сидиш по правиці Отця,\* помилуй нас\.',p.text):
            p_new=p.insert_paragraph_before('Ти, що береш гріхи світу,* прийми нашу молитву.', style='Normal')
            p_new.paragraph_format.space_after = Pt(6)
            p.text='Ти, що сидиш по правиці Отця,* помилуй нас.'
            p.style='Normal'
            p.runs[0].font.name=p_new.runs[0].font.name='Times New Roman'
            p.runs[0].font.size=p_new.runs[0].font.size=152400
            doc.save(filename)
            print(f'Updated: {filename}')

    
def main():
    for filename in glob.glob('*/*.docx'):
        simpleReplaceText(filename)

main()

