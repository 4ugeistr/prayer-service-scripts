from docx.shared import RGBColor
from docx.shared import Pt
RGB_RED = RGBColor(0xff, 0x44, 0x00)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def delete_run(run):
    p = run._element
    p.getparent().remove(p)
    p._p = p._element = None

def copy_run(target_paragraph,run):
    new_run = target_paragraph.add_run(run.text)
    
    #!!!треба перевірити чи буде все копіюваитись без стилю ранів
    #new_run.style = run.style.name

    new_run.bold = run.bold
    new_run.italic = run.italic
    new_run.underline = run.underline
    #new_run.font.size = run.font.size
    #new_run.font.name = run.font.name
    new_run.font.name='Times New Roman'
    new_run.font.size=152400
    new_run.font.color.rgb = run.font.color.rgb
    new_run.font.highlight_color = run.font.highlight_color


def copy_paragraph(target_doc,source_paragraph):
    #target_paragraph = target_doc.add_paragraph(style = source_paragraph.style.name)
    target_paragraph = target_doc.add_paragraph()
    try:
        #if source_paragraph.style.name!='Normal':
        #    print(source_paragraph.style.name)
        #    print(source_paragraph.text[:40])
        #target_paragraph.style = source_paragraph.style
        target_paragraph.style = target_doc.styles[source_paragraph.style.name]
        #if source_paragraph.style.name!='Normal':
        #    print(target_paragraph.style.name)
        target_paragraph.paragraph_format.space_after = Pt(6)
    except KeyError:
        print(f"Warning. Text {source_paragraph.text[:20]} has style{source_paragraph.style}")
    target_paragraph.alignment = source_paragraph.alignment

    for run in source_paragraph.runs:
        new_run = target_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.font.color.rgb = run.font.color.rgb
        new_run.font.highlight_color = run.font.highlight_color
    
    #if source_paragraph.style.name!='Normal':
    #        print(target_paragraph.style.name)
            #print(source_paragraph.text[:40])
   #         print("")

    

def copy_paragraph_list(target_doc,paragraph_list):
    for p in paragraph_list:
        copy_paragraph(target_doc,p)

def copy_paragraph_before(target_doc, paragraph_to_insert_before,source_paragraph):
    target_paragraph = paragraph_to_insert_before.insert_paragraph_before()
    try:
        #target_paragraph.style = source_paragraph.style
        target_paragraph.style = target_doc.styles[source_paragraph.style.name]
        target_paragraph.alignment = source_paragraph.alignment
        #target_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after
        target_paragraph.paragraph_format.space_after = Pt(6)
        #target_paragraph.style.font = source_paragraph.style.font.name
    except KeyError:
        print(f"Warning. Text {source_paragraph.text[:20]} has style{source_paragraph.style}")
    for run in source_paragraph.runs:
        copy_run(target_paragraph,run)
    return target_paragraph

def copy_paragraph_list_before(target_doc, p_to_insert_before,paragraph_list):
    #print("qty of p to insert:",len(paragraph_list))
    for p in paragraph_list:
        copy_paragraph_before(target_doc, p_to_insert_before,p)

BLACK='b'
RED='r' 
def add_text(p,text, color=BLACK):
    r = p.add_run(text)
    r.font.name='Times New Roman'
    r.font.size=152400
    if color == RED:
        r.font.color.rgb = RGB_RED
        r.italic = True

#
#mode = whatever | html
#html - turns red to italics
def format_line(p, handle='', mode=''):
    #handle = "bir"
    if 'b' in handle:
        p.runs[0].font.bold = True
    if 'i' in handle:
        p.runs[0].font.italic = True
    if 'r' in handle:
        p.runs[0].font.color.rgb = RGBColor(0xff, 0x00, 0x00)
    if 'r' in handle and mode == 'html':
        p.runs[0].font.italic = True

    p.runs[0].font.name='Times New Roman'
    p.runs[0].font.size=152400