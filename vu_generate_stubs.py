import re,glob,calendar,docx,os,easygui, shutil, csv,copy
import sys

import ps_docx_utils as pdu
from easygui_timerbox import timerbox
from datetime import datetime
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
import paschalia,get_stichera
import get_saints

from vu_rebuild_octoechos_templates import get_vu_octoechos_variable_parts_from_template as get_octoechos
from vu_rebuild_octoechos_templates import get_template_part_text as get_octoechos_part

from guppy import hpy
h = hpy()

start_time = datetime.now()

RGB_RED = RGBColor(0xff, 0x00, 0x00)


#filenames= glob.glob('*/*/*.txt')
mode = easygui.choicebox('u - Юліанський, g - Григоріанський', 'Вибір календаря', ['u','g'])
#mode='u'



month_no = datetime.now().month+1 if datetime.now().month!=12 else 1
year_no = datetime.now().year if datetime.now().month!=12 else datetime.now().year+1


month_no = timerbox('Вибір місяця', 'Countdown', choices=[month_no, 12 if month_no-1==0 else month_no-1], time=5)

'''
month_no=1
year_no=2026
print(f"WARNING: MONTH OVERRIDE!")
print(f"Processing month: {year_no}-{month_no}")
'''


day_short_dic={"ПН":1,
        "ВТ":2,
        "СР":3,
        "ЧТ":4,
        "ПТ":5,
        "СБ":6,
        "НД":7}
day_short_dic_reversed = {v:k for k,v in day_short_dic.items()}
day_short_dic_string='('+'|'.join(day_short_dic.keys())+')'

day_dic = {"Понеділок":1,
            "Вівторок":2,
            "Середа":3,
            "Четвер":4,
            "П'ятниця":5,
            "Субота":6,
            "Неділя":7}
day_dic_reversed = {v:k for k,v in day_dic.items()}
day_dic["П’ятниця"]=5
day_dic_string='('+'|'.join(day_dic.keys())+')'

month_dic= {'Січень':1,
              'Лютий':2,
              'Березень':3,
              'Квітень':4,
              'Травень':5,
              'Червень':6,
              'Липень':7,
              'Серпень':8,
              'Вересень':9,
              'Жовтень':10,
              'Листопад':11,
              'Грудень':12}
month_dic_reversed = {v:k for k,v in month_dic.items()}
month_dic_string='('+'|'.join([x.lower() for x in month_dic.keys()])+')'
month_w_offset = list(month_dic.values())[4:]+list(month_dic.values())[:4]


if mode == 'u':
    mode_suffix='Юл'
    mode_suffix2='НЮ'
elif mode == 'g':
    mode_suffix='Гр'
    mode_suffix2='ГР'

def get_octoechos_template_files(pattern):
    template_dic = {}
    
    all_files = glob.glob(pattern)
    for f in all_files:
        re_result=re.search(r'docx_resources\\Вечірня-Утреня\\01-Октоїх\\Глас_(\d)\\(\d)-(\w{2})-?.docx',f)
        if re_result:
            #print(re_result.group(2), re_result.group(3))
            doc = docx.Document(f)
            if int(re_result.group(2)) not in template_dic.keys():
                template_dic[int(re_result.group(2))]={}
            if re_result.group(3) in day_short_dic.keys():
                template_dic[int(re_result.group(2))][day_short_dic[re_result.group(3)]]=f
            else:
                print("Something wrong with", f, re_result.group(3))
        else:
            print(f"None found for: {f} using pattern:", r'docx_resources\\Вечірня-Утреня\\01-Октоїх\\Глас_(\d)\\(\d)-(\w{2})-?.docx')
    return template_dic


irmos_list = [1,3,4,5,6,7,8,9]

def get_octoechos_kanon_texts(path):
    doc = docx.Document(path)
    matrix = {}
    for i in range(8):
        matrix[i+1]={}
    for p in doc.paragraphs:
        
        re_result = re.search(r"Глас (\d)",p.text)
        if re_result:
            echos = int(re_result.group(1))
            echos_found = True
            song_found = False
            continue

        if (echos_found or song_found) and re.search(f"{day_dic_string}",p.text):
            week_day_no = day_dic[re.search(f"{day_dic_string}",p.text).group(1)]
            matrix[echos][week_day_no]=[]
            week_day_no_found = True
            echos_found = song_found = False
            continue
        
        if (week_day_no_found or song_found) and re.search("Канон", p.text):
            matrix[echos][week_day_no].append({"label":p.text})
            label_found = True
            week_day_no_found = song_found = False
            continue

        re_result=re.search(r"Пісня (\d)",p.text)
        if (label_found or song_found) and re_result:
            song_no = int(re_result.group(1))
            matrix[echos][week_day_no][-1][song_no]=[p]
            label_found=False
            song_found = True
            continue
        
        if p.text == "":
            continue
        if song_found:
            matrix[echos][week_day_no][-1][song_no].append(p)

    return matrix

def get_kanon_litany(path):
    template_dic={}
    doc = docx.Document(path)
    litany_no = None
    for p in doc.paragraphs:
        re_result = re.search(r"Мала єктенія (\d)",p.text)
        if re_result:
            litany_no=int(re_result.group(1))
            template_dic[litany_no]=[]
            continue
        elif litany_no:
            template_dic[litany_no].append(p)
    return template_dic

def get_menaion_template_files():
    filenames = glob.glob(f'docx_resources\\Вечірня-Утреня\\В,У - Мінея\\{month_no:02}*\\*.docx')
    template_dic={}
    for f in filenames:
        #print("Checking",f)
        pattern=r'(\d{2})-__-(.*?).docx'
        #print(pattern)
        re_result = re.search(pattern,f)
        doc = docx.Document(f)
        template_dic[int(re_result.group(1))]=f
    
    return template_dic

def get_menaion_templates():
    filenames = glob.glob(f'docx_resources/Вечірня-Утреня/В,У - Мінея/{month_no:02}*/*.docx')
    template_dic={}
    for f in filenames:
        print("Checking",f)
        pattern=r'(\d{2})-__-(.*?).docx'
        print(pattern)
        re_result = re.search(pattern,f)
        doc = docx.Document(f)
        template_dic[int(re_result.group(1))]=doc.paragraphs
    return template_dic

def get_resurrection_troparia_texts(path):
    template_dic={}
    doc = docx.Document(path)
    echos = None
    #i=0
    for p in doc.paragraphs:
        re_result = re.search(r"Глас (\d)",p.text)
        if re_result:
            #print("found")
            echos=int(re_result.group(1))
            template_dic[echos]=[]
        #print(p.text)

        re_result = re.search(r"((Тропар|Кондак).*?\(г\. \d\)): (.*)",p.text)
        if echos and re_result:
            #template_dic[echos].append([re_result.group(1),re_result.group(3)])
            template_dic[echos].append(p)
    return template_dic

def get_menaion_troparia_texts(path):
    #print("Building Menaion Troparia matrix")
    template_dic={}
    doc = docx.Document(path)
    key = None #cur_day
    saint={}
    for p in doc.paragraphs:
        re_result = re.search(r"^(\d+)( |\.)(.*)",p.text)
        if re_result:
            #print(f"found day {re_result.group(1)}")
            key=int(re_result.group(1))
            cur_day_heading = re_result.group(2)
            template_dic[key]=[]

        if key:
            if re.search("^Тропар",p.text):
                re_result = re.search(r"(Тропар.*?\(г\. \d\): )(.*)",p.text)
                try:
                    #saint["troparion"]=[re_result.group(1), re_result.group(2)]
                    saint["troparion"]=p
                except IndexError as e:
                    print(p.text, re_result)
                    raise e
            elif re.search("^Кондак",p.text):
                re_result = re.search(r"(Кондак.*?\(г\. \d\): )(.*)",p.text)
                #saint["kondakion"]=[re_result.group(1), re_result.group(2)]
                saint["kondakion"]=p
            else:
                #found saint header
                saint["header"]=p.text    

        if "header" in saint.keys() and "troparion" in saint.keys() and "kondakion" in saint.keys():
            template_dic[key].append(saint)
            saint={}
        #    template_dic[echos].append([re_result.group(1),re_result.group(3)])
    return template_dic

lenten_tr_endings = {
1:'заступництвом безтілесних Твоїх помилуй нас.',
2:'задля молитов Предтечі Твого, помилуй нас.',
3:'силою Хреста Твого, охорони нас.',
4:'задля молитов святих апостолів Твоїх і святителя Миколая, помилуй нас.',
5:'силою Хреста Твого, охорони нас.'}

def get_trinity_troparia_lenten_exapostolaria(path):
    dic = {'tr':{},'ex':{}}
    for i in range(1,9):
        dic['tr'][i] = []
        dic['ex'][i] = []
    doc = docx.Document(path)
    header = None
    echos = None

    for p in doc.paragraphs:
        if "ропарі" in p.text:
            header = 'tr'
            continue
        if "Світильні" in p.text:
            header = 'ex'
            continue

        re_result = re.search(r"Глас (\d)",p.text)
        if re_result:
            echos = int(re_result.group(1))
            continue

        if header and echos:
            if not p.text.startswith('(:)'):
                dic[header][echos].append(p)

    return dic['tr'],dic['ex']



def insert_trinity_troparia(path, date):

    paragraph_list = copy.deepcopy(lent_trinity_troparia[paschalia.get_echos(date,mode)])
    paragraph_list[0].runs[-1].text += " "+lenten_tr_endings[date.weekday() + 1]

    doc = docx.Document(path)
    section_start_found = False
    for p in doc.paragraphs:

        if p.text.startswith("Троїчні"):
            section_start_found = True
            continue

        if not p.text.startswith("Після першого") and section_start_found:
            pdu.delete_paragraph(p)
            continue

        if p.text.startswith("Після першого"):
            pdu.copy_paragraph_list_before(doc,p,paragraph_list)
            break

    doc.save(path)

#для гласів 1, 5, 6, 7
lenten_ex_endings = {
1:'заступництвом безтілесних Твоїх, і спаси мене.',
2:'задля молитов Предтечі Твого, і спаси мене.',
3:'силою Хреста Твого, і спаси мене.',
4:'задля молитов святих апостолів Твоїх і святителя Миколая, і спаси мене.',
5:'силою Хреста Твого, і спаси мене.'}

#для гласів 2, 3, 4, 8
lenten_ex_h_endings = {
1:'заступництвом безтілесних Твоїх, і спаси мене.',
2:'задля молитов, Господи, Предтечі Твого, і спаси мене.',
3:'силою, Господи, Хреста Твого, і спаси мене.',
4:'задля молитов святих апостолів Твоїх і святителя Миколая, і спаси мене.',
5:'силою, Господи, Хреста Твого, і спаси мене.'}

def insert_lent_exapostolaria(path, date):
    echos = paschalia.get_echos(date,mode)
    paragraph_list = copy.deepcopy(lent_exapostolaria[echos])
    if echos in (1,5,6,7):
        paragraph_list[0].runs[-1].text += " "+lenten_ex_endings[date.weekday() + 1]
    elif echos in (2,3,4,8):
        paragraph_list[0].runs[-1].text += " " + lenten_ex_h_endings[date.weekday() + 1]

    doc = docx.Document(path)
    section_start_found = False
    for p in doc.paragraphs:

        if p.text.startswith("Світильний"):
            section_start_found = True
            continue

        if not p.text.startswith("Хвалитні") and section_start_found:
            pdu.delete_paragraph(p)
            continue

        if p.text.startswith("Хвалитні"):
            pdu.copy_paragraph_list_before(doc,p,paragraph_list)
            break

    doc.save(path)

def get_vespers_prokimenon(path):
    template_dic={}
    doc = docx.Document(path)
    day=None
    for p in doc.paragraphs:
        re_result = re.search(f"^{day_dic_string}",p.text)
        if re_result:
            day=re_result.group(1)
            template_dic[day_dic[day]]=[]
        elif day:
            template_dic[day_dic[day]].append(p)
    return template_dic

def get_matrix_full(csv_filename):
    matrix=[]
    with open(csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:
            matrix.append(row)
    return matrix

def get_matrix(csv_filename):
    matrix={}
    with open(csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:
            if row[0]==str(month_no):
                matrix[int(row[1].split('.')[2])]=row[2:]
    return matrix

#need to merge with previous one
def get_dismissal_matrix(dismissal_csv_filename,cur_month):
    matrix={}
    with open(dismissal_csv_filename, newline='', encoding='utf-8') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in spamreader:
            if row[0]==month_dic_reversed[cur_month]:
                matrix[int(row[1].split('.')[0])]=row[2:]
    return matrix



def get_theotokion_troparia_texts(path):
    service_dic = {'Вечірня':'vespers',
                   'Утреня':'orthros'}
    template_dic={}
    doc = docx.Document(path)
    for p in doc.paragraphs:
        re_result = re.search(r"Глас (\d)",p.text)
        if re_result:
            #print("found")
            echos=int(re_result.group(1))
            template_dic[echos]=[]
        #print(p.text)

        re_result = re.search(r"(Вечірня|Утреня) Відпуст",p.text)
        if re_result:
            #print(echos,re_result.group(1))
            service=service_dic[re_result.group(1)]

        re_result = re.search(f"^{day_dic_string}",p.text)
        if re_result:
            day = re_result.group(1)
            weekday_no = day_dic[day]

        re_result = re.search(f"Богородичний",p.text)
        if re_result:
            template_dic[echos].append({'service':service, 'weekday':weekday_no,'p':p})
            
    return template_dic

def insert_prokimenon(path,day):
    doc = docx.Document(path)
    #print(path)
    vespers_prokimenon_found = False
    vespers_prokimenon_end_found = True
    for p in doc.paragraphs:
        re_result = re.search("Прокімен",p.text)
        if re_result:
            vespers_prokimenon_found = True
            #delete_paragraph(p)
        re_result = re.search("(Читання|Сподоби, Господи)",p.text)
        if re_result:
            #print(day)
            for p1 in vespers_prokimenon[day]:
                pdu.copy_paragraph_before(doc,p,p1)
            p.insert_paragraph_before()
            break
        elif vespers_prokimenon_found:
            #print(f"Deleting {p.text}")
            pdu.delete_paragraph(p)
    doc.save(path)

def insert_boh_hospod_echos(path,date):
    doc = docx.Document(path)
    for p in doc.paragraphs:
        re_result = re.search("Бог Господь",p.text)
        if re_result:
            #print("found")
            if date.weekday()+1==7:
                echos = paschalia.get_echos(date,mode)
            else:
                #echos = int(re.search("(\d)",templates_menaion[date.day][0]['troparion'][0]).group(1))
                echos = int(re.search(r"(\d)",troparia_menaion[date.day][0]['troparion'].text).group(1))

            for r in p.runs:
                if '?' in r.text:
                    #print(r.text)
                    r.font.highlight_color=None
                    r.text = re.sub(r'(\?+)',str(echos),r.text)
            break
    doc.save(path)


def insert_aliluia_echos(path,date):
    doc = docx.Document(path)
    for p in doc.paragraphs:
        re_result = re.search("^Алилуя на глас",p.text)
        if re_result:
            echos = paschalia.get_echos(date, mode)

            for r in p.runs:
                r.text = re.sub(r'(\d)',str(echos),r.text)
            break
    doc.save(path)

def insert_first_sessional_hymn(path,date):
    print("Inserting sessional hymn for:", date)
    section_start_found=False
    #section_end_ = False
    doc = docx.Document(path)
    echos = paschalia.get_echos(date,mode)
    wd = date.weekday()+1

    dic = get_octoechos(echos,wd)
    paragraph_list = get_octoechos_part(dic,'сідальний1')[1:3]


    for p in doc.paragraphs:
        if p.text.startswith("Після першого"):
            #print("found sid header")
            section_start_found = True
            continue

        if len(p.text)>0 and section_start_found:
            #print("deleting:", p.text[:20])
            pdu.delete_paragraph(p)

        if len(p.text)==0 and section_start_found:
            #print("inserting",paragraph_list[0].text[:20])
            pdu.copy_paragraph_list_before(doc,p,paragraph_list)
            break
    doc.save(path)

def insert_gv_echos(path):
    #print("Inserting gv for ",path)
    doc = docx.Document(path)

    first_gv_echos = None
    for p in doc.paragraphs:

        re_result = re.search(r'\(г. (\d)',p.text)
        if re_result:
            first_gv_echos=re_result.group(1)
            #print("Found gv echos:", first_gv_echos)
            break

    for p in doc.paragraphs:
        if first_gv_echos and 'Два перші стихи 140-го псалма' in p.text:
            #print("Changing:",p.runs[0].text)
            p.runs[0].text = re.sub(r'Два перші стихи 140-го псалма, на глас \d.',f'Два перші стихи 140-го псалма, на глас {first_gv_echos}.',p.runs[0].text )
            #print("Changed:", p.runs[0].text)
            break

    doc.save(path)

def insert_prefix_to_paragraph(doc,p,text="Слава: ",formatting='b'):
    p_new = pdu.copy_paragraph_before(doc,p,p)
    #print("iptp:    ",p_new.text)
    new_run = p_new.add_run(text)
    if 'b' in formatting:
        new_run.font.bold = True
    new_run.font.name='Times New Roman'
    new_run.font.size=152400
    
    bkp=p_new.runs[:-1]
    for r in p_new.runs:
        pdu.delete_run(r)
        
    pdu.copy_run(p_new,new_run)
    for r in bkp:
        pdu.copy_run(p_new,r)
    return p_new

def get_troparia_theotokion(echos, date, service):
    #day_details = paschalia.get_day_details(date, mode)
    if paschalia.get_day_details(date, mode)[0]==('pascha'):
        day = 7
    else:
        day = date.weekday()+1


    #print(echos, date, service)
    res = filter(lambda t: t['service'] == service and t['weekday'] == day, templates_theotokion_dic[echos])
    l = list(res)
    p = l[0]["p"]
    return p
    #return list(filter(lambda t: t['service'] == service and t['weekday'] == date.weekday()+1, templates_theotokion_dic[echos]))[0]["p"]

def get_troparia_block(date,service):
    #if date.day!=1:
    #    return 0


    
    troparia_block=[]
    troparia=None
    if ordo_matrix[date.day][1]=='y':
        print("Something wrong")
        raise Exception("Should have a ready template")
    saint_counter=0
    #print("preparing troparia:",date.day, service)
    for i in list(range(4)):
        #if date.day==2:
        #    print(f"processing troparia for day{d}. item {i}")
        try:
            if not ordo_matrix[date.day][2:][i]:
                continue
        except IndexError as e:
            print('erred on:',i)
            raise e
            
            
        if ordo_matrix[date.day][2:][i] == 'resurrection':
            #if date.day==2:
            #    print(f"inserting troparia for day{d}. item {i}")
            troparia = troparia_resurrection[paschalia.get_echos(date, mode)][0]
        if ordo_matrix[date.day][2:][i] == 'special':
            troparia = troparia_special[date.day][0]["troparion"]
        if ordo_matrix[date.day][2:][i] == 'triodion':
            if i==3:
                troparia=insert_prefix_to_paragraph(None,troparia_triodion[date.day][0]["troparion"],text="І нині: ")
            else:
                troparia = troparia_triodion[date.day][0]["troparion"]
        if ordo_matrix[date.day][2:][i] == 'saint':
            #print(date.day, i)
            try:
                if i==2:
                    #print("adding Slava")
                    #print(templates_menaion[date.day][saint_counter]["troparion"].text)
                    troparia=insert_prefix_to_paragraph(None,troparia_menaion[date.day][saint_counter]["troparion"],text="Слава: ")
                    #print(troparia.text)
                else:
                    troparia=troparia_menaion[date.day][saint_counter]["troparion"]
            except IndexError as e:
                print(date.day, e)
                print("Check if troparion is in the file")
                print("Check if menaion matrix is populated")
                raise e

            saint_counter+=1

        if ordo_matrix[date.day][2:][2]:
            prefix = "І нині: "
        else:
            prefix = "Слава, і нині: "


        if ordo_matrix[date.day][2:][i] == 'theotokos':
            echos = int(re.search(r"(\d)",troparia_block[-1].text).group(1))

            troparia = insert_prefix_to_paragraph(None, get_troparia_theotokion(echos, date, service),text=prefix)

        if ordo_matrix[date.day][2:][i] == 'feast':
            if troparia_feast and date.day in troparia_feast:
                troparia = insert_prefix_to_paragraph(None, troparia_feast[date.day][0]['troparion'],text=prefix)
            else:
                print(f"Увага: нема святкового тропаря для {date.day}")
                continue

        troparia_block.append(troparia)
    
    #???
    troparia_block.append(troparia_block[-1].insert_paragraph_before())
    #print([item.text[:36] for item in troparia_block])

    for p in troparia_block:
        re_result = re.search(r"^(.*?)(Богородичний |Тропар |Кондак )(\(г. \d\): )(.*)",p.text)
        if re_result:
            for r in p.runs:
                pdu.delete_run(r)
            #якщо є Слава і Нині - переносимо
            if re_result.group(1):
                r_new = p.add_run(re_result.group(1))
                format_run(r_new,'b')
            #якщо це не Богородичний - переносимо глас.
            if not (re_result.group(2).startswith('Б')):
                r_new = p.add_run(re_result.group(3))
                format_run(r_new,'ri')
            #Переносимо все решта
            r_new = p.add_run(re_result.group(4))
            format_run(r_new,'')

    #print([item.text[:36] for item in troparia_block])
    return troparia_block

def insert_troparia(path,date):
    print(date.day)
    doc = docx.Document(path)
    troparia_block = {"orthros": get_troparia_block(date,'orthros'),
                      #"orthros_dism": get_troparia_block(date, 'orthros')
                      "vespers": get_troparia_block(date,'vespers')}
    #troparia_block=get_troparia_block(date,'orthros')
    troparion_found = False
    troparion_end_found = False
    troparion_label_count=0
    service = False
    for p in doc.paragraphs:
        #if troparion_found and date.day==8:
        #    print(p.text[:20])
        if re.search("ВЕЧІРНЯ",p.text):
            service = 'vespers'
        if re.search("УТРЕНЯ",p.text):
            service = 'orthros'
        
        re_result = re.search("Тропарі",p.text)
        if re_result:
            #print(f"Day {date.day}. Troparion label for service = {service} found")
            troparion_found = True
            troparion_end_found = False
            troparion_label_count+=1
            continue
        
        re_result = re.search("(Великий відпуст|.ктенія усильного благання|Мала .ктенія)",p.text)

        if troparion_found and re_result:
            troparion_end_found=True
            troparia_block_to_copy = copy.deepcopy(troparia_block[service])

            #Додаємо (2 р.) до першого тропаря на Бог Господь.
            if service == 'orthros' and troparion_label_count ==2:
                add_text(troparia_block_to_copy[0]," (2 р.)","r")

            #міняємо богородичний на богородичний вечірні/
            if service == 'orthros' and re_result.group(1).startswith('М'):
                troparia_block_to_copy[-2]=copy.deepcopy(troparia_block['vespers'][-2])

            pdu.copy_paragraph_list_before(doc,p,troparia_block_to_copy)
            troparion_found=False
        elif  troparion_found and not troparion_end_found and ordo_matrix[date.day][1]!='y':
            #pass
            pdu.delete_paragraph(p)
    if troparion_label_count!=3 and date.weekday()+1 !=7:
        print(f"WARNING! Day {date.day}: troparion_label_count={troparion_label_count}")
    doc.save(path)
    
def format_line(p, handle=''):
    #handle = "bir"
    if 'b' in handle:
        p.runs[0].font.bold = True
    if 'i' in handle:
        p.runs[0].font.italic = True
    if 'r' in handle:
        p.runs[0].font.color.rgb = RGBColor(0xff, 0x44, 0x00)
    p.runs[0].font.name='Times New Roman'
    p.runs[0].font.size=152400

def format_run(r, handle=''):
    #handle = "bir"
    if 'b' in handle:
        r.font.bold = True
    else:
        r.font.bold = False
        
    if 'i' in handle:
        r.font.italic = True
    else:
        r.font.italic = False

    if 'r' in handle:
        r.font.color.rgb = RGBColor(0xff, 0x44, 0x00)
    else:
        r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    r.font.highlight_color=None
    r.font.name='Times New Roman'
    r.font.size=152400

def insert_header(path,date):
    #print(date)
    doc = docx.Document(path)
    for p in doc.paragraphs:
        re_result = re.search("ВЕЧІРНЯ",p.text)
        if not re_result:
            pdu.delete_paragraph(p)
        else:
            #Шапка в форматі "Травень 10 Середа"
            day_name = day_dic_reversed[date.weekday()+1]
            month_name = month_dic_reversed[month_no]
            week_no=paschalia.get_week(date,"","")[:2]
            
            txt =" ".join([month_name,str(date.day)+",",day_name])
            p_new = p.insert_paragraph_before(txt)
            format_line(p_new, '')

            #Субота, Неділя, Тиждень etc...

            special_dates=[{"date":datetime(year_no,12,25),
                             "holiday":"Різдво",
                             "holiday_locative":"Різдві",
                             "holiday_instrumental":"Різдвом"},
                            {"date":datetime(year_no+1,1,6),
                             "holiday":"Богоявленні",
                             "holiday_locative":"Богоявленні",
                             "holiday_instrumental":"Богоявленням"}]
            for sd in special_dates:
                diff = (sd["date"]-date).days
                if abs(diff)<=7 and date.weekday()+1 in [6,7]:
                    if diff>0:
                        p_new = p.insert_paragraph_before(f"{day_name} перед {sd['holiday_instrumental']}")
                    else:
                        p_new = p.insert_paragraph_before(f"{day_name} по {sd['holiday_locative']}")
                    format_line(p_new, '')
            
            if date.weekday()+1 in [6,7]:
                p_new = p.insert_paragraph_before(f"{day_name} {week_no} по П'ятидесятниці")
                format_line(p_new, '')

            if date.weekday()+1 in [1]:
                p_new = p.insert_paragraph_before(f"Тиждень {week_no} по П'ятидесятниці")
                format_line(p_new, '')

            #перелік святих
            lst = filter(lambda l:int(l[0])==date.month and int(l[1])==date.day,saint_matrix[1:])
            for l in lst:
                p_new=p.insert_paragraph_before(l[9])
                format_line(p_new, ''.join(l[2:5]))
            break
    doc.save(path)

def insert_header_from_dismissal_matrix(path,date):
    #print(date)
    doc = docx.Document(path)
    for p in doc.paragraphs:
        re_result = re.search("ВЕЧІРНЯ",p.text)
        if not re_result:
            pdu.delete_paragraph(p)
        else:
            day_name = day_dic_reversed[date.weekday()+1]
            month_name = month_dic_reversed[month_no]
            txt =" ".join([month_name,str(date.day)+",",day_name])
            p_new = p.insert_paragraph_before(txt)
            format_line(p_new, 'b')
            p_new = p.insert_paragraph_before(dismissal_matrix[date.day][2])
            format_line(p_new, '')
            break
    doc.save(path)    


def stretch_texts(qty_of_verses_requested, texts):
    result = []
    if qty_of_verses_requested > len(texts):
        repetitions_per_text, remaining_verses = divmod(qty_of_verses_requested, len(texts))
        for t in texts:
            result.append([t]*repetitions_per_text)
        for i in range(remaining_verses):
            result[i].append(result[i][0])
        return sum(result,[])
    else:
        return texts[:qty_of_verses_requested]

# робимо всі "Священик" червоними та italic
BLACK='b'
RED='r'   
def add_text(p,text, color=BLACK):
    r = p.add_run(text)
    r.font.name='Times New Roman'
    r.font.size=152400
    if color == RED:
        r.font.color.rgb = RGB_RED
        r.italic = True
        
def insert_dismissal(path,date):
    #print("Searching for dismissal")
    doc = docx.Document(path)
    shoutout_found = None
    for p in doc.paragraphs:
        re_result=re.search(r"\(3 р\.\)(\.)? Благослов(и|и́)\.",p.text)
        if re_result:
            shoutout_found = True
            #print(date.day, "Благослови found!")
            continue

        if shoutout_found:
            #print(date.day, "Відпуст inserting")
            p_new=p.insert_paragraph_before(dismissal_matrix[date.day][9])
            p_new.paragraph_format.space_after = Pt(6)
            #print(date.day,p_new.text)
            format_line(p_new, '')
            pdu.delete_paragraph(p)
            shoutout_found = False

    for p in doc.paragraphs:
        if re.search(f'^Священник:',p.text):
            re_result=re.search(f'^(Священник:)( .+?)(якого є храм)(.*?)$',p.text)
            p_bak=p.text
            p.clear()

            try:
                if re_result:
                    add_text(p,re_result.group(1),color=RED)
                    add_text(p,re_result.group(2))
                    add_text(p,re_result.group(3),color=RED)
                    add_text(p,re_result.group(4))
                else:
                    re_result=re.search(f'^(Священник:)(.*?)$',p_bak)
                    add_text(p,re_result.group(1),color=RED)
                    add_text(p,re_result.group(2))
            except:
                print(p_bak)
                raise
    doc.save(path)


#obsolete values
stichos_list_old=[
    "Стих: Виведи з в'язниці мою душу",
    "Стих: Мене обступлять праведники",
    "Стих: З глибин взиваю до тебе, Господи",
    "Стих: Нехай будуть твої вуха уважні",
    "Стих: Коли ти, Господи, зважатимеш на беззаконня",
    "Стих: Задля імени твого надіюсь на тебе, Господи",
    "Стих: Від ранньої сторожі до ночі",
    "Стих: Бо в Господа милість і відкуплення велике в нього",
    "Стих: Хваліте Господа всі народи",
    "Стих: Велике бо до нас його милосердя"
]


stichos_list=[
    "Стих: Ви́веди з в’язни́ці мою́ ду́шу",
    "Стих: Мене́ обсту́плять пра́ведники",
    "Стих: З глиби́н взива́ю до Те́бе, Го́споди",
    "Стих: Неха́й бу́дуть Твої́ ву́ха ува́жні",
    "Стих: Коли́ Ти, Го́споди, зважа́тимеш на беззако́ння, ",
    "Стих: За́для І́мени Твого́ наді́юсь на Те́бе",
    "Стих: Від ра́нньої сторо́жі до но́чі",
    "Стих: Бо в Го́спода ми́лість і відку́плення вели́ке в Ньо́го",
    "Стих: Хвалі́те Го́спода всі наро́ди",
    "Стих: Вели́ке бо до нас Його́ милосе́рдя"
]

def insert_menaion_stichera(path,date,template_type):
    #print("Вставляємо стихири з Мінеї:", path)
    doc = docx.Document(path)
    stichera_no=None
    look_for_slava=False
    k=date.day
    delta = None
    dogmatikos = triodion_doxa = None
    #print(f"inserting {k}")
    #required_stichera_qty = stichera_gv_matrix[k][1]

    if date.weekday() + 1 == 6:
        #print(f"Стихири: пропускаємо {k}, бо субота, обернений порядок стихир")
        #return -2
        if stichera_gv_matrix[k][1]=='3':
            delta = 3
        #else:
        #    print(f"Стихири: пропускаємо {k}, бо субота і стихир {stichera_gv_matrix[k][1]}")
            #return -2

    #TODO: зайва умова, перенесена на рівень вище.
    if stichera_gv_matrix[k][1].isnumeric() and stichera_gv_matrix[k][1]!='0':
        #stichos_dic =init_stichera_dic(stichera_gv_matrix[k][1])
        n= int(stichera_gv_matrix[k][1])
        '''
        if date.weekday() + 1 in (1,2,3,4,5,7):
            stichos_dic = {stichos_list[-n:][i]: i for i in range(n) }
        if date.weekday() + 1 in (6,) and template_type=="октоїх":
            stichos_dic = {stichos_list[-n - delta :-delta][i]: i for i in range(n) }
        '''
        if date.weekday() + 1 in (6,) and template_type=="октоїх" and stichera_gv_matrix[k][1]=='3':
            stichos_dic = {stichos_list[-n - delta :-delta][i]: i for i in range(n) }
            stichos_dic_old = {stichos_list_old[-n - delta :-delta][i]: i for i in range(n) }
        else:
            stichos_dic = {stichos_list[-n:][i]: i for i in range(n) }
            stichos_dic_old = {stichos_list_old[-n:][i]: i for i in range(n)}
        #print(len(stichos_dic),stichos_dic)

        #fallback to using old values


    #Якщо в матриці "0" - виходимо з процедури
    else:
        #print(f"Для {date.day} в матриці стихир: {stichera_gv_matrix[k][1]}. Пропускаємо.")
        return 0

    #print(stichos_dic)

    stichos_dic_string ='('+'|'.join(stichos_dic.keys())+')'
    stichos_dic_string_old = '(' + '|'.join(stichos_dic_old.keys()) + ')'
    print(date.day,stichos_dic_string)
    print(date.day, stichos_dic_string_old)

    if not 'gv_stichera' in stichera_matrix[k] and stichera_gv_matrix[k][3]: 

            stichera_matrix[k]['gv_stichera'] = generic_stichera_matrix[stichera_gv_matrix[k][3]]['gv_stichera']
            stichera_matrix[k]['gv_doxa'] = generic_stichera_matrix[stichera_gv_matrix[k][3]]['gv_doxa']
            stichera_matrix[k]['gv_theotokion'] = generic_stichera_matrix[stichera_gv_matrix[k][3]]['gv_theotokion']

            for p in (stichera_matrix[k]['gv_stichera']+stichera_matrix[k]['gv_doxa']):
                for r in p.runs:
                    if re.search('(ім’я)',r.text):
                        r.text = r.text.replace('(ім’я)', stichera_gv_matrix[k][4])
                        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        r.font.italic = False
                        r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

            print(f"Стихири: використано стихири заг. служби для {k}")
    
    if 'gv_stichera' in stichera_matrix[k]:
        texts = stretch_texts(n,stichera_matrix[k]['gv_stichera'])
    else:
        print(f"Стихири: пропускаємо {k}, немає стихир ГВ.")
        return -1
    
    delete_empty = False
    for p in doc.paragraphs:
        #stichera_no=None
        re_result = re.search(f"{stichos_dic_string}",p.text)
        re_result_old = re.search(f"{stichos_dic_string_old}", p.text)

        if re_result:
            #print(f"Day {k}, found string:",re_result.group(1)[:20])
            stichera_no=stichos_dic[re_result.group(1)]
            delete_empty = True
            continue

        if re_result_old:
            #print(f"Day {k}, found string:",re_result.group(1)[:20])
            stichera_no=stichos_dic_old[re_result_old.group(1)]
            delete_empty = True
            continue

        #print(f"stichera_no: {stichera_no}; look_for_slava={look_for_slava}")
        if stichera_no or stichera_no==0:
            pdu.copy_paragraph_before(doc,p,texts[stichera_no])
            pdu.delete_paragraph(p)
            #print(f"Inserted stichera {stichera_no} for {v}")
            stichera_no=None
            look_for_slava=True
            continue

        
        if p.text == "" and delete_empty:
            pdu.delete_paragraph(p)
            continue

        re_result = re.search("^(Слава|І нині|Сла́ва|І ни́ні)",p.text)
        if re_result and look_for_slava:
            print("Found Слава/І нині:", p.text[:40])
            if re.search('догмат',p.text):
                dogmatikos = p
            else:
                triodion_doxa = p
                print("Hit triodion_doxa")
            pdu.delete_paragraph(p)

        re_result = re.search("^(Вхід|Світло тихе)",p.text)
        if re_result:
            look_for_slava=False
            if "gv_doxa" in stichera_matrix[k]:
                #print(f"Inserted doxa for {k}")
                pdu.copy_paragraph_before(doc,p,stichera_matrix[k]["gv_doxa"][0])
            #if datetime(year_no,month_no,k).weekday()+1 in (1,2,3,4,5):
            #    copy_paragraph_before(doc,p,doc.add_paragraph("Догмат"))



            #!!!!!
            #Некоректно заміняється приспів. віпрацвати на аних місяця
            if dogmatikos:
                if dogmatikos.runs[0].text.find("Слава")!=-1 and 'gv_doxa' in stichera_matrix[k]:
                    dogmatikos.runs[0].text = dogmatikos.runs[0].text.replace("Слава і","І").replace("Слава, і","І").replace("Слава: І","І")
                pdu.copy_paragraph_before(doc,p,dogmatikos)

            elif triodion_doxa:
                if triodion_doxa.runs[0].text.find("Слава")!=-1 and 'gv_doxa' in stichera_matrix[k]:
                    triodion_doxa.runs[0].text = triodion_doxa.runs[0].text.replace("Слава і","І").replace("Слава, і","І").replace("Слава: І","І")
                elif triodion_doxa.runs[0].text.find("Слава")==-1 and not 'gv_doxa' in stichera_matrix[k]:
                    triodion_doxa.runs[0].text = triodion_doxa.runs[0].text.replace("І нині", "Слава, І нині")

                pdu.copy_paragraph_before(doc,p,triodion_doxa)


            elif datetime(year_no,month_no,k).weekday()+1 in (3,5):
                if  "gv_theotokion" in stichera_matrix[k] and len(stichera_matrix[k]["gv_theotokion"])==2:
                    #print(f"Inserted theo for {k}")
                    pdu.copy_paragraph_before(doc,p,stichera_matrix[k]["gv_theotokion"][1])
                else:
                    p.insert_paragraph_before("ХРЕСТОБОГОРОДИЧНИЙ")
            else:
                #print(f"Inserted theo for {k}")
                try:
                    pdu.copy_paragraph_before(doc,p,stichera_matrix[k]["gv_theotokion"][0])
                except KeyError:
                    #print("Warning, failed to find stichera_matrix[k]['gv_theotokion'][0]")
                    print(f"Стихири: для {k} не знайдено богородичний ГВ в файлі з стихирами")
                    p.insert_paragraph_before("БОГОРОДИЧНИЙ")
            break         

    doc.save(path)

def insert_kanon(path,date):
    doc = docx.Document(path)
    song8_end_found = False
    week_day = date.weekday()+1
    echos = paschalia.get_echos(date,mode)
    kanon_no = 1
    #визначаємо № канону дня, що беремо
    #в піст - 0
    #в інші дні - 1
    '''
    if date>datetime(2024,11,15) == 5:
        kanon_no == 0
    else:
        kanon_no == 1
    '''

    #print(f"Inserting kanon for {date.day} {day_dic_reversed[date.weekday()+1]}")

    kanon_found = False
    song8_end_found = False

    for p in doc.paragraphs:
        
        if re.search("Канон",p.text):
            kanon_found = True
            #print("---знайшли початок канону")
            continue
    
        re_result = re.search ("Пісня Богородиці", p.text)
        if re_result and kanon_found:
            #print("---знайшли 9-ту пісню")

            

            pdu.copy_paragraph_list_before(doc,p,templates_kanon_dic[echos][week_day][kanon_no][1])
            pdu.copy_paragraph_list_before(doc,p,templates_kanon_dic[echos][week_day][kanon_no][3])
            pdu.copy_paragraph_list_before(doc,p,kanon_litany_list[1])
            pdu.copy_paragraph_list_before(doc,p,templates_kanon_dic[echos][week_day][kanon_no][4])
            pdu.copy_paragraph_list_before(doc,p,templates_kanon_dic[echos][week_day][kanon_no][5])
            pdu.copy_paragraph_list_before(doc,p,templates_kanon_dic[echos][week_day][kanon_no][6])
            pdu.copy_paragraph_list_before(doc,p,kanon_litany_list[2])
            pdu.copy_paragraph_list_before(doc,p,templates_kanon_dic[echos][week_day][kanon_no][7])
            pdu.copy_paragraph_list_before(doc,p,templates_kanon_dic[echos][week_day][kanon_no][8])
            break
            
        if kanon_found and not song8_end_found:
            pdu.delete_paragraph(p)
    
    for p in doc.paragraphs:
        if re.search(r"(Пісня \d|Мала .ктенія)",p.text):
            p.style="Heading 3"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #p.paragraph_format.space_after = Pt(6)

    #print("WARNING: не знайшли закінчення канону!")
    doc.save(path)

def insert_resurrection_gospel_parts(path,date):
    doc = docx.Document(path)
    
    gospel_no = paschalia.get_resurrection_gospel(date,mode)

    gospel_stichera_label_found = None

    for p in doc.paragraphs:

        re_result = re.search("Священ{1,2}ик: Від ВКАЗАТИ святого Євангелія читання.",p.text)
        if re_result:
            for r in p.runs:
                if "ВКАЗАТИ" in r.text:
                    r.text = r.text.replace("ВКАЗАТИ",resurrection_gospel_matrix[gospel_no]['gospel_apostle'])
                    format_run(r,'')
            continue
        
        re_result = re.search("ТЕКСТ ЄВАНГЕЛІЯ",p.text)
        if re_result:
            pdu.copy_paragraph_list_before(doc,p,resurrection_gospel_matrix[gospel_no]['gospel'])
            pdu.delete_paragraph(p)
            continue

        re_result = re.search("ВСТАВИТИ СВІТИЛЬНИЙ",p.text)
        if re_result:
            pdu.copy_paragraph_list_before(doc,p,resurrection_gospel_matrix[gospel_no]['exapostolarion'])
            pdu.delete_paragraph(p)
            continue

        re_result = re.search("Стихира євангельська",p.text)
        if re_result:
            gospel_stichera_label_found=True
            continue
        
        re_result = re.search(r"Слава \(г. \): ",p.text)
        if gospel_stichera_label_found and re_result:
            pdu.copy_paragraph_list_before(doc,p,resurrection_gospel_matrix[gospel_no]['stichera'])
            pdu.delete_paragraph(p)
            gospel_stichera_label_found = None
            continue

    doc.save(path)

folder_name=f'..\\ps_drafts\\vu\\{year_no}-{month_no:02}-{mode}'
if not os.path.exists(folder_name):
    os.makedirs(folder_name)

print(year_no, month_no)

def get_saints_for_filename(month,day,multiline=False):
    index = 10
    separator = ','
    saint_string=""
    for row in saint_matrix[1:]:
        if int(row[0])==month and int(row[1])==day:
            if row[index]:
                if saint_string:
                    saint_string+=separator+row[index]
                else:
                    saint_string=row[index]
    return saint_string

def get_polyeleos_saints_for_filename(month,day,multiline=False):
    index = 10
    separator = ','
    saint_string=""
    for row in saint_matrix[1:]:
        if int(row[0])==month and int(row[1])==day and row[6] and row[6] in '+*#':
            if row[index]:
                if saint_string:
                    saint_string+=separator+row[index]
                else:
                    saint_string=row[index]
    return saint_string

draft_dic={}

def create_stubs():

    print("commencing creation")
    
    draft_dic={}
    for d in range(1,calendar.monthrange(year_no, month_no)[1]+1):
    #for d in range(1,5):    
        draft_dic[d]=[]
        print("Дата:",d,datetime(year_no, month_no, d).weekday()+1,mode)
        day_details = paschalia.get_day_details(datetime(year_no,month_no,d),mode)
        print(d, day_details)
        filename_triodion_row = None
        filename_menaion_string = get_saints_for_filename(month_no,d)

        for row in filename_triodion_matrix:

            if day_details[0] == 'pentecost' and day_details[2]:
                if day_details[0] == row[0] and row[2] and day_details[2] == int(row[2]) and day_details[3] == int(row[3]):
                    filename_triodion_row = row
            else:
                if day_details[0] == row[0] and day_details[1] == int(row[1]) and day_details[3] == int(row[3]) and row[12]:
                    filename_triodion_row = row


        dest_filename=f'{folder_name}\\{d:02}-{day_short_dic_reversed[datetime(year_no, month_no, d).weekday()+1]}'
        
        
        if (datetime(year_no, month_no, d).weekday()+1==7 and paschalia.get_echos(datetime(year_no,month_no,d),mode)) and not (day_details[0]=='pascha' and day_details[1]==1):
            dest_filename+=f'-Гл.{paschalia.get_echos(datetime(year_no,month_no,d),mode)}'
        
        if filename_triodion_row and filename_triodion_row[12]:
            dest_filename+=f'-{filename_triodion_row[12]}'
            if get_polyeleos_saints_for_filename(month_no,d):
                dest_filename+=f'-{get_polyeleos_saints_for_filename(month_no,d)}'
            dest_filename+=f'.docx'
        elif filename_menaion_string:
            dest_filename+=f'-{filename_menaion_string}.docx'
        else:
            print(f"УВАГА: відсутнє коротке ім'я для {d}")
            dest_filename+=f'-___.docx'
        
        #print(expected_template_path)

        #Сирні Середа, П'ятниця
        if day_details[0]=='pentecost' and day_details[3]in (3,5,6) and day_details[2]==1:
            expected_lent_template_path = f"docx_resources\\Вечірня-Утреня\\В,У - Пісна Тріодь\\тиждень-0\\0-{day_details[3]}-ВУ.docx"
            if expected_lent_template_path in lent_templates:
                src_filename=expected_lent_template_path
                draft_dic[d].append('сирний т-нь')
            #print(d, "Шаблон Воскресіння")
        elif day_details[0]=='lent':
            expected_lent_template_path = f"docx_resources\\Вечірня-Утреня\\В,У - Пісна Тріодь\\тиждень-{day_details[1]}\\{day_details[1]}-{day_details[3]}-ВУ.docx"    
            if expected_lent_template_path in lent_templates and (day_details[3]!=7 or day_details[1] in [6,7]):
                src_filename=expected_lent_template_path
                draft_dic[d].append('піст')
                #print(d, "Шаблон Великого Посту")
            else:
                src_filename = templates_octoechos_dic[paschalia.get_echos(datetime(year_no, month_no, d), mode)][datetime(year_no, month_no, d).weekday() + 1]
                draft_dic[d].append('октоїх')
        elif day_details[0] in ('pascha'):
            expected_template_path = f"docx_resources\\Вечірня-Утреня\\В,У - Квітна Тріодь\\pascha-{day_details[1]}\\{day_details[1]}-{day_details[3]}-ВУ.docx"
            #print(expected_template_path)
            if expected_template_path in pascha_pentecost_templates:
                src_filename=expected_template_path
                draft_dic[d].append('пасха')

        elif day_details[0] == 'pentecost':
            expected_template_path = f"docx_resources\\Вечірня-Утреня\\В,У - Квітна Тріодь\\pentecost-{day_details[1]}\\{day_details[1]}-{day_details[3]}-ВУ.docx"
            #print(expected_template_path)
            if expected_template_path in pascha_pentecost_templates:
                src_filename=expected_template_path
                draft_dic[d].append('50-ця')

            elif d in templates_menaion_dic.keys() and datetime(year_no, month_no, d).weekday()+1!=7:
                src_filename=templates_menaion_dic[d]
                draft_dic[d].append('мінея')
                #print(d, "Шаблон Мінеї")
            else:        
                src_filename=templates_octoechos_dic[paschalia.get_echos(datetime(year_no,month_no,d),mode)][datetime(year_no, month_no, d).weekday()+1]
                draft_dic[d].append('октоїх')
                #print(d, "Шаблон Октоїха")
        #else:
        #    src_filename = templates_octoechos_dic[paschalia.get_echos(datetime(year_no, month_no, d), mode)][datetime(year_no, month_no, d).weekday() + 1]
        #    draft_dic[d].append('октоїх')
        try:
            #print("Trying to copy:",src_filename,dest_filename)
            shutil.copy2(src_filename,dest_filename)
        except NameError as e:
            print(e)
            print("Помилка для дня ", d)
            raise e
        except OSError as e:
            print(e)
            print("Помилка імені файлу дня ", d)
            raise e

        draft_dic[d].append(dest_filename)
    print("Завершено створення чернеток УВ!")
    return draft_dic

def update_ektenia_before_kanon(path,date):
    doc = docx.Document(path)
    saints_str = get_saints.get_saints_full(date)

    for p in doc.paragraphs:
        if "Спаси́, Бо́же, наро́д Твій" in p.text or "Спаси, Боже, народ Твій" in p.text or "Спаси́, Бо́же, наро́д Твій" in p.text:
            re_result = re.search(r'^(.*?)(\(ім’я, що його́ є храм цей\))(.*?)(і всіх .вяти́х:)(.*?)$',p.text)

            if re_result:
                p.text = re_result.group(1)
                format_run(p.runs[0],'')
                new_run = p.add_run(re_result.group(2))
                format_run(new_run, 'ri')
                new_run = p.add_run(', і '+saints_str+', і всіх cвяти́х:'+re_result.group(5))
                format_run(new_run, '')

                '''
                ektenia_text = re_result.group(1)
                ektenia_text += '<red>'+re_result.group(2)+'</red>'+', і '
                ektenia_text += saints_str
                ektenia_text += ', і всіх cвяти́х:'
                ektenia_text += re_result.group(5)
                print(path)
                print(ektenia_text)
                '''

            else:
                print(path, "Ektenia, RE failed")
            break
    doc.save(path)

def update_stubs(draft_dic):
    #drafts = glob.glob(f'{folder_name}\\*.docx')

    for d,desc in draft_dic.items():
        day_details = paschalia.get_day_details(datetime(year_no, month_no,d),mode)
        #print("Дата:",d,datetime(year_no, month_no, d).weekday()+1,mode)
        insert_header_from_dismissal_matrix(desc[1],datetime(year_no, month_no, d))


        if not(day_details[0]=='lent' and day_details[1]==7):
            insert_dismissal(desc[1],datetime(year_no, month_no, d))
        
        if desc[0] in ('неділя','октоїх','пасха','50-ця'):
            #вставити тропарі вечірні, утрені
            if ordo_matrix[d][1]!='y':
                #print(f"Inserting troparia for {datetime(year_no, month_no, d)}")
                insert_troparia(desc[1],datetime(year_no, month_no, d))

            #вставити глас для Бог Господь
            insert_boh_hospod_echos(desc[1],datetime(year_no, month_no, d))
            
            if datetime(year_no, month_no, d).weekday()+1==7:
                #print("Inserting gospel for", datetime(year_no, month_no, d))
                insert_resurrection_gospel_parts(desc[1],datetime(year_no, month_no, d))

            #вставити тропарі вкінці утрені
        if desc[0] in ('октоїх') and (datetime(year_no, month_no, d).weekday()+1)!=7:
            insert_kanon(desc[1],datetime(year_no, month_no, d))

        if desc[0]=='мінея':
            #Якщо мінейний шаблон - вставити прокімен
            insert_prokimenon(desc[1], datetime(year_no, month_no, d).weekday()+1)

        elif stichera_gv_matrix[d][1].isnumeric() and stichera_gv_matrix[d][1]!='0':
            #Якщо шаблон не мінеї і вказані стихири в матриці - вставити стихири ГВ
            insert_menaion_stichera(desc[1],datetime(year_no, month_no, d),desc[0])

        if desc[0]=='піст' and datetime(year_no, month_no, d).weekday()+1==6:
            pass
        day_details = paschalia.get_day_details(datetime(year_no, month_no, d), mode)
        if desc[0]=='піст' and datetime(year_no, month_no, d).weekday()+1 in (1,2,3,4,5) and day_details[1] in (1,2,3,4,5,6):
            insert_trinity_troparia(desc[1],datetime(year_no, month_no, d))
            insert_lent_exapostolaria(desc[1], datetime(year_no, month_no, d))
            insert_aliluia_echos(desc[1], datetime(year_no, month_no, d))
            insert_first_sessional_hymn(desc[1], datetime(year_no, month_no, d))

        # вставити глас для Господи Воззвах
        insert_gv_echos(desc[1])
        # додати святого дня в Єктенію "Спаси Господи" перед каноном (свята, Піст)
        update_ektenia_before_kanon(desc[1], datetime(year_no, month_no, d))


    print("Завершено оновлення чернеток УВ!")
    
def get_files_in_dir():
    draft_dic={}
    path = easygui.diropenbox()
    file_list = glob.glob(f'{path}/*.docx')
    #for item in file_list:
    #    print(item)
    for item in file_list:
        try:
            day=int(item.split('\\')[-1][:2])
            draft_dic[day]=['',item]
        except Exception as e:
            raise e
    return draft_dic

def sanitize_spaces(path):
    doc = docx.Document(path)
    for p in doc.paragraphs:
        for r in p.runs:
            r.text = r.text.replace('\xa0',' ')
    doc.save(path)


lent_templates = glob.glob(r'docx_resources\Вечірня-Утреня\В,У - Пісна Тріодь\*\*.docx')
pascha_pentecost_templates = glob.glob(r'docx_resources\Вечірня-Утреня\В,У - Квітна Тріодь\*\*.docx')

ordo_matrix = get_matrix(f"matrices/vu_тропарі{mode_suffix2}.csv")
stichera_gv_matrix = get_matrix(f"matrices/vu_стихириМінеї{mode_suffix2}.csv")
dismissal_matrix = get_dismissal_matrix(f'matrices/Відпусти{mode_suffix}.csv', month_no)
saint_matrix = get_matrix_full("matrices/Місяцеслов-БД.csv")
filename_triodion_matrix = get_matrix_full("matrices/Місяцеслов-БД-Тріодь.csv")
templates_octoechos_dic = get_octoechos_template_files(r'docx_resources\Вечірня-Утреня\01-Октоїх\*\*.docx')
templates_menaion_dic = get_menaion_template_files()
troparia_resurrection = get_resurrection_troparia_texts('docx_resources/воскресні_тропарі.docx')



troparia_menaion_filename = f'{month_w_offset[month_no - 1]:02}-{month_dic_reversed[month_no].upper()}.docx'
troparia_menaion = get_menaion_troparia_texts(glob.glob(f'docx_resources/Тропарі - Мінея/{troparia_menaion_filename}')[0])

if glob.glob(f'docx_resources\\Вечірня-Утреня\\tmp_тропарі_Тріоді\\тріодь-{month_no:02}-{mode_suffix2}.docx'):
    troparia_triodion = get_menaion_troparia_texts(
        glob.glob(f'docx_resources\\Вечірня-Утреня\\tmp_тропарі_Тріоді\\тріодь-{month_no:02}-{mode_suffix2}.docx')[0])
else:
    troparia_triodion = None

if glob.glob(f'docx_resources\\Вечірня-Утреня\\tmp_тропарі_святкові\\тропарі-святкові-{month_no:02}.docx'):
    troparia_feast = get_menaion_troparia_texts(
        glob.glob(f'docx_resources\\Вечірня-Утреня\\tmp_тропарі_святкові\\тропарі-святкові-{month_no:02}.docx')[0])
elif glob.glob(f'docx_resources\\Вечірня-Утреня\\tmp_тропарі_святкові\\тропарі-святкові-{month_no:02}-{mode_suffix2}.docx'):
    troparia_feast = get_menaion_troparia_texts(
        glob.glob(f'docx_resources\\Вечірня-Утреня\\tmp_тропарі_святкові\\тропарі-святкові-{month_no:02}-{mode_suffix2}.docx')[0])
else:
    troparia_feast = None

if glob.glob(f'docx_resources\\Вечірня-Утреня\\tmp_тропарі_спеціальні\\тропарі-спеціальні-{month_no:02}.docx'):
    troparia_special = get_menaion_troparia_texts(
        glob.glob(f'docx_resources\\Вечірня-Утреня\\tmp_тропарі_спеціальні\\тропарі-спеціальні-{month_no:02}.docx')[0])
elif glob.glob(f'docx_resources\\Вечірня-Утреня\\tmp_тропарі_спеціальні\\тропарі-спеціальні-{month_no:02}-{mode_suffix2}.docx'):
    troparia_special = get_menaion_troparia_texts(
        glob.glob(f'docx_resources\\Вечірня-Утреня\\tmp_тропарі_спеціальні\\тропарі-спеціальні-{month_no:02}-{mode_suffix2}.docx')[0])
else:
    troparia_special = None


vespers_prokimenon = get_vespers_prokimenon(f'docx_resources/Вечірня-Утреня/vu_прокімени.docx')
print("vespers_prokimenon",sys.getsizeof(vespers_prokimenon))

templates_theotokion_dic = get_theotokion_troparia_texts(
    'docx_resources/Вечірня-Утреня/vu_богородичні-тропарі.docx')
print("templates_theotokion_dic",sys.getsizeof(templates_theotokion_dic))

templates_kanon_dic = get_octoechos_kanon_texts('docx_resources/Вечірня-Утреня/05a_ОКТОЇХ_КАНОНИ.docx')
print("templates_kanon_dic",sys.getsizeof(templates_kanon_dic))

kanon_litany_list = get_kanon_litany("docx_resources/Вечірня-Утреня/Канон - Мала єтенія.docx")

stichera_matrix = get_stichera.get_stichera_matrix(
    glob.glob(f'docx_resources/Стихири - Мінея/Мінея_{month_no:02}*.docx')[0])
print("stichera_matrix",sys.getsizeof(kanon_litany_list))

generic_stichera_matrix = get_stichera.get_generic_stichera_matrix(
    'docx_resources/Вечірня-Утреня/Стихири ГВ загальної служби.docx')
print("generic_stichera_matrix", sys.getsizeof(generic_stichera_matrix))

resurrection_gospel_matrix = get_stichera.get_resurrection_gospel_matrix(
    "docx_resources/Вечірня-Утреня/11-ВОСКРЕСНІ ЄВАНГЕЛІЯ.docx")
print("resurrection_gospel_matrix",sys.getsizeof(resurrection_gospel_matrix))

lent_trinity_troparia, lent_exapostolaria = get_trinity_troparia_lenten_exapostolaria(
    r'docx_resources\Вечірня-Утреня\13-ТРОЇЧНІ ТРОПАРІ - СВІТИЛЬНІ.docx')
print("lent_trinity_troparia, lent_exapostolaria", sys.getsizeof(lent_trinity_troparia)+sys.getsizeof(lent_trinity_troparia)  )


# dogmatika = get_stichera.get_dogmatika("Догмати.docx")
print(f"Finished building dictionaries: {(datetime.now() - start_time).total_seconds()}")

choice_list=[
    "1. Згенерувати чернетки",
    "2. Оновити Відпусти",
    "3. Оновити гласи у заголовках",
    "4. Оновити троїчні тропарі та Світильні"
]

#@profile
if __name__ == "__main__":
    start_time = datetime.now()
    print(h.heap())

    action = easygui.choicebox('Виберіть операцію:', 'Вибір операції', choice_list)
    #action="1"
    print(action)

    if action[0]=="1":
        draft_dic = create_stubs()
        for d, desc in draft_dic.items():
            if len(desc)<2:
                print(d,desc)
                raise Exception("Немає дескриптора дня")
        #print(draft_dic)
        update_stubs(draft_dic)
    elif action[0]=="2":
        draft_dic=get_files_in_dir()
        #print(len(draft_dic))
        #print(draft_dic[1])
        for d,desc in draft_dic.items():
            insert_dismissal(desc[1],datetime(year_no, month_no, d))
        print("Завершено оновлення відпустів")
    elif action[0] == "3":
        draft_dic = get_files_in_dir()
        # print(len(draft_dic))
        # print(draft_dic[1])
        for d, desc in draft_dic.items():
            insert_boh_hospod_echos(desc[1], datetime(year_no, month_no, d))
            insert_gv_echos(desc[1])
        print("Завершено оновлення гласів у заголовках")

    elif action[0] == "4":
        draft_dic = get_files_in_dir()
        # print(len(draft_dic))
        # print(draft_dic[1])
        for d, desc in draft_dic.items():
            day_details = paschalia.get_day_details(datetime(year_no, month_no, d),mode)
            if day_details[0] == 'lent' and datetime(year_no, month_no, d).weekday() + 1 in (1, 2, 3, 4, 5) and day_details[1] in (1, 2, 3, 4, 5, 6):
                insert_trinity_troparia(desc[1], datetime(year_no, month_no, d))
                insert_lent_exapostolaria(desc[1], datetime(year_no, month_no, d))
        print("Завершено оновлення Троїчних тропарів та Світильних")
    else:
        print("Нічого не вибрано.")
    print("All done!")
    print(f"Time spent: {(datetime.now() - start_time).total_seconds()}")


