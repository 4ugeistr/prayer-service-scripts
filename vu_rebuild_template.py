import docx
from datetime import datetime

def get_variable_parts_from_template(path):
    #path = find_octoechos_template(echos,day)
    
    doc = docx.Document(path)
    matrix = []
    service = None
    key=None



    for p in doc.paragraphs:
        
        if p.text.lower() == 'вечірня':
            service = 'вечірня'
        if p.text.lower() == 'утреня':
            service = 'утреня'

        if key:
            for item in list_of_lines_that_end_parts:
                if p.text.startswith(item):
                    key= None
                    continue

        #SPAGHETTI CODE. Redo later
        if not key and p.text =='Стихири':
            
            key = 'стихири_ГВ' if service == 'вечірня' else 'стихири_хвалитні'
            matrix.append({'key':key,'text':[p]})
            continue

        if not key and p.text.startswith('Прокімен'):
            key = f'прокімен_{service}'
            matrix.append({'key':key,'text':[p]})
            continue

        if not key and p.text.startswith('Читання'):
            key = f'паремії'
            matrix.append({'key':key,'text':[p]})
            continue

        if not key and p.text.startswith('Стихири на Литії'):
            key = f'стихири_литія'
            matrix.append({'key':key,'text':[p]})
            continue

        if not key and p.text.startswith('Стихири на стиховні'):
            key = f'стиховня_{service}'
            matrix.append({'key':key,'text':[p]})
            continue

        if not key and p.text.startswith('Тропарі'):
            key = f'тропарі_{service}'
            matrix.append({'key':key,'text':[p]})
            continue

        if not key and p.text.startswith('Після Вечірні'):
            key = f'тропар_вкінці_{service}'
            matrix.append({'key':key,'text':[p]})
            continue

        if not key and p.text.startswith('Перший сідальний'):
            key = f'сідальний1'
            matrix.append({'key':key,'text':[p]})
            continue
        
        if not key and p.text.startswith('Другий сідальний'):
            key = f'сідальний2'
            matrix.append({'key':key,'text':[p]})
            continue


        #Іпакої (неділі)
        if not key and p.text.startswith('Іпакой'):
            key = f'іпакой'
            matrix.append({'key':key,'text':[p]})
            continue



        #Величання
        if not key and p.text.startswith('Величання'):
            key = f'величання'
            matrix.append({'key':key,'text':[p]})
            continue

        #сідальний_пополієлейний
        if not key and p.text.startswith('Сідальний'):
            key = f'сідальний_пополієлейний'
            matrix.append({'key':key,'text':[p]})
            continue

        #"І щоб удостоїтися" + Євангеліє
        if not key and p.text.startswith('І щоб удостоїтися'):
            key = f'і_щоб_удостоїтися_євангеліє'
            matrix.append({'key': key, 'text': [p]})
            continue


        if not key and p.text.startswith('Степенна пісня'):
            key = f'степенна'
            matrix.append({'key':key,'text':[p]})
            continue

        #post_50
        if not key and p.text.startswith('Після 50-го псалма'):
            key = f'стихири_після_50'
            matrix.append({'key': key, 'text': [p]})
            continue



        if not key and p.text.startswith('Канон'):
            key = f'канон'
            matrix.append({'key':key,'text':[p]})
            continue


        if not key and p.text.startswith('Пісня 9'):
            key = f'пісня9'
            matrix.append({'key':key,'text':[p]})
            continue

        if not key and p.text.startswith('Світильний'):
            key = f'світильний'
            matrix.append({'key':key,'text':[p]})
            continue
    
        if not key and p.text==('Тропар'):
            key = f'воскресний_тропар'
            matrix.append({'key':key,'text':[p]})
            continue

        #!!!!!!!! Чи це працюватиме, з послідовними змінними частинами???


        if key:
            matrix[-1]['text'].append(p)

    return matrix







list_of_lines_that_end_parts = ['Світло тихе',
                                'Вхід', 
                                'Сподоби, Господи', #кінець прокімена седмичного дня
                                'Читання', #кінець святкового дня
                                'Господеві помолімся.', #кінець прокімена неділі
                                'Литійні молитви', #кінець литійних стихир
                                'Пісня Симеона', #кінець стиховні вечірні
                                "Великий відпуст", #кінець тропарів вечірні
                                "Благословення хлібів", #кінець тропарів на вечірні з литією
                                "Утреня", #кінець вечірні
                                "УТРЕНЯ", #кінець вечірні
                                "Тропарі", #кінець початку утрені
                                "По тропарях читає читець чергову катизму, після якої диякон перед св. дверми виголошує:", #кінець тропарів 
                                'Відтак читець читає другу чергову катизму, диякон виголошує другу малу єктенію:',

                                'Мала єктенія', #також кінець тропарів
                                'Полієлей', #кінець другого сідального
                                'Ангельський хор', #кінець другого сідального в період без полієлею
                                'Степенна', #кінець іпакоя
                                'Євангеліє', #кінець степенної
                                'Господеві помолімся.', #кінець прокімна
                                #'Мала єктенія',
                                'Псалом 50',
                                'Канон', #кінець єктенії в свята / неділі / піст
                                'Пісня Богородиці',#кінець канону 1-8
                                'Пісня 9',#кінець Пісні Богородиці
                                'Достойно',
                                '«Свят Господь Бог',
                                'Псалом 148', #кінець світильного
                                'Хвалитні', #кінець світильного
                                "Велике славослов'я",  #кінець хвалитних стихир
                                'Після стиховні читаємо:', #кінець стиховні утрені седмичного дня
                                "Тропар", #кінець Великого славослів'я
                                "Єктенія усильного благання", #кінець воскресного тропаря, кінець веірнього прокімена
]


list



def dump_list_to_text_file(data_list):
    filename='test.txt'
    with open(filename, 'w', encoding='utf8') as f:
        for entry in data_list:
            f.write('\nkey:\n')
            f.write(f"{entry['key']}\n")
            f.write('text:\n')
            for paragraph in entry['text']:
                f.write(paragraph.text + '\n')
    print(f'Dumped result to {filename}')

if __name__ == "__main__":
    start_time = datetime.now()
    
    #print("Hello World!")
    path = r"drafts\\07_vu\\15-ВТ-Володимира_Великого.docx"
    doc = docx.Document(path)
    matrix = get_variable_parts_from_template(path)
    

    for item in matrix:
        print(item['key'])

    dump_list_to_text_file(matrix)
    #print(item['key'])

    end_time = datetime.now()
    elapsed_time = (end_time - start_time).total_seconds()
